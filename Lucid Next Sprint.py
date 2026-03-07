"""Generate Lucid Next Sprint planning doc."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

STEEL_BLUE = RGBColor(0x5B, 0x8D, 0xB8)
DARK = RGBColor(0x1a, 0x1d, 0x21)
GRAY = RGBColor(0x5a, 0x62, 0x70)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'none'))
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = STEEL_BLUE
    run.font.name = 'Calibri'
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    return p

def body(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ": ")
        r1.font.bold = True
        r1.font.color.rgb = DARK
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r2 = p.add_run(text)
        r2.font.color.rgb = GRAY
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.color.rgb = GRAY
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E4E8EC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def status_table(label, value, color='5B8DB8'):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(1.4)
    table.columns[1].width = Inches(4.4)
    c0, c1 = table.cell(0, 0), table.cell(0, 1)
    shade_cell(c0, color)
    shade_cell(c1, 'FFFFFF')
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(label)
    r0.font.bold = True
    r0.font.color.rgb = WHITE
    r0.font.size = Pt(10)
    r0.font.name = 'Calibri'
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(value)
    r1.font.size = Pt(10)
    r1.font.color.rgb = DARK
    r1.font.name = 'Calibri'
    doc.add_paragraph()

# ── COVER ──────────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(8)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("Lucid — Next Sprint Plan")
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = DARK
r.font.name = 'Calibri'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Prepared March 5, 2026  ·  For implementation week of March 6")
rs.font.size = Pt(11)
rs.font.italic = True
rs.font.color.rgb = GRAY
rs.font.name = 'Calibri'

divider()

# ── SECTION 1: BUG FIXES & IMMEDIATE CHANGES ──────────────────────────────
h1("Part 1: Bug Fixes & App Improvements")
body("These three changes go into the existing Lucid app immediately when the weekly limit resets.")

# Bug 1
h2("1. Auto-Reset Stuck Voice Analyzer Timer")
h3("Problem")
body("The voice analyzer timer can get stuck at a partial reading (e.g., 12/30 seconds) for over an hour without advancing. This poisons subsequent analysis — data from 60+ minutes ago should not contribute to \"now\" readings.")
h3("Expected Behavior")
bullet("If the analyzer timer has not advanced for more than 60 minutes, automatically reset it to 0")
bullet("Log a silent reset event (no user-facing notification needed)")
bullet("Resume normal passive monitoring after the reset")
h3("Technical Note")
body("Check `last_progress_time` vs `current_time` in the voice analyzer loop. If delta > 3600s and timer is between 1–29s, trigger reset.", italic=True)

divider()

# Bug 2
h2("2. Morning Overlay: Full Screen (Not Top Half)")
h3("Problem")
body("The morning readiness overlay currently only occupies the top half of the Lucid app window.")
h3("Expected Behavior")
bullet("The morning overlay should fill the entire app window height")
bullet("No partial coverage — it should feel intentional and immersive, not clipped")
h3("Technical Note")
body("Check the overlay container CSS — likely a fixed height or max-height constraint needs to be removed or changed to 100%.", italic=True)

divider()

# Bug 3
h2("3. Morning Overlay: Label Above Score")
h3("Problem")
body("The morning readiness overlay shows the numeric score but lacks a clear label identifying what it represents.")
h3("Expected Behavior")
bullet('Add the text "Morning Readiness Score" directly above the numeric score')
bullet("Style: same font, smaller size (e.g., 14px), muted color (body text gray #5a6270)")
bullet("The score number itself stays large and prominent")

divider()

# Bug 4
h2("4. False 'All Rings Closed' Notifications During Meetings")
h3("Problem")
body("The 'All rings have closed' notification fires repeatedly even when the user is silent — e.g., listening to someone else speak in a meeting. This means the speaker gate is letting through audio from other people and counting it toward ring completion.")
h3("Expected Behavior")
bullet("Ring completion should only trigger from verified user speech")
bullet("Audio from other speakers (colleagues, etc.) should be rejected by the speaker gate before it contributes to any ring or metric")
bullet("If the speaker gate is working correctly, this bug should reduce once Features 6 and 7 (voice profile enhancement + anti-momentum rejection) are implemented — but investigate independently")
h3("Technical Note")
body("Check whether ring-closing logic runs before or after speaker gate filtering. If rings are being credited from unverified audio segments, move the gate check upstream so only gate-passed audio contributes to ring progress. Also verify the notification deduplication — if the same ring-close event fires multiple times, add a cooldown (e.g., no repeat within 30 minutes).", italic=True)

divider()

# Bug 5
h2("5. Remove Time Capsule Section")
h3("What to remove")
bullet("Delete the check_time_capsules method from insight_engine.py")
bullet("Remove time_capsules DB methods (add_time_capsule, get_time_capsules) from database.py")
bullet("Remove the TIME CAPSULE card from index.html (the <!-- Time Capsule (Feature #9) --> block)")
bullet("Remove the /api/capsules route from routes.py")
bullet("Remove any capsule-related CSS from engagement.css")
h3("Technical Note")
body("Five files touch this feature: insight_engine.py, database.py, index.html, routes.py, engagement.css. Remove in that order and verify the app loads without errors before rebuilding.", italic=True)

divider()

# ── SECTION 2: NEW FEATURES ────────────────────────────────────────────────
h1("Part 2: New Features to Build")
body('These features came out of the March 2026 external audit and concept image session. Reference images are in the Gemini conversation at gemini.google.com/app/3e10a65caef714ba')

# Feature 1
h2("Feature 1: Adaptive Notification Timing")
body("Best concept: Image #3 (analytics bar chart showing hourly response rates with 2-3 PM zone highlighted).")
h3("What it does")
bullet("Lucid learns which time windows you are most likely to respond to a check-in notification")
bullet("Uses a rolling personal model — not a generic schedule")
bullet("Sends notifications during your peak receptivity window instead of at fixed times")
h3("Key UI elements to build")
bullet("Settings toggle: \"Adaptive timing: ON\" with steel blue switch")
bullet("Mini weekly heatmap in settings showing your learned pattern (Mon–Sun × 8AM–8PM)")
bullet("Analytics screen: horizontal bar chart of response rate by hour, your peak zone highlighted")
h3("Implementation notes")
body("Store notification open timestamps. Build a simple hourly histogram per user. Prefer 7-day rolling window. Pick the top 2-hour window for daily notification delivery.", italic=True)

divider()

# Feature 2
h2("Feature 2: Streak Insurance (Resilience Day)")
body("Best concept: Image #3 (the 'Streak protected ✓' confirmation screen).")
h3("What it does")
bullet("Users get 1 free \"resilience day\" per week — a streak skip that doesn't break their streak")
bullet("If they miss a day, a notification fires: \"Protect your streak?\" with two options")
bullet("After using a resilience day, a confirmation screen shows the streak intact with a shield badge")
h3("Key UI elements to build")
bullet("Streak screen: large streak number + shield badge + \"1 resilience day available\" text")
bullet("Push notification with \"Protect streak\" (blue) and \"Let it reset\" (gray) action buttons")
bullet("Confirmation screen: streak number + green \"Streak protected ✓\" + weekly budget bar (1/1 used)")
h3("Implementation notes")
body("Track resilience_days_used per week (reset Monday). If missed_day detected and resilience_days_used < 1, queue the protection notification. Cap at 1 use per week.", italic=True)

divider()

# Feature 3
h2("Feature 3: Voice Season (90-Day Arc Narrative)")
body("Best concept: Image #1 (the full-screen 'Voice Season 1 Complete' card with 3-phase arc).")
h3("What it does")
bullet("Frames each 90-day period as a \"Voice Season\" with three phases: Discovery (days 1-30), Patterns (days 31-60), Prediction (days 61-90)")
bullet("Delivers a narrative arc — users feel they are building toward something, not just logging indefinitely")
bullet("At season completion: a celebratory full-screen card with stats (total readings, key discoveries)")
h3("Key UI elements to build")
bullet("Home screen widget: 'Day 47 of 90' progress bar + phase labels (completed/active/locked)")
bullet("Phase transition notification: 'You're entering the Patterns phase'")
bullet("Season complete card: dark background, Playfair Display title, 3-phase arc graphic, stats row")
bullet("Season 2 auto-starts after Season 1 completes")
h3("Implementation notes")
body("Season start = first recording date. Store season_number and season_start_date in DB. Phase = floor((day - 1) / 30). Prediction phase unlocks after day 60.", italic=True)

divider()

# Feature 4
h2("Feature 4: Voice Wellness Report (Therapist/Physician PDF Export)")
body("Best concept: Image #1 (clean clinical cover page with title, date range, patient name field).")
h3("What it does")
bullet("Generates a PDF report covering the last 90 days of voice data")
bullet("Designed to be shared with a therapist, physician, or employee assistance program")
bullet("Contains: cover page, 90-day stress trend chart, Grove themes table, key annotations")
h3("Key UI elements to build")
bullet("'Export Report' button in settings or history view")
bullet("Cover page: 'Voice Wellness Report', date range, 'Prepared by: Lucid Voice Monitor'")
bullet("Page 2: 90-day line chart (stress over time) with labeled annotations (vacations, milestones)")
bullet("Page 3: Grove Themes Analysis table (theme, frequency, trend arrow)")
bullet("Footer note on all pages: 'For clinical review. Extracted from voice pattern analysis.'")
h3("Implementation notes")
body("Use ReportLab or WeasyPrint for PDF generation. Pull last 90 days from SQLite. Grove themes = top keywords from daily_entries. Export to ~/Desktop or user-chosen path.", italic=True)

divider()

# Feature 5
h2("Feature 5: Weekly Wrapped")
body("Best concept: full-screen summary card — dark background, day circles, key stats, share option.")
h3("What it does")
bullet("Every Friday (or after 7 recorded days), generates a weekly summary card")
bullet("Shows: calmest day, avg stress vs prior week, streak, top Grove theme")
bullet("Shareable as a private card (send to one trusted contact) or save to camera roll")
h3("Key UI elements to build")
bullet("Full-screen card: 'Your Week in Voice', 5 day circles (colored for recorded days), key stats")
bullet("Shareable social card: square format, minimal waveform, week number, caption line")
bullet("Share sheet: 'Share privately' (to trusted contact) + 'Save to camera roll'")
h3("Implementation notes")
body("Trigger on Fridays at 6 PM, or when 7 days of recordings complete. Pull week's readings from DB. Calculate: avg_stress, stress_delta_vs_last_week, calmest_day, grove_theme_top.", italic=True)

divider()

# Feature 6
h2("Feature 6: Morning Recording → Voice Profile Enhancement")
body("Use the daily morning voice profile recording to continuously refine the user's speaker profile, improving accuracy over time.")
h3("What it does")
bullet("After each morning enrollment recording, extract the speaker embedding and use it to update the stored voice profile")
bullet("Keeps the profile calibrated to your current voice (accounts for seasonal changes, mic changes, etc.)")
bullet("Prevents the profile from drifting — each morning is a small recalibration")
h3("Implementation notes")
body("After morning recording: extract embedding → compare to stored profile (cosine similarity). If similarity > 0.7, merge: 90% old + 10% new embedding. If similarity unexpectedly low (illness, different mic), log warning but do NOT update. Persist updated embedding to profile store.", italic=True)

divider()

# Feature 7
h2("Feature 7: Anomaly Rejection — Anti-Momentum for Isolated Speaker Acceptances")
body("Mirror of the existing momentum recovery logic, but in reverse. Rejects isolated acceptances surrounded by consecutive rejections (e.g., a colleague briefly matching the threshold in a meeting).")
h3("What it does")
bullet("Tracks a rolling window of speaker verification verdicts (accept/reject)")
bullet("If a single acceptance is sandwiched by ≥2 consecutive rejections on both sides, retroactively marks it as rejected")
bullet("Prevents false data from colleagues or background voices from polluting your analysis")
h3("Logic")
bullet("Pattern to catch: [reject, reject, ..., accept, reject, reject, ...] where the acceptance is isolated")
bullet("Inverse of momentum recovery, which catches: [accept, accept, ..., reject, accept, accept, ...]")
bullet("Configurable: min consecutive rejections required on each side (default: 2)")
h3("Implementation notes")
body("Add anti_momentum_window config param (default: 2). In speaker gate logic, maintain rolling verdict buffer. After each verdict, check if the previous acceptance is now isolated by ≥N rejections on each side. If so, flip that verdict and exclude the segment from analysis. Log when anti-momentum fires for debugging.", italic=True)

divider()

# ── SECTION 3: REFERENCE ──────────────────────────────────────────────────
h1("Reference")
body("All 33 concept images are in the Gemini conversation. Scroll to find them by feature name in the prompt text.")
status_table("Gemini conversation", "gemini.google.com/app/3e10a65caef714ba")
status_table("Concept images folder", "/Users/zacharypoll/Desktop/Documents/Claude Code/Lucid Business Documents/feature-concepts/")
status_table("Audit report", "Lucid External Audit Report.docx (same folder)")

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = footer.add_run("Lucid  ·  'Clarity through voice.'  ·  March 2026")
rf.font.size = Pt(9)
rf.font.italic = True
rf.font.color.rgb = GRAY
rf.font.name = 'Calibri'

output = "/Users/zacharypoll/Desktop/Lucid Next Sprint.docx"
doc.save(output)
print(f"Saved: {output}")
