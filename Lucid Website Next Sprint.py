"""Generate Lucid Website Next Sprint planning doc."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

STEEL_BLUE = RGBColor(0x5B, 0x8D, 0xB8)
DARK = RGBColor(0x1a, 0x1d, 0x21)
GRAY = RGBColor(0x5a, 0x62, 0x70)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = STEEL_BLUE
    run.font.name = 'Calibri'
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    return p

def body(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ": ")
        r1.font.bold = True
        r1.font.color.rgb = DARK
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r2 = p.add_run(text)
        r2.font.color.rgb = GRAY
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.color.rgb = GRAY
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E4E8EC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def status_table(label, value, color='5B8DB8'):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.4)
    c0, c1 = table.cell(0, 0), table.cell(0, 1)
    shade_cell(c0, color)
    shade_cell(c1, 'FFFFFF')
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(label)
    r0.font.bold = True
    r0.font.color.rgb = WHITE
    r0.font.size = Pt(10)
    r0.font.name = 'Calibri'
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(value)
    r1.font.size = Pt(10)
    r1.font.color.rgb = DARK
    r1.font.name = 'Calibri'
    doc.add_paragraph()

# ── COVER ──────────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(8)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("Lucid Website — v22 & v23 Sprint Plan")
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = DARK
r.font.name = 'Calibri'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Prepared March 5, 2026  ·  For implementation week of March 6")
rs.font.size = Pt(11)
rs.font.italic = True
rs.font.color.rgb = GRAY
rs.font.name = 'Calibri'

divider()

# ── OVERVIEW ──────────────────────────────────────────────────────────────
h1("Overview & Two-Phase Plan")
body("This sprint produces two new website versions. Complete Phase 1 fully before starting Phase 2.")

h3("Phase 1 — Create v22 (section/content changes)")
body("Copy v21-ultimate.html → v22-ultimate.html. Apply all section reordering, removals, label edits, new components, and carousel. v21 is never modified.")

h3("Phase 2 — Create v23 (hero video upgrade)")
body("Copy the completed v22-ultimate.html → v23-ultimate.html. Apply only the hero video background change and spokesperson video removal. v22 is not touched again.")

status_table("Source (read-only)", ".../actual-websites/v21-ultimate.html")
status_table("Phase 1 output", ".../actual-websites/v22-ultimate.html")
status_table("Phase 2 output", ".../actual-websites/v23-ultimate.html")
status_table("Hero video", "/Users/zacharypoll/Desktop/lucid-hero-v2-clean.mp4")
status_table("Components dir", ".../lucid-website/components/")

body("Full paths: /Users/zacharypoll/Desktop/Documents/Claude Code/Lucid/lucid-website/actual-websites/", italic=True)

divider()

# ── TARGET SECTION ORDER ────────────────────────────────────────────────────
h1("Target Section Order (v22 and v23)")
body("After all Phase 1 edits, both v22 and v23 should flow in this exact order:")
bullet("1. Hero — 'Your Mind Is 9.8 Years Behind Your Body'")
bullet("2. Trust Strip — 9.8 years | 40M | 1x/year | 25 sec (moved up)")
bullet("3. 'Your Voice Has Been Leaking Your Mental State Your Entire Life' (moved up)")
bullet("4. System Comparison — 'The System You're In vs. The System You Deserve' (unchanged)")
bullet("5. Readiness — 'What Continuous Monitoring Actually Looks Like' (unchanged)")
bullet("6. [NEW] Your Mental Biomarkers (embedded from lucid-biomarker-cards.html)")
bullet("7. [NEW] Carousel — Recovery Pulse + Stress Throughout Your Day (replaces features section)")
bullet("8. How It Works and all remaining sections (unchanged)")
body("v23 is identical to v22 except the hero section uses a full-width background video instead of the current design.", italic=True)

divider()

# ── PHASE 1 ────────────────────────────────────────────────────────────────
h1("PHASE 1: Create v22-ultimate.html")

h2("Step 0: Duplicate v21 → v22")
body("Before making any changes, copy v21-ultimate.html to v22-ultimate.html in the same directory. All Phase 1 edits go into v22 only. v21 is never touched.", italic=True)

divider()

h2("Step 1: Remove the Viral Stat Section")
h3("What to delete")
body("Find and delete the entire section with class 'viral-stat'. Look for the comment '<!-- ============ VIRAL STAT (NEW - v14) ============ -->'.")
h3("HTML identifier")
body('<section class="viral-stat"> ... </section>', italic=True)
h3("Confirm deleted")
bullet("The quote 'The average American waits 9.8 years between first symptoms...' is gone")
bullet("The PHQ-9 / 1999 reference is gone")
bullet("The Wang et al. / Kroenke et al. citation line is gone")

divider()

h2("Step 2: Move Trust Strip to Immediately After Hero")
h3("Current position")
body("After deleting the viral stat, the trust strip will naturally be second. Confirm it sits directly after the hero's closing </section> tag with nothing between them.")
h3("If it is not in position 2")
body("Cut the entire trust strip section and paste it immediately after </section> of the hero, before any other section.", italic=True)
h3("HTML identifier")
body('<section class="trust-strip"> ... </section>', italic=True)

divider()

h2("Step 3: Change the '9.8 years' Label in the Trust Strip")
h3("What to change")
body("Inside the trust strip, find the trust-item that shows '9.8 years'. Its label reads 'Average Delay to Diagnosis'. Change ONLY that label.")
h3("Find")
body('<div class="trust-label">Average Delay to Diagnosis</div>', italic=True)
h3("Replace with")
body('<div class="trust-label">Average Delay to Start Improvement</div>', italic=True)
h3("Leave unchanged")
bullet("The '9.8 years' number stays the same")
bullet("The other three items (40M, 1x/year, 25 sec) are untouched")

divider()

h2("Step 4: Move 'Your Voice Has Been Leaking' to Section 3")
h3("Current position")
body("This section is currently the 5th section — after System Comparison. It has class 'secret-section'.")
h3("Target position")
body("Cut this entire section and paste it AFTER the trust strip and BEFORE the system comparison section.")
h3("HTML identifier")
body('<section class="secret-section section" id="science"> ... </section>', italic=True)
h3("Resulting order after this step")
bullet("Hero → Trust Strip → Secret ('Your Voice Has Been Leaking') → System Comparison → Readiness")

divider()

h2("Step 5: Remove the Features Section")
h3("What to delete")
body("Find and delete the entire features section. Look for the comment '<!-- ============ HEALTH FEATURES (Tabbed) ============ -->'.")
h3("HTML identifier")
body('<section class="features-section section" id="features"> ... </section>', italic=True)
h3("Confirm deleted")
bullet("Heading 'What Used to Require a Clinical Referral. Now It's on Your Mac.' is gone")
bullet("Feature tabs (Daily Detection, Early Warning, Raw Biomarkers, On Your Machine) are gone")

divider()

h2("Step 6: Embed 'Your Mental Biomarkers' Section After Readiness")
h3("Placement")
body("Insert the new section AFTER the readiness section (class 'readiness-section') and BEFORE the new carousel section.")
h3("Source file")
body("/Users/zacharypoll/Desktop/Documents/Claude Code/Lucid/lucid-website/components/lucid-biomarker-cards.html", italic=True)
h3("How to embed")
bullet("Open lucid-biomarker-cards.html — read its <style> block and <body> content in full")
bullet("Copy its CSS into v22's <style> section (add after existing styles; watch for class name conflicts)")
bullet("Wrap the body HTML in a <section class='section'> using a background color that matches the component's design")
bullet("The component already contains the 'Your Mental Biomarkers' heading — do not add a duplicate wrapper heading")
h3("Confirm")
bullet("Section visible between Readiness and Carousel when v22 is opened in a browser")
bullet("Cards render correctly with no CSS conflicts with surrounding sections")

divider()

h2("Step 7: Build the Recovery Pulse / Stress Timeline Carousel")
h3("Placement")
body("This section goes where the features section was — after Mental Biomarkers and before How It Works.")
h3("Source files")
bullet("Recovery Pulse: .../lucid-website/components/lucid-recovery-pulse.html")
bullet("Stress Throughout Your Day: .../lucid-website/components/lucid-stress-timeline.html")

h3("Carousel structure")
bullet("Full-width section with a short heading (e.g. 'See Your Data In Action')")
bullet("Left arrow (<) and right arrow (>) buttons flanking the component display area, vertically centered")
bullet("One component visible at a time — CSS fade or slide transition between them")
bullet("Two dot indicators below: active dot #5B8DB8 (steel blue), inactive #e4e8ec")

h3("Hover-to-animate behavior")
bullet("Default state: all CSS animations PAUSED (animation-play-state: paused) — shows a static preview")
bullet("On mouseenter: animations resume (animation-play-state: running) — charts draw, waveforms pulse, numbers count up")
bullet("On mouseleave: animations pause and reset so the next hover triggers a fresh animation (remove and re-add animation classes)")

h3("Embedding the components")
bullet("Read each component's HTML source — copy its <style> block into v22's <style> section")
bullet("Place each component's body HTML inside a .carousel-slide div with data-slide='0' and data-slide='1'")
bullet("Inline JS at bottom of <body>: tracks currentSlide (0 or 1), switches slides on arrow click, updates dots")
bullet("On mouseenter of active slide: resume animations. On mouseleave: pause and reset.")

h3("Arrow button styling")
bullet("48px × 48px, border-radius: 50%")
bullet("Default arrow color: #9aa2ae — hover: #5B8DB8")
bullet("Background: transparent or #f0f2f4")

h3("Confirm")
bullet("Right arrow: Recovery Pulse → Stress Timeline. Left arrow: back.")
bullet("Dot indicators update on each switch")
bullet("Hover triggers animation; mouse-out resets to static preview")
bullet("Both components display correctly at 1440px viewport width")

divider()

h2("Step 8: Add FAQ Item to Technology Section")
h3("What to add")
body("In the FAQ or Technology section of the page, add the following new question-and-answer entry:")
bullet("Q: What kind of application is this?")
bullet("A: This is a desktop application, not a mobile application. We utilize advanced machine learning algorithms to provide you with the most accurate and comprehensive data about yourself possible — data your phone isn't capable of processing locally.")
h3("Placement")
body("Add this as the first item in the Technology category of the FAQ section, before any existing technology questions.", italic=True)
h3("Confirm")
bullet("The new Q&A is visible in the Technology section of v22")
bullet("Answer text matches the approved copy exactly")

divider()

h2("Phase 1 Complete — Verify v22 Before Continuing")
h3("Open v22-ultimate.html in a browser and confirm all of the following")
bullet("Section order: Hero → Trust Strip → 'Your Voice Has Been Leaking' → System Comparison → Readiness → Mental Biomarkers → Carousel → How It Works")
bullet("Viral stat block is completely absent")
bullet("Features section (tabbed) is completely absent")
bullet("Trust strip shows: 9.8 years / Average Delay to Start Improvement")
bullet("Mental Biomarkers section renders correctly")
bullet("Carousel arrows work; hover-to-animate works on both slides")
bullet("v21 is untouched. Do not proceed to Phase 2 until v22 is verified.")

divider()

# ── PHASE 2 ────────────────────────────────────────────────────────────────
h1("PHASE 2: Create v23-ultimate.html")

h2("Step 0: Duplicate v22 → v23")
body("Copy the verified v22-ultimate.html to v23-ultimate.html in the same directory. All Phase 2 changes go into v23 only. v22 is never touched again.", italic=True)

divider()

h2("Phase 2 — Change 1: Hero Full-Width Video Background")
h3("Current hero (inherited from v22)")
body("Two-column layout: left = headline + CTAs, right = spokesperson video. Background is a dark CSS gradient.")
h3("New design")
body("The entire hero section background becomes the lucid-hero-v2-clean.mp4 video playing full-width. Text overlays it centered.")
h3("Implementation")
bullet("Copy the video: cp /Users/zacharypoll/Desktop/lucid-hero-v2-clean.mp4 .../actual-websites/images/lucid-hero-v2-clean.mp4")
bullet("Add a <video> as the FIRST child of .hero: src='images/lucid-hero-v2-clean.mp4', autoplay, muted, loop, playsinline")
bullet("Style the video: position:absolute; top:0; left:0; width:100%; height:100%; object-fit:cover; z-index:0")
bullet("Set .hero { position:relative; overflow:hidden; }")
bullet(".hero-inner and .hero-content: position:relative; z-index:1")
bullet("Add a dark overlay for readability: ::before or overlay div with background:rgba(0,0,0,0.45)")
bullet("Switch hero layout to single centered column: remove two-column grid; .hero-content: text-align:center; max-width:720px; margin:0 auto")
h3("Confirm")
bullet("Background video autoplays muted on page load, fills the full hero area")
bullet("Headline and CTAs are legible over the video")

divider()

h2("Phase 2 — Change 2: Remove Spokesperson Video")
h3("What to remove")
body("Delete the entire .hero-image div and all its contents — the spokesperson video and the play button.")
h3("HTML to delete")
body('<div class="hero-image fade-in fade-in-delay-2"> ... </div>', italic=True)
body("Contains: <video id='hero-video' src='images/privacy-clip-1.mp4' ...> and <button id='hero-play-btn' ...>", italic=True)
h3("After removing")
bullet("Only the text content (headline, description, CTAs, fine print) remains in the hero")
bullet("No split-column layout — single centered column over the background video")

divider()

h2("Phase 2 Complete — Final Verification")
h3("Open v23-ultimate.html in a browser and confirm")
bullet("Hero background video (lucid-hero-v2-clean.mp4) autoplays muted on load")
bullet("Headline, description, and CTAs are visible and readable over the video")
bullet("No spokesperson video visible anywhere on the page")
bullet("All Phase 1 changes are still intact (section order, carousel, biomarkers, etc.)")
bullet("Three files exist in actual-websites/: v21-ultimate.html (unchanged), v22-ultimate.html (content changes), v23-ultimate.html (video hero)")

divider()

# ── REFERENCE ────────────────────────────────────────────────────────────────
h1("Reference")
status_table("v21 (untouched source)", ".../actual-websites/v21-ultimate.html")
status_table("v22 (content changes)", ".../actual-websites/v22-ultimate.html")
status_table("v23 (video hero)", ".../actual-websites/v23-ultimate.html")
status_table("Hero video source", "/Users/zacharypoll/Desktop/lucid-hero-v2-clean.mp4")
status_table("Copy video to", ".../actual-websites/images/lucid-hero-v2-clean.mp4")
status_table("Biomarker component", ".../lucid-website/components/lucid-biomarker-cards.html")
status_table("Recovery Pulse", ".../lucid-website/components/lucid-recovery-pulse.html")
status_table("Stress Timeline", ".../lucid-website/components/lucid-stress-timeline.html")

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = footer.add_run("Lucid  ·  'Clarity through voice.'  ·  March 2026  ·  Website v22 + v23 Sprint")
rf.font.size = Pt(9)
rf.font.italic = True
rf.font.color.rgb = GRAY
rf.font.name = 'Calibri'

output = "/Users/zacharypoll/Desktop/Lucid Website Next Sprint.docx"
doc.save(output)
print(f"Saved: {output}")
