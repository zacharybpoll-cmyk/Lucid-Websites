#!/usr/bin/env python3
"""
Lucid Software Architecture Audit Report Generator
Produces a professional .docx report with charts and detailed findings.
"""

import io
import os
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Brand Colors ──────────────────────────────────────────────────────────
STEEL_BLUE = RGBColor(0x5B, 0x8D, 0xB8)
DARK_TEXT = RGBColor(0x1A, 0x1D, 0x21)
BODY_TEXT = RGBColor(0x5A, 0x62, 0x70)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF8, 0xF9, 0xFA)
RED_ACCENT = RGBColor(0xE0, 0x4B, 0x4B)
AMBER_ACCENT = RGBColor(0xF5, 0xA6, 0x23)
GREEN_ACCENT = RGBColor(0x4C, 0xAF, 0x50)

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Lucid_Architecture_Audit_2026-03.docx")

# ── Scoring Data ──────────────────────────────────────────────────────────
CATEGORIES = [
    {"name": "Architecture & Structure",       "weight": 0.18, "score": 6.5},
    {"name": "Code Quality",                   "weight": 0.14, "score": 5.8},
    {"name": "Security & Data Privacy",        "weight": 0.16, "score": 7.5},
    {"name": "Performance & Efficiency",       "weight": 0.12, "score": 5.0},
    {"name": "Testing & Reliability",          "weight": 0.14, "score": 5.5},
    {"name": "Maintainability & Tech Debt",    "weight": 0.10, "score": 5.5},
    {"name": "Accessibility & UX Engineering", "weight": 0.08, "score": 4.5},
    {"name": "Documentation & Dev Experience", "weight": 0.08, "score": 6.0},
]
OVERALL_SCORE = 5.94

# ── Subcategory Data ──────────────────────────────────────────────────────
SUBCATEGORIES = {
    "Architecture & Structure": [
        ("1A", "Separation of Concerns", 7, "Clean 3-tier: Electron shell -> FastAPI backend -> SQLite"),
        ("1B", "File Organization & Modularity", 5, "routes.py monolith (1,887 LOC); app.js monolith (3,949 LOC); router migration underway"),
        ("1C", "API Design & Contract Clarity", 8, "Pydantic schemas on all 61+ endpoints, custom exception hierarchy, dependency injection"),
        ("1D", "Data Layer Architecture", 5, "Centralized Database class with RLock, WAL mode, auto-backup; no migration system"),
    ],
    "Code Quality": [
        ("2A", "Readability & Naming", 6, "Good utility naming; FE-015 style inline comments; CSS semantic names"),
        ("2B", "Complexity Management", 5, "app.js: 42 intervals/timeouts, 24 event listeners with only 3 cleanup references"),
        ("2C", "Error Handling", 5, "96 except Exception across 20 Python files; .catch(() => {}) x6 in frontend"),
        ("2D", "Consistency", 6, "Bridge pattern centralizes API; canonical color palette; some global duplication"),
    ],
    "Security & Data Privacy": [
        ("3A", "Input Validation", 8, "Pydantic validation on all endpoints; sanitizeHTML() in frontend"),
        ("3B", "Injection Prevention", 9, "All parameterized queries confirmed; zero SQL injection vectors"),
        ("3C", "Data Privacy & PII", 7, "Crash log PII redaction; all processing local; SQLite unencrypted on disk"),
        ("3D", "IPC & Network Security", 6, "CORS localhost-only; minimal preload (7 lines); SSRF prevention; overly permissive CORS headers"),
    ],
    "Performance & Efficiency": [
        ("4A", "Resource Usage", 5, "Targets: idle <500MB, active <1.5GB — high for a menubar app"),
        ("4B", "Polling & Animation Efficiency", 4, "42 setInterval/setTimeout unmanaged; 16+ RAF calls; carousel loops when paused"),
        ("4C", "Startup Performance", 5, "Splash <2s, ready <15s; ML model loading is bottleneck"),
        ("4D", "Caching Strategy", 6, "Daily summary cache; config caching; no broader strategy"),
    ],
    "Testing & Reliability": [
        ("5A", "Test Coverage", 5, "4,305 LOC tests across 10 files; ~22% coverage ratio"),
        ("5B", "Test Quality", 6, "Good breadth: DB, acoustic, engagement, patterns, insights, speaker, notifications, API"),
        ("5C", "Error Resilience", 7, "Circuit breaker on LLM; retry-after headers; require_initialized(); auto-backup"),
        ("5D", "CI/CD & Automation", 3, "No CI/CD pipeline; manual verification only"),
    ],
    "Maintainability & Tech Debt": [
        ("6A", "Dead Code & Unused Files", 5, "timeline.css (empty), audio-processor.js (stub); duplicate globals"),
        ("6B", "Technical Debt Awareness", 7, "Hardcoded values documented; tasks/ directory maintained"),
        ("6C", "Dependency Health", 5, "package-lock.json exists; audit commands documented; no enforced schedule"),
    ],
    "Accessibility & UX Engineering": [
        ("7A", "ARIA & Semantic HTML", 3, "Only 13 aria-/role= attributes across 1,424-line SPA; critically low"),
        ("7B", "Keyboard Navigation", 3, "No keyboard handlers or focus management; no skip links"),
        ("7C", "Visual Design Quality", 7, "Strong brand identity; canonical zone colors; CSS variables for theming"),
    ],
    "Documentation & Dev Experience": [
        ("8A", "Inline Documentation", 6, "Module-level docstrings on Python files; not systematic"),
        ("8B", "API Documentation", 7, "FastAPI auto-generates OpenAPI; 61 endpoints catalogued"),
        ("8C", "Onboarding & Setup", 5, "CLAUDE.md setup docs aimed at AI dev, not human onboarding"),
    ],
}

# ── Detailed Findings ─────────────────────────────────────────────────────
FINDINGS = {
    "Architecture & Structure": {
        "key_findings": [
            ("Clean 3-Tier Separation", "STRENGTH", "Electron shell communicates with FastAPI backend via localhost HTTP; SQLite provides persistence. Each tier has clear responsibility boundaries."),
            ("Frontend Monolith", "HIGH RISK", "app.js at 3,949 LOC is a single-file SPA controller managing 10+ views, 42 timers, and 24 event listeners. No module system, no view lifecycle management."),
            ("Router Migration In Progress", "MODERATE", "Legacy routes.py (1,887 LOC) being decomposed into 10 modular router files under api/routers/. Migration ~60% complete."),
            ("No Database Migration System", "MODERATE", "Schema evolution via CREATE TABLE IF NOT EXISTS. No version tracking, no rollback capability. Risky for production updates."),
            ("Strong API Contracts", "STRENGTH", "Pydantic validation on all 61+ endpoints. Custom exception hierarchy (LucidError -> DatabaseNotReady, ModelNotLoaded, ServiceNotReady) with retry-after headers."),
        ],
        "evidence": [
            "app.js:1-3949 — Single file with 42 setInterval/setTimeout calls",
            "routes.py:1-1887 — Legacy monolithic route handler",
            "api/routers/ — 10 modular router files (analysis, dashboard, lab, readings, reports, settings, speaker, voice_profile, health, analytics)",
            "database.py — RLock thread safety, WAL mode, 13 tables",
        ],
    },
    "Code Quality": {
        "key_findings": [
            ("Timer Accumulation Risk", "HIGH RISK", "42 intervals/timeouts in app.js with no centralized timer management. No cleanup on view switches — timers accumulate as users navigate."),
            ("Event Listener Leaks", "HIGH RISK", "24 addEventListener calls with only 3 cleanup/dispose references. Listeners accumulate on repeated view switches, causing memory leaks."),
            ("Broad Exception Handling", "MODERATE", "96 except Exception occurrences across 20 Python files. While preventing crashes, this masks specific errors and complicates debugging."),
            ("Silent Frontend Failures", "MODERATE", ".catch(() => {}) appears 6 times — errors are swallowed silently with no logging or user feedback."),
            ("Bridge Pattern", "STRENGTH", "All frontend-backend communication flows through bridge.js (255 LOC), centralizing API calls and providing a consistent interface."),
        ],
        "evidence": [
            "app.js — grep -c 'setInterval\\|setTimeout' returns 42",
            "app.js — grep -c 'addEventListener' returns 24; grep -c 'removeEventListener\\|dispose\\|cleanup' returns 3",
            "Backend Python files — grep -rc 'except Exception' returns 96 across 20 files",
            "Frontend JS — grep -rc '.catch(() => {})' returns 6 occurrences",
        ],
    },
    "Security & Data Privacy": {
        "key_findings": [
            ("Zero SQL Injection Vectors", "STRENGTH", "All database queries use parameterized statements. Manual audit of database.py (1,528 LOC) confirmed no string interpolation in queries."),
            ("SSRF Prevention", "STRENGTH", "_validate_webhook_url() blocks localhost, loopback, and private IP ranges before making outbound requests."),
            ("API Token Security", "STRENGTH", "Per-session token via secrets.token_urlsafe(32), localhost-only binding, validated on every request."),
            ("Crash Log PII Redaction", "STRENGTH", "Sanitization removes embeddings, home paths, and base64 strings from crash reports before logging."),
            ("Unencrypted SQLite", "LOW RISK", "Database file on disk is unencrypted. Contains voice-derived health metrics. Acceptable for local-only app but limits enterprise deployment."),
            ("Overly Permissive CORS", "LOW RISK", "CORS allows * methods and * headers. While localhost-only, should be restricted to specific methods."),
        ],
        "evidence": [
            "database.py — All 150+ SQL statements use ? placeholders",
            "webhook_manager.py — _validate_webhook_url() with IP validation",
            "main.py — secrets.token_urlsafe(32) for session tokens",
            "logging_config.py — PII sanitization regex patterns",
        ],
    },
    "Performance & Efficiency": {
        "key_findings": [
            ("Unmanaged Animation Loops", "HIGH RISK", "16+ requestAnimationFrame calls in app.js; carousel RAF loop runs indefinitely even when the carousel is paused or off-screen."),
            ("Overlapping Polling", "MODERATE", "Multiple concurrent polling strategies: status, data, voice scan, trends. No coordination or cancellation on navigation."),
            ("High Memory Footprint", "MODERATE", "Idle target <500MB, active <1.5GB for a menubar app. ML model stack (PyTorch + Whisper + Silero VAD + SpeechBrain + DAM) drives baseline."),
            ("Startup Bottleneck", "MODERATE", "15-second ready target driven by sequential ML model loading. No lazy loading or parallel initialization."),
            ("Daily Summary Cache", "STRENGTH", "Effective caching prevents redundant recomputation of daily aggregates."),
        ],
        "evidence": [
            "app.js — 16 requestAnimationFrame calls; carousel animation never cancelled",
            "app.js — 5 separate setInterval polling loops (status, data, voice, trends, config)",
            "CLAUDE.md — documented memory targets: idle <500MB, active <1.5GB",
            "analysis_orchestrator.py — sequential model loading in startup()",
        ],
    },
    "Testing & Reliability": {
        "key_findings": [
            ("Reasonable Test Breadth", "STRENGTH", "10 test files covering database, acoustic features, engagement, pattern detection, insight engine, speaker gate, notifications, API, and audit."),
            ("Circuit Breaker Pattern", "STRENGTH", "LLM calls (insight engine) protected by circuit breaker with exponential backoff. Prevents cascade failures from API outages."),
            ("Low Coverage Ratio", "MODERATE", "4,305 LOC tests for ~32,000 LOC source = ~22% ratio. Frontend has zero test coverage."),
            ("No CI/CD Pipeline", "HIGH RISK", "Tests exist but run manually. No automated gate preventing regressions from reaching production."),
            ("No E2E Tests", "MODERATE", "No integration tests covering the full Electron -> FastAPI -> SQLite path. Frontend behavior untested."),
        ],
        "evidence": [
            "10 test files in python/tests/ totaling 4,305 LOC",
            "insight_engine.py — CircuitBreaker class with state management",
            "No .github/workflows/, no CI config files found",
            "No Playwright/Cypress/Selenium test files found",
        ],
    },
    "Maintainability & Tech Debt": {
        "key_findings": [
            ("Dead Files Present", "LOW RISK", "timeline.css is 1 line (empty). audio-processor.js (26 lines) is an unused AudioWorklet stub. engagement.js has empty renderMilestones()."),
            ("Duplicate Global State", "MODERATE", "Global variables duplicated between routes.py and dependencies.py. State synchronization is fragile."),
            ("Debt Awareness Good", "STRENGTH", "Hardcoded config values documented in CLAUDE.md. tasks/ directory tracks known issues. Technical debt is visible, not hidden."),
            ("Repetitive Error Handling", "MODERATE", "20+ endpoints use nearly identical try-except-return pattern. Opportunity for decorator-based error handling."),
        ],
        "evidence": [
            "timeline.css — 1 line, empty file",
            "audio-processor.js — 26 lines, unused AudioWorklet stub",
            "routes.py + dependencies.py — duplicate global variable declarations",
            "CLAUDE.md — hardcoded values section with 15+ documented constants",
        ],
    },
    "Accessibility & UX Engineering": {
        "key_findings": [
            ("Critically Low ARIA Coverage", "HIGH RISK", "Only 13 aria-/role= attributes across the entire 1,424-line SPA. For an app with 10+ views, modals, and interactive components, this blocks inclusive access."),
            ("No Keyboard Navigation", "HIGH RISK", "No keyboard event handlers for navigation. No focus management on view switches. No skip links. Users cannot operate the app without a mouse."),
            ("No Modal Accessibility", "MODERATE", "Settings modal lacks role='dialog', aria-modal, and focus trapping. Screen readers cannot distinguish modal from page content."),
            ("Strong Visual Design", "STRENGTH", "Canonical zone color palette with CSS variables. Steel blue accent, Inter/Playfair Display fonts. Consistent visual identity across all views."),
        ],
        "evidence": [
            "index.html — grep -c 'aria-\\|role=' returns 13",
            "app.js — grep -c 'keydown\\|keyup\\|keypress\\|focus' returns 0 navigation-related handlers",
            "index.html — settings modal div has no role or aria attributes",
            "main.css — CSS custom properties for theming (--steel-blue, --dark-text, etc.)",
        ],
    },
    "Documentation & Dev Experience": {
        "key_findings": [
            ("Auto-Generated API Docs", "STRENGTH", "FastAPI provides automatic OpenAPI/Swagger documentation for all 61+ endpoints. Supplemented by shared/lucid-api-reference.md."),
            ("AI-Targeted Setup Docs", "MODERATE", "CLAUDE.md documents deploy procedures, test commands, and performance benchmarks — but is optimized for AI-assisted development, not human developer onboarding."),
            ("Inconsistent Docstrings", "LOW RISK", "Module-level docstrings present on Python files. Function-level documentation inconsistent. No JSDoc coverage on frontend."),
            ("Good Operational Docs", "STRENGTH", "Performance targets, memory budgets, and architecture decisions documented. tasks/ directory maintained."),
        ],
        "evidence": [
            "FastAPI auto-generates /docs and /openapi.json endpoints",
            "CLAUDE.md — 200+ lines of deployment and configuration documentation",
            "Python modules — module-level docstrings present; function-level ~40% coverage",
            "No JSDoc comments found in frontend JS files",
        ],
    },
}

# ── Recommendations Per Category ──────────────────────────────────────────
RECOMMENDATIONS = {
    "Architecture & Structure": [
        ("Split app.js into view modules", "Extract each view (dashboard, settings, lab, reports, etc.) into its own module with a shared view lifecycle manager that handles init/destroy. Target: no single JS file over 500 LOC.", "P0"),
        ("Implement database migrations", "Integrate Alembic with a migrations/ directory. Create an initial migration from the current schema. All future schema changes go through versioned migrations with up/down support.", "P1"),
        ("Complete router migration", "Move the remaining ~40 endpoints from routes.py into the modular api/routers/ files. Delete routes.py once empty. Target: routes.py at 0 LOC.", "P1"),
        ("Add view lifecycle management", "Create a ViewManager class that tracks the active view and calls init()/destroy() on transitions. All timers, listeners, and RAF loops register through ViewManager for automatic cleanup.", "P0"),
    ],
    "Code Quality": [
        ("Create TimerRegistry", "Build a centralized timer management class that wraps setInterval/setTimeout. Each timer gets a view scope. On view switch, all timers for the departing view are automatically cleared.", "P0"),
        ("Fix event listener leaks", "Audit all 24 addEventListener calls. Add corresponding removeEventListener in a dispose() function for each view. Use AbortController where supported for clean bulk removal.", "P0"),
        ("Replace broad exception handlers", "Audit all 96 except Exception occurrences. Replace with specific exception types (DatabaseNotReady, ModelNotLoaded, etc.). Log the original exception with traceback for debugging.", "P1"),
        ("Eliminate silent catches", "Replace all 6 .catch(() => {}) occurrences with .catch(err => console.error('context:', err)) at minimum. Consider user-visible error states for critical operations.", "P1"),
    ],
    "Security & Data Privacy": [
        ("Restrict CORS headers", "Replace allow_methods=['*'] and allow_headers=['*'] with explicit lists. Only GET, POST, PUT, DELETE are needed. Only Content-Type and Authorization headers are required.", "P2"),
        ("Consider SQLite encryption", "Evaluate SQLCipher for database-at-rest encryption. While local-only mitigates risk, encrypted storage strengthens the security posture for enterprise and compliance scenarios.", "P2"),
        ("Audit innerHTML usage", "Replace innerHTML assignments in grove.js, correlation.js, and reports.js with textContent or sanitized DOM construction. Even local apps should follow defense-in-depth.", "P1"),
        ("Token display security", "Change settings UI to show API token only once on generation, then mask it. Add a 'regenerate' button instead of persistent display.", "P2"),
    ],
    "Performance & Efficiency": [
        ("Fix carousel RAF loop", "Add a guard to cancelAnimationFrame when the carousel is paused or not visible. Save the RAF ID and cancel it in the pause handler.", "P0"),
        ("Consolidate polling strategies", "Replace the 5 separate setInterval polling loops with a single orchestrated polling manager that coordinates data fetching. Cancel all polls on view exit.", "P1"),
        ("Lazy-load ML models", "Load models on-demand rather than all at startup. Whisper and DAM can load in the background after the UI is ready. Use a loading indicator for first analysis.", "P1"),
        ("Profile memory usage", "Instrument the app with periodic memory snapshots (process.memoryUsage()). Log to a rolling buffer. Identify the top 3 memory growth patterns over a 24-hour session.", "P2"),
    ],
    "Testing & Reliability": [
        ("Set up CI/CD", "Add a GitHub Actions workflow that runs pytest on every push and PR. Block merges on test failure. Add a badge to README.", "P0"),
        ("Add frontend test coverage", "Start with critical paths: bridge.js API contract tests, view initialization tests, timer cleanup verification. Use Jest or Vitest.", "P1"),
        ("Add E2E integration tests", "Implement Playwright tests covering: onboarding flow, first recording, dashboard data display, report generation, settings changes.", "P2"),
        ("Increase backend coverage to 50%", "Prioritize untested areas: thread safety under contention, audio capture stream handling, edge cases in score_engine.py, and webhook delivery.", "P1"),
    ],
    "Maintainability & Tech Debt": [
        ("Remove dead files", "Delete timeline.css (empty), audio-processor.js (unused stub). Remove empty renderMilestones() from engagement.js. Clean FIRST_LIGHT_TASKS array.", "P0"),
        ("Resolve global duplication", "Consolidate duplicate global variables between routes.py and dependencies.py into a single source of truth in dependencies.py.", "P1"),
        ("Extract error handling decorator", "Create a @handle_errors decorator for FastAPI routes that replaces the repeated try/except/return pattern across 20+ endpoints.", "P1"),
        ("Consolidate CSS files", "Merge the 12 CSS files into a structured system: base.css (reset + variables), components.css (shared), and per-view files. Delete empty/near-empty files.", "P2"),
    ],
    "Accessibility & UX Engineering": [
        ("ARIA landmark audit", "Add role='main', role='navigation', role='complementary' to the primary layout regions. Add aria-label to all interactive elements. Target: 100+ ARIA attributes.", "P0"),
        ("Implement keyboard navigation", "Add keydown handlers for Tab, Enter, Escape, and arrow keys. Implement focus management on view switches (focus first interactive element). Add skip links.", "P0"),
        ("Modal accessibility", "Add role='dialog', aria-modal='true', and focus trapping to the settings modal. Return focus to trigger element on close.", "P1"),
        ("Canvas accessibility", "Add aria-label and role='img' to all waveform/visualization canvases. Provide text alternatives for chart data.", "P1"),
    ],
    "Documentation & Dev Experience": [
        ("Create human onboarding guide", "Write a CONTRIBUTING.md covering: prerequisites, setup steps, development workflow, testing, and deployment. Separate from CLAUDE.md (AI-targeted).", "P1"),
        ("Add architecture decision records", "Document key decisions (why Electron, why SQLite over Postgres, why local-only) in an ADR directory. Prevents re-litigating settled decisions.", "P2"),
        ("Systematic JSDoc coverage", "Add JSDoc to all exported functions in frontend JS files. Start with bridge.js, app.js public functions, and gauges.js.", "P2"),
        ("Add inline architecture diagrams", "Create Mermaid diagrams for: data flow (audio -> analysis -> storage), API topology, and frontend view hierarchy.", "P2"),
    ],
}

# ── Risk Heatmap Items ────────────────────────────────────────────────────
# (name, likelihood 1-5, impact 1-5)
RISK_ITEMS = [
    ("app.js monolith / timer leaks",       4, 4),
    ("No CI/CD pipeline",                   5, 4),
    ("Accessibility gaps",                   3, 5),
    ("Event listener memory leaks",          4, 3),
    ("No database migrations",               3, 4),
    ("Silent error swallowing",              3, 3),
    ("RAF loops running when paused",        4, 2),
    ("routes.py incomplete migration",       2, 3),
    ("Unencrypted SQLite",                   2, 3),
    ("CORS overly permissive",               1, 2),
]

# ── LOC Data ──────────────────────────────────────────────────────────────
BACKEND_LOC = [
    ("database.py", 1528), ("insight_engine.py", 1169), ("analysis_orchestrator.py", 762),
    ("score_engine.py", 645), ("report_generator.py", 613), ("engagement.py", 606),
    ("acoustic_features.py", 575), ("linguistic_features.py", 572), ("speaker_verifier.py", 540),
    ("notifications.py", 528), ("pattern_detector.py", 479), ("analytics.py", 295),
    ("speaker_gate.py", 266), ("active_assessment.py", 250), ("dam_analyzer.py", 247),
    ("audio_capture.py", 200), ("speech_buffer.py", 148), ("meeting_detector.py", 127),
    ("baseline_calibrator.py", 115), ("supabase_client.py", 113), ("burnout_calculator.py", 112),
    ("plda_scorer.py", 110), ("webhook_manager.py", 98), ("vad_processor.py", 92),
    ("overlap_detector.py", 74), ("logging_config.py", 59),
]

API_LOC = [
    ("routes.py (legacy)", 1887), ("analysis.py", 571), ("dashboard.py", 559),
    ("lab.py", 428), ("schemas.py", 298), ("readings.py", 269), ("reports.py", 261),
    ("voice_profile.py", 243), ("settings.py", 193), ("health.py", 172),
    ("speaker.py", 118), ("dependencies.py", 82), ("analytics.py", 33),
    ("exceptions.py", 31), ("constants.py", 7),
]

FRONTEND_JS_LOC = [
    ("app.js", 3949), ("gauges.js", 1165), ("onboarding.js", 921), ("lab.js", 857),
    ("active_assessment.js", 701), ("sculptor.js", 550), ("trends.js", 550),
    ("reports.js", 494), ("anxiety_timeline.js", 414), ("correlation.js", 268),
    ("bridge.js", 255), ("timeline.js", 215), ("layout.js", 147), ("grove.js", 130),
    ("voice_profile.js", 111), ("engagement.js", 50), ("audio-processor.js", 26),
]

FRONTEND_CSS_LOC = [
    ("main.css", 3038), ("engagement.css", 2404), ("lab.css", 826),
    ("active_assessment.css", 744), ("reports.css", 428), ("sculptor.css", 167),
    ("trends.css", 120), ("voice_profile.css", 118), ("anxiety_timeline.css", 76),
    ("correlation.css", 40), ("gauges.css", 10), ("timeline.css", 1),
]

TEST_LOC = [
    ("test_audit_additions.py", 596), ("test_pattern_detector.py", 543),
    ("test_insight_engine.py", 533), ("test_engagement.py", 531),
    ("test_notifications.py", 515), ("test_database.py", 511),
    ("test_api.py", 389), ("test_speaker_gate.py", 384),
    ("test_acoustic.py", 218), ("conftest.py", 85),
]

PYTHON_DEPS = [
    ("fastapi", "0.128.8", "Web framework"),
    ("uvicorn", "0.40.0", "ASGI server"),
    ("pydantic", "2.12.5", "Data validation"),
    ("torch", "2.10.0", "ML framework"),
    ("torchaudio", "2.10.0", "Audio processing"),
    ("silero-vad", "6.2.0", "Voice activity detection"),
    ("sounddevice", "0.5.5", "Audio capture"),
    ("soundfile", "0.13.1", "Audio I/O"),
    ("librosa", "0.11.0", "Audio analysis"),
    ("transformers", ">=4.41.0,<5.0.0", "ML models"),
    ("speechbrain", "1.0.3", "Speaker verification"),
    ("openai-whisper", ">=20240930", "Speech-to-text"),
    ("spacy", ">=3.7.0", "NLP"),
    ("sentence-transformers", "3.3.1", "Embeddings"),
    ("numpy", "2.3.5", "Numerical computing"),
    ("scipy", "1.17.0", "Scientific computing"),
    ("psutil", "7.2.2", "System monitoring"),
    ("httpx", ">=0.27.0", "HTTP client"),
    ("reportlab", ">=4.0.0", "PDF generation"),
    ("numba", "0.63.1", "JIT compilation"),
    ("plyer", "2.1.0", "Native notifications"),
    ("accelerate", "latest", "ML acceleration"),
]

NPM_DEPS = [
    ("electron", "^35.0.0", "Desktop framework"),
    ("@electron/packager", "^19.0.5", "Build tool"),
    ("electron-updater", "^6.8.3", "Auto-update"),
]


# ══════════════════════════════════════════════════════════════════════════
# Helpers
# ══════════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="5B8DB8"):
    """Create a professional table with header styling."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.color.rgb = DARK_TEXT
            if r % 2 == 1:
                set_cell_shading(cell, "F0F2F4")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def score_color(score):
    """Return color based on score value."""
    if score >= 7:
        return "#4CAF50"
    elif score >= 5:
        return "#F5A623"
    else:
        return "#E04B4B"


def score_color_rgb(score):
    if score >= 7:
        return GREEN_ACCENT
    elif score >= 5:
        return AMBER_ACCENT
    else:
        return RED_ACCENT


def risk_label(text):
    if text == "STRENGTH":
        return GREEN_ACCENT
    elif text == "HIGH RISK":
        return RED_ACCENT
    elif text == "MODERATE":
        return AMBER_ACCENT
    else:
        return BODY_TEXT


def add_heading_styled(doc, text, level=1):
    """Add a heading with brand styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_TEXT
    return h


def chart_to_image(fig, dpi=150):
    """Convert matplotlib figure to bytes for docx embedding."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════
# Chart Generation
# ══════════════════════════════════════════════════════════════════════════

def create_radar_chart():
    """Generate a radar chart of category scores."""
    labels = [c["name"].replace(" & ", "\n& ").replace("UX ", "UX\n") for c in CATEGORIES]
    scores = [c["score"] for c in CATEGORIES]

    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    scores_plot = scores + [scores[0]]
    angles += [angles[0]]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, scores_plot, color="#5B8DB8", alpha=0.15)
    ax.plot(angles, scores_plot, color="#5B8DB8", linewidth=2.5, marker="o", markersize=7)

    # Reference lines
    ref_10 = [10] * (len(labels) + 1)
    ref_5 = [5] * (len(labels) + 1)
    ax.plot(angles, ref_10, color="#E4E8EC", linewidth=0.8, linestyle="--")
    ax.plot(angles, ref_5, color="#E4E8EC", linewidth=0.8, linestyle="--")

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, size=8, color="#5A6270")
    ax.set_ylim(0, 10)
    ax.set_yticks([2, 4, 6, 8, 10])
    ax.set_yticklabels(["2", "4", "6", "8", "10"], size=7, color="#999")
    ax.spines["polar"].set_color("#E4E8EC")
    ax.grid(color="#E4E8EC", linewidth=0.5)
    ax.set_title(f"Overall Score: {OVERALL_SCORE}/10", size=14, color="#1A1D21",
                 fontweight="bold", pad=20)
    fig.tight_layout()
    return chart_to_image(fig)


def create_bar_chart():
    """Generate a horizontal bar chart of weighted scores."""
    names = [c["name"] for c in CATEGORIES]
    scores = [c["score"] for c in CATEGORIES]
    weighted = [c["score"] * c["weight"] for c in CATEGORIES]

    fig, ax = plt.subplots(figsize=(7, 4))
    y_pos = np.arange(len(names))
    colors = [score_color(s) for s in scores]

    bars = ax.barh(y_pos, scores, color=colors, height=0.6, alpha=0.85)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(names, size=8.5, color="#5A6270")
    ax.set_xlim(0, 10)
    ax.set_xlabel("Score (out of 10)", size=9, color="#5A6270")
    ax.axvline(x=OVERALL_SCORE, color="#5B8DB8", linestyle="--", linewidth=1.5, label=f"Overall: {OVERALL_SCORE}")

    for i, (s, w) in enumerate(zip(scores, weighted)):
        ax.text(s + 0.15, i, f"{s}", va="center", ha="left", size=8.5, color="#1A1D21", fontweight="bold")

    ax.legend(loc="lower right", fontsize=8)
    ax.invert_yaxis()
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#E4E8EC")
    ax.spines["bottom"].set_color("#E4E8EC")
    ax.tick_params(colors="#999")
    fig.tight_layout()
    return chart_to_image(fig)


def create_risk_heatmap():
    """Generate a risk heatmap (likelihood vs impact)."""
    fig, ax = plt.subplots(figsize=(7, 5.5))

    # Background gradient zones
    for x in range(1, 6):
        for y in range(1, 6):
            risk = x * y
            if risk >= 12:
                color = "#FFCDD2"
            elif risk >= 6:
                color = "#FFF9C4"
            else:
                color = "#C8E6C9"
            ax.add_patch(plt.Rectangle((x - 0.5, y - 0.5), 1, 1, facecolor=color, edgecolor="#E4E8EC", linewidth=0.5))

    # Plot risk items with jitter to avoid overlap
    np.random.seed(42)
    for name, likelihood, impact in RISK_ITEMS:
        jx = likelihood + np.random.uniform(-0.2, 0.2)
        jy = impact + np.random.uniform(-0.2, 0.2)
        risk = likelihood * impact
        if risk >= 12:
            color = "#E04B4B"
        elif risk >= 6:
            color = "#F5A623"
        else:
            color = "#4CAF50"
        ax.plot(jx, jy, "o", color=color, markersize=10, markeredgecolor="white", markeredgewidth=1.5)
        ax.annotate(name, (jx, jy), textcoords="offset points", xytext=(8, 4),
                    fontsize=6.5, color="#1A1D21", fontweight="medium")

    ax.set_xlim(0.5, 5.5)
    ax.set_ylim(0.5, 5.5)
    ax.set_xticks([1, 2, 3, 4, 5])
    ax.set_xticklabels(["Very Low", "Low", "Medium", "High", "Very High"], size=8)
    ax.set_yticks([1, 2, 3, 4, 5])
    ax.set_yticklabels(["Very Low", "Low", "Medium", "High", "Very High"], size=8)
    ax.set_xlabel("Likelihood", size=10, color="#5A6270", fontweight="bold")
    ax.set_ylabel("Impact", size=10, color="#5A6270", fontweight="bold")
    ax.set_title("Risk Heat Map", size=13, color="#1A1D21", fontweight="bold", pad=12)

    legend_elements = [
        mpatches.Patch(facecolor="#FFCDD2", edgecolor="#ccc", label="Critical (>=12)"),
        mpatches.Patch(facecolor="#FFF9C4", edgecolor="#ccc", label="Moderate (6-11)"),
        mpatches.Patch(facecolor="#C8E6C9", edgecolor="#ccc", label="Low (<6)"),
    ]
    ax.legend(handles=legend_elements, loc="upper left", fontsize=7, framealpha=0.9)

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return chart_to_image(fig)


# ══════════════════════════════════════════════════════════════════════════
# Document Sections
# ══════════════════════════════════════════════════════════════════════════

def create_cover_page(doc):
    """Create the cover page."""
    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LUCID")
    run.font.size = Pt(42)
    run.font.color.rgb = STEEL_BLUE
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Software Architecture Audit")
    run.font.size = Pt(24)
    run.font.color.rgb = DARK_TEXT

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("March 2026")
    run.font.size = Pt(14)
    run.font.color.rgb = BODY_TEXT

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 1.0  |  Confidential")
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_TEXT
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared for Lucid Leadership Team")
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_TEXT

    doc.add_page_break()


def create_executive_summary(doc):
    """Create the executive summary page."""
    add_heading_styled(doc, "Executive Summary", level=1)

    # Overall score callout
    p = doc.add_paragraph()
    run = p.add_run(f"Overall Architecture Score: {OVERALL_SCORE} / 10")
    run.font.size = Pt(18)
    run.font.color.rgb = STEEL_BLUE
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Adequate with notable strengths in security; gaps in performance and accessibility")
    run.font.size = Pt(11)
    run.font.color.rgb = BODY_TEXT
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    doc.add_paragraph(
        "This report presents a comprehensive architecture audit of the Lucid voice wellness application, "
        "a desktop Electron app with a FastAPI backend and SQLite persistence layer. The codebase comprises "
        "approximately 36,600 lines of source code across Python, JavaScript, CSS, and HTML. "
        "Eight categories were evaluated using a weighted scoring rubric, producing an overall score of "
        f"{OVERALL_SCORE}/10."
    )

    add_heading_styled(doc, "Top 3 Strengths", level=2)
    strengths = [
        ("Security (7.5/10)", "Zero SQL injection vectors. All parameterized queries. SSRF prevention, crash log PII redaction, localhost-only API tokens with per-session rotation."),
        ("Architecture (6.5/10)", "Clean 3-tier separation (Electron -> FastAPI -> SQLite). Pydantic validation on all 61+ endpoints. Custom exception hierarchy with dependency injection."),
        ("Analysis Pipeline", "Sophisticated ML pipeline (VAD -> Speaker Gate -> Acoustic -> Linguistic -> DAM -> Scoring) with circuit breakers and graceful degradation under failure conditions."),
    ]
    for title, desc in strengths:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = GREEN_ACCENT
        run = p.add_run(desc)
        run.font.size = Pt(10)
        run.font.color.rgb = BODY_TEXT

    add_heading_styled(doc, "Top 3 Risks", level=2)
    risks = [
        ("Frontend Monolith", "app.js at 3,949 LOC with 42 unmanaged timers, 24 event listeners with only 3 cleanup references. Memory leaks accumulate on view navigation."),
        ("No CI/CD Pipeline", "4,305 LOC of tests exist but run manually. No automated gate prevents regressions from reaching production builds."),
        ("Accessibility (4.5/10)", "Only 13 ARIA attributes across a 1,424-line SPA with 10+ views. No keyboard navigation. Blocks inclusive access for users with disabilities."),
    ]
    for title, desc in risks:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RED_ACCENT
        run = p.add_run(desc)
        run.font.size = Pt(10)
        run.font.color.rgb = BODY_TEXT

    add_heading_styled(doc, "Immediate Actions", level=2)
    actions = [
        "Set up CI/CD with automated pytest on every commit",
        "Split app.js into view-scoped modules with lifecycle management",
        "Add centralized timer registry with automatic cleanup on view switches",
    ]
    for action in actions:
        p = doc.add_paragraph(action, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()


def create_scoring_dashboard(doc):
    """Create the scoring dashboard with charts."""
    add_heading_styled(doc, "Scoring Dashboard", level=1)

    # Radar chart
    radar_img = create_radar_chart()
    doc.add_picture(radar_img, width=Inches(4.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Score table
    rows = []
    for c in CATEGORIES:
        weighted = round(c["score"] * c["weight"], 2)
        rows.append((c["name"], f"{int(c['weight']*100)}%", f"{c['score']}", f"{weighted}"))
    rows.append(("OVERALL", "100%", "", f"{OVERALL_SCORE}"))

    add_styled_table(doc, ["Category", "Weight", "Score", "Weighted"], rows,
                     col_widths=[3.0, 0.8, 0.8, 0.9])

    doc.add_paragraph()

    # Bar chart
    bar_img = create_bar_chart()
    doc.add_picture(bar_img, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def create_codebase_overview(doc):
    """Create the codebase overview section."""
    add_heading_styled(doc, "Codebase Overview", level=1)

    add_heading_styled(doc, "Architecture", level=2)
    doc.add_paragraph(
        "Lucid follows a 3-tier desktop architecture: an Electron shell provides the native macOS "
        "container and system tray integration; a FastAPI backend runs as a subprocess handling all "
        "business logic, ML inference, and data persistence; SQLite provides the storage layer with "
        "WAL mode for concurrent read access."
    )
    doc.add_paragraph(
        "The frontend is a single-page application (SPA) served as static files from FastAPI. "
        "Communication between Electron and the backend uses HTTP via a bridge.js abstraction layer. "
        "The Electron preload script is minimal (7 lines), exposing only a window.lucid namespace "
        "via contextBridge."
    )

    add_heading_styled(doc, "Tech Stack", level=2)
    stack_rows = [
        ("Electron", "^35.0.0", "Desktop container, system tray, auto-update"),
        ("FastAPI", "0.128.8", "Backend API, static file serving, OpenAPI docs"),
        ("SQLite", "3.x (WAL)", "Persistence, 13 tables, auto-backup"),
        ("PyTorch", "2.10.0", "ML inference (DAM, PLDA, embeddings)"),
        ("Whisper", ">=20240930", "Speech-to-text transcription"),
        ("Silero VAD", "6.2.0", "Voice activity detection"),
        ("SpeechBrain", "1.0.3", "Speaker verification (ECAPA-TDNN)"),
        ("spaCy", ">=3.7.0", "NLP / linguistic feature extraction"),
    ]
    add_styled_table(doc, ["Technology", "Version", "Role"], stack_rows,
                     col_widths=[1.5, 1.2, 3.5])

    doc.add_paragraph()

    add_heading_styled(doc, "Lines of Code Summary", level=2)
    loc_rows = [
        ("Python Backend", "27 modules", "10,324"),
        ("Python API Layer", "17 files", "5,153"),
        ("Python Tests", "10 files", "4,305"),
        ("JavaScript", "17 files", "9,712"),
        ("CSS", "12 files", "7,972"),
        ("HTML", "3 files", "~1,810"),
        ("Electron (main.js)", "1 file", "640"),
        ("Total", "87 files", "~36,600"),
    ]
    add_styled_table(doc, ["Layer", "Files", "LOC"], loc_rows,
                     col_widths=[2.5, 1.2, 1.2])

    doc.add_paragraph()

    add_heading_styled(doc, "ML Model Stack", level=2)
    doc.add_paragraph(
        "The analysis pipeline processes audio through a multi-stage pipeline: Voice Activity Detection "
        "(Silero VAD) -> Speaker Gate (ECAPA-TDNN) -> Acoustic Feature Extraction (librosa + custom) -> "
        "Linguistic Analysis (Whisper + spaCy) -> Decision Audio Model (Kintsugi DAM, 700MB checkpoint) -> "
        "Score Engine. A circuit breaker protects LLM calls (insight generation) with exponential backoff."
    )

    doc.add_paragraph()

    add_heading_styled(doc, "Database Schema", level=2)
    doc.add_paragraph(
        "SQLite database with WAL (Write-Ahead Logging) mode for concurrent read access. "
        "Thread safety provided by RLock in the centralized Database class. Auto-backup on startup."
    )
    schema_rows = [
        ("readings", "Primary data", "Timestamped voice analysis readings (stress, energy, mood scores)"),
        ("daily_summaries", "Aggregation", "Pre-computed daily averages and trend data"),
        ("baselines", "Calibration", "Per-user voice baseline measurements for personalization"),
        ("tags", "Metadata", "User-applied tags for readings and time periods"),
        ("briefings", "Insights", "Generated daily/weekly wellness briefings"),
        ("echoes", "History", "Historical wellness snapshots for trend visualization"),
        ("compass_entries", "Tracking", "Wellness compass directional tracking data"),
        ("webhooks", "Integration", "Registered webhook URLs for external notifications"),
        ("notification_prefs", "Settings", "User notification preferences and quiet hours"),
        ("voice_profile", "Identity", "Speaker verification embeddings and thresholds"),
        ("self_assessments", "Validation", "User self-reported wellness for calibration"),
        ("goals", "Engagement", "User-defined wellness goals and progress"),
        ("active_sessions", "Runtime", "Currently active analysis sessions"),
    ]
    add_styled_table(doc, ["Table", "Category", "Purpose"], schema_rows,
                     col_widths=[1.5, 1.0, 3.7])

    doc.add_paragraph()

    add_heading_styled(doc, "Thread Safety Architecture", level=2)
    doc.add_paragraph(
        "The backend employs 12+ locks across critical subsystems to ensure thread safety in the "
        "concurrent FastAPI environment. Key patterns include:"
    )
    lock_items = [
        "database.py: RLock for all DB operations (allows reentrant access within same thread)",
        "analysis_orchestrator.py: _state_lock (RLock), _analysis_lock (Lock), meeting_lock (Lock) for state machine coordination",
        "speaker_gate.py: Lock for speaker verification state",
        "notifications.py: Lock for rate limiting (max 4/hour) and quiet hours enforcement",
        "Edge case handlers: EDGE-001 (120s analysis watchdog), EDGE-002 (DB integrity check on startup), EDGE-003 (2s audio reconnect debounce)",
    ]
    for item in lock_items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(9.5)

    doc.add_page_break()


def create_category_section(doc, cat_idx):
    """Create a detailed findings section for one category."""
    cat = CATEGORIES[cat_idx]
    name = cat["name"]
    score = cat["score"]

    add_heading_styled(doc, f"{cat_idx + 1}. {name}", level=1)

    # Score callout
    p = doc.add_paragraph()
    run = p.add_run(f"  Score: {score} / 10  ")
    run.font.size = Pt(16)
    run.font.color.rgb = score_color_rgb(score)
    run.bold = True

    # Subcategory table
    if name in SUBCATEGORIES:
        subcats = SUBCATEGORIES[name]
        rows = [(s[0], s[1], str(s[2]), s[3]) for s in subcats]
        add_styled_table(doc, ["ID", "Subcategory", "Score", "Notes"], rows,
                         col_widths=[0.5, 1.8, 0.6, 3.5])
        doc.add_paragraph()

    # Key findings
    if name in FINDINGS:
        findings = FINDINGS[name]

        add_heading_styled(doc, "Key Findings", level=2)
        for title, severity, description in findings["key_findings"]:
            p = doc.add_paragraph()
            # Severity badge
            run = p.add_run(f"[{severity}] ")
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = risk_label(severity)
            # Title
            run = p.add_run(f"{title}: ")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_TEXT
            # Description
            run = p.add_run(description)
            run.font.size = Pt(10)
            run.font.color.rgb = BODY_TEXT

        add_heading_styled(doc, "Evidence", level=2)
        for evidence in findings["evidence"]:
            p = doc.add_paragraph(evidence, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = BODY_TEXT
                run.italic = True

        # Recommendations
        if name in RECOMMENDATIONS:
            add_heading_styled(doc, "Recommendations", level=2)
            for rec_title, rec_detail, rec_priority in RECOMMENDATIONS[name]:
                p = doc.add_paragraph()
                run = p.add_run(f"[{rec_priority}] ")
                run.bold = True
                run.font.size = Pt(9)
                if rec_priority == "P0":
                    run.font.color.rgb = RED_ACCENT
                elif rec_priority == "P1":
                    run.font.color.rgb = AMBER_ACCENT
                else:
                    run.font.color.rgb = BODY_TEXT
                run = p.add_run(f"{rec_title} — ")
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = DARK_TEXT
                run = p.add_run(rec_detail)
                run.font.size = Pt(10)
                run.font.color.rgb = BODY_TEXT

    doc.add_page_break()


def create_risk_heatmap_section(doc):
    """Create the risk heatmap page."""
    add_heading_styled(doc, "Risk Heat Map", level=1)

    doc.add_paragraph(
        "The following heat map plots identified risks by likelihood of occurrence (x-axis) against "
        "potential impact on the product (y-axis). Items in the red zone require immediate attention; "
        "yellow items should be addressed within 90 days."
    )

    heatmap_img = create_risk_heatmap()
    doc.add_picture(heatmap_img, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Risk inventory table
    rows = []
    for name, likelihood, impact in RISK_ITEMS:
        risk_score = likelihood * impact
        if risk_score >= 12:
            level = "Critical"
        elif risk_score >= 6:
            level = "Moderate"
        else:
            level = "Low"
        rows.append((name, str(likelihood), str(impact), str(risk_score), level))

    rows.sort(key=lambda r: int(r[3]), reverse=True)
    add_styled_table(doc, ["Risk", "Likelihood", "Impact", "Score", "Level"], rows,
                     col_widths=[2.5, 0.8, 0.8, 0.6, 0.8])

    doc.add_page_break()


def create_roadmap(doc):
    """Create the improvement roadmap."""
    add_heading_styled(doc, "Prioritized Improvement Roadmap", level=1)

    phases = [
        ("Phase 1: Immediate (0-30 days)", "HIGH", [
            ("Set up CI/CD pipeline", "Add GitHub Actions or similar with pytest on every push. Block merges on test failure.", "Testing"),
            ("Split app.js into modules", "Extract view controllers into separate files (dashboard.js, settings.js, etc.) with a view lifecycle manager.", "Architecture"),
            ("Centralize timer management", "Create a TimerRegistry class that tracks all setInterval/setTimeout calls and auto-clears on view switch.", "Performance"),
            ("Fix event listener leaks", "Add removeEventListener calls on view switches. Implement a dispose() pattern for each view.", "Code Quality"),
            ("Remove dead files", "Delete timeline.css (empty), audio-processor.js (unused stub). Clean FIRST_LIGHT_TASKS.", "Maintainability"),
            ("Fix carousel RAF loop", "Cancel requestAnimationFrame when carousel is paused or view is not active.", "Performance"),
        ]),
        ("Phase 2: Short-term (30-90 days)", "MEDIUM", [
            ("Add database migration system", "Integrate Alembic for version-tracked schema migrations with rollback capability.", "Architecture"),
            ("Complete routes.py migration", "Move remaining endpoints from routes.py (1,887 LOC) into modular router files.", "Architecture"),
            ("Accessibility remediation", "Add ARIA landmarks, roles, and labels to all views. Implement keyboard navigation and focus management.", "Accessibility"),
            ("Structured error handling", "Replace bare except Exception with specific exception types. Add error context/logging.", "Code Quality"),
            ("Add database indices", "Profile slow queries and add indices on timestamp, date, and foreign key columns.", "Performance"),
            ("Systematic polling cancellation", "Cancel all polling intervals when navigating away from a view.", "Performance"),
        ]),
        ("Phase 3: Medium-term (90-180 days)", "LOW", [
            ("E2E test suite", "Add Playwright tests covering critical user flows (onboarding, recording, insights, reports).", "Testing"),
            ("Performance profiling", "Profile memory usage, identify leaks, optimize ML model loading (lazy load, parallel init).", "Performance"),
            ("Developer documentation", "Create human-targeted onboarding guide, architecture decision records, and contribution guidelines.", "Documentation"),
            ("Consolidate CSS", "Merge 12 CSS files into a structured system (base, components, views). Remove unused selectors.", "Maintainability"),
            ("Complete type coverage", "Bring Python type hints to 100%. Add JSDoc to all exported frontend functions.", "Code Quality"),
            ("Restrict CORS", "Limit CORS to specific HTTP methods and headers instead of wildcards.", "Security"),
        ]),
    ]

    for phase_title, priority, items in phases:
        add_heading_styled(doc, phase_title, level=2)

        rows = [(item[0], item[1], item[2]) for item in items]
        add_styled_table(doc, ["Action", "Details", "Category"], rows,
                         col_widths=[1.8, 3.2, 1.2])
        doc.add_paragraph()

    doc.add_page_break()


def create_appendix_a(doc):
    """Appendix A: File Inventory with LOC."""
    add_heading_styled(doc, "Appendix A: File Inventory", level=1)

    add_heading_styled(doc, "Python Backend Modules", level=2)
    rows = [(f[0], str(f[1])) for f in BACKEND_LOC]
    rows.append(("Total", str(sum(f[1] for f in BACKEND_LOC))))
    add_styled_table(doc, ["File", "LOC"], rows, col_widths=[3.0, 1.0])

    doc.add_paragraph()

    add_heading_styled(doc, "API Layer", level=2)
    rows = [(f[0], str(f[1])) for f in API_LOC]
    rows.append(("Total", str(sum(f[1] for f in API_LOC))))
    add_styled_table(doc, ["File", "LOC"], rows, col_widths=[3.0, 1.0])

    doc.add_paragraph()

    add_heading_styled(doc, "Frontend JavaScript", level=2)
    rows = [(f[0], str(f[1])) for f in FRONTEND_JS_LOC]
    rows.append(("Total", str(sum(f[1] for f in FRONTEND_JS_LOC))))
    add_styled_table(doc, ["File", "LOC"], rows, col_widths=[3.0, 1.0])

    doc.add_paragraph()

    add_heading_styled(doc, "Frontend CSS", level=2)
    rows = [(f[0], str(f[1])) for f in FRONTEND_CSS_LOC]
    rows.append(("Total", str(sum(f[1] for f in FRONTEND_CSS_LOC))))
    add_styled_table(doc, ["File", "LOC"], rows, col_widths=[3.0, 1.0])

    doc.add_paragraph()

    add_heading_styled(doc, "Test Suite", level=2)
    rows = [(f[0], str(f[1])) for f in TEST_LOC]
    rows.append(("Total", str(sum(f[1] for f in TEST_LOC))))
    add_styled_table(doc, ["File", "LOC"], rows, col_widths=[3.0, 1.0])

    doc.add_page_break()


def create_appendix_b(doc):
    """Appendix B: Dependency Inventory."""
    add_heading_styled(doc, "Appendix B: Dependency Inventory", level=1)

    add_heading_styled(doc, "Python Dependencies (48 packages)", level=2)
    doc.add_paragraph(
        "Key dependencies from requirements.txt. Full list includes transitive dependencies."
    )
    rows = [(d[0], d[1], d[2]) for d in PYTHON_DEPS]
    add_styled_table(doc, ["Package", "Version", "Purpose"], rows,
                     col_widths=[2.0, 1.2, 3.0])

    doc.add_paragraph()

    add_heading_styled(doc, "npm Dependencies", level=2)
    rows = [(d[0], d[1], d[2]) for d in NPM_DEPS]
    add_styled_table(doc, ["Package", "Version", "Purpose"], rows,
                     col_widths=[2.0, 1.2, 3.0])

    doc.add_paragraph()

    add_heading_styled(doc, "ML Models", level=2)
    model_rows = [
        ("Kintsugi DAM 3.1", "~700 MB", "Decision Audio Model for stress/affect detection"),
        ("Whisper base", "~140 MB", "OpenAI speech-to-text transcription"),
        ("Silero VAD", "~2 MB", "Voice activity detection"),
        ("ECAPA-TDNN", "~30 MB", "Speaker verification embeddings (SpeechBrain)"),
        ("en_core_web_sm", "~12 MB", "spaCy English NLP model"),
    ]
    add_styled_table(doc, ["Model", "Size", "Purpose"], model_rows,
                     col_widths=[2.0, 1.0, 3.2])

    doc.add_page_break()


def create_appendix_c(doc):
    """Appendix C: Methodology & Scoring Scale."""
    add_heading_styled(doc, "Appendix C: Methodology & Scoring Scale", level=1)

    add_heading_styled(doc, "Scoring Scale", level=2)
    scale_rows = [
        ("9-10", "Exceptional", "Industry-leading practices; minimal improvement needed"),
        ("7-8", "Good", "Solid practices with minor gaps; low risk"),
        ("5-6", "Adequate", "Functional but with notable gaps; moderate risk"),
        ("3-4", "Below Average", "Significant gaps requiring attention; high risk"),
        ("1-2", "Critical", "Fundamental issues blocking progress; immediate action required"),
    ]
    add_styled_table(doc, ["Range", "Rating", "Description"], scale_rows,
                     col_widths=[0.8, 1.2, 4.2])

    doc.add_paragraph()

    add_heading_styled(doc, "Methodology", level=2)
    doc.add_paragraph(
        "This audit was conducted through systematic manual review of the complete Lucid codebase "
        "(~36,600 LOC). The methodology included:"
    )
    methods = [
        "Static analysis of all Python, JavaScript, CSS, and HTML source files",
        "Grep-based pattern detection for security vulnerabilities, error handling patterns, and code smells",
        "Dependency inventory and version audit of all Python (pip) and Node (npm) packages",
        "Architecture mapping of module dependencies, API contracts, and data flow",
        "Accessibility audit against WCAG 2.1 Level AA guidelines",
        "Test coverage analysis comparing test LOC to source LOC by module",
        "Performance pattern review of timers, animation loops, polling strategies, and memory management",
        "Security review covering OWASP Top 10, input validation, injection prevention, and data privacy",
    ]
    for method in methods:
        doc.add_paragraph(method, style="List Bullet")

    doc.add_paragraph()

    add_heading_styled(doc, "Weight Rationale", level=2)
    doc.add_paragraph(
        "Category weights reflect the relative importance for a health-focused desktop application "
        "handling sensitive voice data:"
    )
    weight_rows = [
        ("Architecture & Structure", "18%", "Foundation; affects all other categories"),
        ("Security & Data Privacy", "16%", "Health data sensitivity; regulatory considerations"),
        ("Code Quality", "14%", "Daily developer productivity and bug prevention"),
        ("Testing & Reliability", "14%", "User trust; health data integrity"),
        ("Performance & Efficiency", "12%", "Desktop UX; always-on background processing"),
        ("Maintainability & Tech Debt", "10%", "Long-term velocity"),
        ("Accessibility & UX Engineering", "8%", "Inclusive access; regulatory compliance potential"),
        ("Documentation & Dev Experience", "8%", "Team scaling and onboarding"),
    ]
    add_styled_table(doc, ["Category", "Weight", "Rationale"], weight_rows,
                     col_widths=[2.5, 0.7, 3.0])

    doc.add_paragraph()

    add_heading_styled(doc, "Limitations", level=2)
    doc.add_paragraph(
        "This audit is based on static code review and does not include runtime profiling, "
        "penetration testing, or user testing. Scores reflect the state of the codebase as of "
        "March 2026. Dynamic behavior (memory leaks under load, startup time variance, etc.) "
        "would require instrumented runtime analysis for precise measurement."
    )


# ══════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = BODY_TEXT

    # Heading styles
    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = DARK_TEXT

    # Build document
    print("Building cover page...")
    create_cover_page(doc)

    print("Building executive summary...")
    create_executive_summary(doc)

    print("Building scoring dashboard...")
    create_scoring_dashboard(doc)

    print("Building codebase overview...")
    create_codebase_overview(doc)

    print("Building detailed findings (8 categories)...")
    for i in range(len(CATEGORIES)):
        print(f"  {i+1}. {CATEGORIES[i]['name']}...")
        create_category_section(doc, i)

    print("Building risk heat map...")
    create_risk_heatmap_section(doc)

    print("Building improvement roadmap...")
    create_roadmap(doc)

    print("Building appendices...")
    create_appendix_a(doc)
    create_appendix_b(doc)
    create_appendix_c(doc)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"\nReport saved to: {OUTPUT_PATH}")
    print("Done!")


if __name__ == "__main__":
    main()
