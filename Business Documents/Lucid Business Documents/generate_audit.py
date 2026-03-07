"""
Lucid External Audit Report Generator
Generates a comprehensive .docx audit report from 3 company perspectives.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "/Users/zacharypoll/Desktop/Documents/Claude Code/Lucid Business Documents/Lucid External Audit Report.docx"

# Lucid brand colors
STEEL_BLUE = RGBColor(0x5B, 0x8D, 0xB8)
DARK_TEXT = RGBColor(0x1A, 0x1D, 0x21)
BODY_TEXT = RGBColor(0x5A, 0x62, 0x70)
MID_GRAY = RGBColor(0xF0, 0xF2, 0xF4)
DETAIL_GRAY = RGBColor(0xE4, 0xE8, 0xEC)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = RGBColor(0xD6, 0xE6, 0xF4)


def set_cell_bg(cell, hex_color):
    """Set cell background color via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), kwargs.get(side, 'none'))
        border.set(qn('w:sz'), kwargs.get('sz', '4'))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get('color', '5B8DB8'))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_page_break(doc):
    doc.add_page_break()


def style_heading(para, level=1, color=None):
    para.clear()
    run = para.add_run()
    if level == 1:
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = DARK_TEXT if color is None else color
    elif level == 2:
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = STEEL_BLUE if color is None else color
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = DARK_TEXT if color is None else color
    return run


def add_h1(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = DARK_TEXT if color is None else color
    # Add bottom border via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '5B8DB8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = STEEL_BLUE
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DARK_TEXT
    return p


def add_body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run_bold = p.add_run(bold_prefix + ": ")
        run_bold.font.bold = True
        run_bold.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_rubric_table(doc, criteria):
    """3-column rubric table: #, Criterion, What It Measures"""
    table = doc.add_table(rows=1 + len(criteria), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    headers = ['#', 'Criterion', 'What It Measures']
    widths = [Inches(0.4), Inches(2.0), Inches(4.0)]
    for i, (cell, hdr) in enumerate(zip(table.rows[0].cells, headers)):
        cell.width = widths[i]
        set_cell_bg(cell, '5B8DB8')
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(hdr)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for row_idx, (num, criterion, measure) in enumerate(criteria):
        row = table.rows[row_idx + 1]
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        row.cells[2].width = widths[2]
        bg = 'F0F2F4' if row_idx % 2 == 0 else 'FFFFFF'
        for cell in row.cells:
            set_cell_bg(cell, bg)
        # Number
        p = row.cells[0].paragraphs[0]
        p.clear()
        r = p.add_run(str(num))
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = STEEL_BLUE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Criterion
        p = row.cells[1].paragraphs[0]
        p.clear()
        r = p.add_run(criterion)
        r.font.size = Pt(10)
        r.font.bold = True
        # Measure
        p = row.cells[2].paragraphs[0]
        p.clear()
        r = p.add_run(measure)
        r.font.size = Pt(10)
    doc.add_paragraph()


def add_score_card(doc, auditor, title, company, round_num, scores, labels, commentary):
    """Auditor bio + score table."""
    # Bio card
    bio_table = doc.add_table(rows=1, cols=1)
    bio_table.style = 'Table Grid'
    cell = bio_table.rows[0].cells[0]
    set_cell_bg(cell, 'D6E6F4')
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(f"{auditor}  |  {title}  |  Round {round_num}")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_TEXT
    doc.add_paragraph()

    add_body(doc, commentary)

    # Score table
    score_table = doc.add_table(rows=2, cols=5)
    score_table.style = 'Table Grid'
    score_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [Inches(1.6), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0)]

    # Header row
    header_cells = score_table.rows[0].cells
    header_labels = ['Criterion', labels[0], labels[1], labels[2], labels[3]]
    for i, (cell, lbl) in enumerate(zip(header_cells, header_labels)):
        cell.width = col_widths[i]
        set_cell_bg(cell, '5B8DB8')
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(lbl)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Score row
    score_cells = score_table.rows[1].cells
    total = sum(scores)
    score_vals = ['Score /25'] + [str(s) for s in scores]
    for i, (cell, val) in enumerate(zip(score_cells, score_vals)):
        cell.width = col_widths[i]
        set_cell_bg(cell, 'F0F2F4')
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(val)
        r.font.size = Pt(11)
        r.font.bold = (i > 0)
        r.font.color.rgb = STEEL_BLUE if i > 0 else BODY_TEXT
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Total score callout
    total_p = doc.add_paragraph()
    total_p.paragraph_format.space_before = Pt(6)
    total_p.paragraph_format.space_after = Pt(12)
    r = total_p.add_run(f"Total Score: {total}/100")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = STEEL_BLUE
    doc.add_paragraph()


def add_summary_table(doc, data):
    """Full 9-auditor summary table."""
    headers = ['Auditor', 'Company', 'Round', 'C1', 'C2', 'C3', 'C4', 'Total']
    col_widths = [Inches(1.5), Inches(1.3), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.7)]

    table = doc.add_table(rows=1 + len(data), cols=8)
    table.style = 'Table Grid'

    # Header
    for i, (cell, hdr) in enumerate(zip(table.rows[0].cells, headers)):
        cell.width = col_widths[i]
        set_cell_bg(cell, '5B8DB8')
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(hdr)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    company_colors = {
        'Instagram': 'FFF0F5',
        'Function Health': 'F0F8F0',
        'Oura Ring': 'F0F4FF',
    }
    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx + 1]
        bg = company_colors.get(row_data[1], 'FFFFFF')
        for i, (cell, val) in enumerate(zip(row.cells, row_data)):
            cell.width = col_widths[i]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.clear()
            is_total = (i == 7)
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.bold = is_total
            if is_total:
                r.font.color.rgb = STEEL_BLUE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()


def add_path_to_80(doc, auditor_name, intro_text, current_scores, improvements, criterion_labels):
    """Renders 'Path to 80/100' section with a styled improvement table."""
    # Section header as tinted box
    header_tbl = doc.add_table(rows=1, cols=1)
    header_tbl.style = 'Table Grid'
    hcell = header_tbl.rows[0].cells[0]
    set_cell_bg(hcell, 'D6E6F4')
    hp = hcell.paragraphs[0]
    hp.clear()
    hr = hp.add_run(f"Path to 80/100 — {auditor_name}")
    hr.font.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = DARK_TEXT
    doc.add_paragraph()

    # Intro sentence in auditor's voice
    add_body(doc, intro_text, italic=True)

    # Build table
    num_rows = 1 + len(improvements) + 1  # header + data + total
    tbl = doc.add_table(rows=num_rows, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [Inches(3.2), Inches(1.5), Inches(0.7), Inches(0.9)]

    # Header row
    header_labels_tbl = ['Improvement', 'Criterion', '+Points', 'New Score']
    for i, (cell, lbl) in enumerate(zip(tbl.rows[0].cells, header_labels_tbl)):
        cell.width = col_widths[i]
        set_cell_bg(cell, '5B8DB8')
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(lbl)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows — track running criterion scores
    running_scores = list(current_scores)
    for row_idx, imp in enumerate(improvements):
        running_scores[imp['criterion_idx']] += imp['points']
        row = tbl.rows[row_idx + 1]
        bg = 'F0F2F4' if row_idx % 2 == 0 else 'FFFFFF'
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_bg(cell, bg)
        # Improvement text
        p = row.cells[0].paragraphs[0]
        p.clear()
        r = p.add_run(imp['text'])
        r.font.size = Pt(10)
        # Criterion label
        p = row.cells[1].paragraphs[0]
        p.clear()
        r = p.add_run(criterion_labels[imp['criterion_idx']])
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = BODY_TEXT
        # Points gain
        p = row.cells[2].paragraphs[0]
        p.clear()
        r = p.add_run(f"+{imp['points']}")
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = STEEL_BLUE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # New criterion score
        p = row.cells[3].paragraphs[0]
        p.clear()
        r = p.add_run(str(running_scores[imp['criterion_idx']]))
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = STEEL_BLUE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Projected Total row
    total_row = tbl.rows[len(improvements) + 1]
    last_cell = total_row.cells[3]
    for cell in total_row.cells:
        set_cell_bg(cell, 'D6E6F4')
    total_row.cells[0].merge(total_row.cells[2])
    p = total_row.cells[0].paragraphs[0]
    p.clear()
    r = p.add_run("Projected Total Score")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARK_TEXT
    p = last_cell.paragraphs[0]
    p.clear()
    r = p.add_run("80 / 100")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = STEEL_BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def build_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_TEXT

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ─────────────────────────────────────────
    # COVER PAGE
    # ─────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = cover_title.add_run("Lucid: External Audit Report")
    rt.font.size = Pt(32)
    rt.font.bold = True
    rt.font.color.rgb = DARK_TEXT

    doc.add_paragraph()

    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = cover_sub.add_run("Three-Perspective Competitive Analysis")
    rs.font.size = Pt(16)
    rs.font.color.rgb = BODY_TEXT

    cover_co = doc.add_paragraph()
    cover_co.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cover_co.add_run("Instagram  ·  Function Health  ·  Oura Ring")
    rc.font.size = Pt(14)
    rc.font.bold = True
    rc.font.color.rgb = STEEL_BLUE

    for _ in range(4):
        doc.add_paragraph()

    cover_date = doc.add_paragraph()
    cover_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = cover_date.add_run("March 2026  |  Confidential")
    rd.font.size = Pt(11)
    rd.font.color.rgb = BODY_TEXT
    rd.font.italic = True

    add_page_break(doc)

    # ─────────────────────────────────────────
    # TABLE OF CONTENTS (manual)
    # ─────────────────────────────────────────
    add_h1(doc, "Table of Contents")

    toc_items = [
        ("Executive Summary", "3"),
        ("Part 1: Instagram Perspective", "4"),
        ("  Scoring Rubric", "4"),
        ("  Round 1 — Alex Chen, Senior Growth PM", "5"),
        ("  Round 2 — Maya Rodriguez, Consumer Engagement Lead", "6"),
        ("  Round 3 — Jordan Kim, Behavioral Science Lead", "7"),
        ("Part 2: Function Health Perspective", "8"),
        ("  Scoring Rubric", "8"),
        ("  Round 1 — Dr. Sarah Liu, Head of Biomarker Innovation", "9"),
        ("  Round 2 — Marcus Thompson, Chief Science Officer", "10"),
        ("  Round 3 — Priya Patel, Director of Member Health Outcomes", "11"),
        ("Part 3: Oura Ring Perspective", "12"),
        ("  Scoring Rubric", "12"),
        ("  Round 1 — Thomas Berg, Head of Product", "13"),
        ("  Round 2 — Elena Virtanen, VP of Science & Research", "14"),
        ("  Round 3 — Kai Nakamura, Director of Member Experience", "15"),
        ("Appendix: Cross-Perspective Top 10 Recommendations", "16"),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run_item = p.add_run(item)
        run_item.font.size = Pt(11)
        if "Part" in item or "Executive" in item or "Appendix" in item:
            run_item.font.bold = True
            run_item.font.color.rgb = DARK_TEXT
        else:
            run_item.font.color.rgb = BODY_TEXT
        # dots + page
        run_dots = p.add_run(f"{'.' * max(2, 55 - len(item))} {page}")
        run_dots.font.size = Pt(11)
        run_dots.font.color.rgb = DETAIL_GRAY

    add_page_break(doc)

    # ─────────────────────────────────────────
    # EXECUTIVE SUMMARY
    # ─────────────────────────────────────────
    add_h1(doc, "Executive Summary")

    add_h3(doc, "What Is Lucid?")
    add_body(doc, (
        "Lucid is a passive voice wellness monitor that runs silently on macOS, capturing ambient "
        "speech throughout the workday to derive a real-time Stress Index and four supporting biomarkers: "
        "voice energy, pitch variability, speech rate, and a cognitive load proxy. Unlike wearables or "
        "active journaling apps, Lucid requires no conscious user action — it listens only when the user "
        "is already speaking (calls, meetings, ambient dictation) and synthesizes those signals into a "
        "rolling Readiness Score and historical timeline. Key features include Grove (weekly voice-pattern "
        "journaling), Waypoints (personalized milestones), Echoes (pattern memory across sessions), and a "
        "Weekly Wrapped narrative summary."
    ))

    add_h3(doc, "How This Audit Works")
    add_body(doc, (
        "This report simulates nine independent audits across three companies — Instagram, Function Health, "
        "and Oura Ring — each conducting three iterative rounds of evaluation. Each company applies a custom "
        "four-criterion rubric scored 0–25 per criterion (100 points total). Auditors are fictional but "
        "representative personas drawn from each company's product, science, and growth disciplines. "
        "Each subsequent round reads prior rounds and builds upon or challenges earlier findings, producing "
        "progressively deeper analysis. Scores generally increase across rounds as reviewers identify "
        "strengths overlooked by earlier auditors."
    ))

    add_h3(doc, "Score Summary")

    summary_data = [
        ["Alex Chen", "Instagram", "1", "14", "12", "4", "16", "46"],
        ["Maya Rodriguez", "Instagram", "2", "15", "14", "7", "17", "53"],
        ["Jordan Kim", "Instagram", "3", "16", "15", "9", "19", "59"],
        ["Dr. Sarah Liu", "Function Health", "1", "13", "9", "6", "11", "39"],
        ["Marcus Thompson", "Function Health", "2", "15", "11", "8", "13", "47"],
        ["Priya Patel", "Function Health", "3", "16", "13", "9", "15", "53"],
        ["Thomas Berg", "Oura Ring", "1", "17", "13", "5", "14", "49"],
        ["Elena Virtanen", "Oura Ring", "2", "18", "14", "7", "15", "54"],
        ["Kai Nakamura", "Oura Ring", "3", "19", "16", "8", "17", "60"],
    ]
    add_summary_table(doc, summary_data)

    add_h3(doc, "Three Cross-Perspective Top Recommendations")
    add_bullet(doc, "Integrate with Oura Ring and Apple Health to correlate voice stress with HRV and sleep quality — the single highest-impact feature request across all three companies.", bold_prefix="Cross-Sensor Correlation")
    add_bullet(doc, "Introduce intimate/trusted-circle sharing and therapist-ready PDF exports to address both social engagement (Instagram's concern) and clinical utility (Function Health's demand) without compromising the mental health context.", bold_prefix="Structured Sharing & Clinical Export")
    add_bullet(doc, "Build a 90-day longitudinal narrative (\"Voice Season\") that reframes linear stress scores as an arc of recovery and resilience — directly addressing Oura's engagement-depth critique and Function Health's longitudinal data thesis.", bold_prefix="Longitudinal Voice Season Narrative")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # PART 1: INSTAGRAM
    # ═══════════════════════════════════════════════════════
    add_h1(doc, "Part 1: Instagram Perspective")

    add_body(doc, (
        "Instagram evaluates Lucid through the lens of engagement architecture, hook cycles, and "
        "long-term retention mechanics. The three auditors — drawn from Growth, Consumer Engagement, "
        "and Behavioral Science — bring progressively deeper analytical frameworks to bear on Lucid's "
        "passive-use design challenge."
    ))

    add_h2(doc, "Scoring Rubric — Instagram")
    instagram_rubric = [
        (1, "Hook Cycle Strength", "Trigger → Action → Variable Reward → Investment loop quality"),
        (2, "Notification & Re-engagement Architecture", "Timeliness, personalization, rate, re-engagement mechanics"),
        (3, "Social Proof & Viral Mechanics", "Sharing, community, network effects, word-of-mouth"),
        (4, "Habit Formation & Retention Depth", "Daily ritual anchoring, loss aversion, streak design"),
    ]
    add_rubric_table(doc, instagram_rubric)
    instagram_labels = ["Hook Cycle", "Notification", "Social/Viral", "Habit/Retention"]

    # ── Round 1: Alex Chen ──
    add_h2(doc, "Round 1 — Alex Chen, Senior Growth PM")

    add_score_card(
        doc,
        auditor="Alex Chen",
        title="Senior Growth PM, Instagram Stories",
        company="Instagram",
        round_num=1,
        scores=[14, 12, 4, 16],
        labels=instagram_labels,
        commentary=(
            "Alex Chen joined Instagram in 2019 after two years at Snap and runs growth loops and "
            "retention cohort analysis for the Stories product. He evaluates apps through the lens of "
            "\"compulsion architecture\" — his term for the systems that keep users returning daily "
            "without being asked. This is his first review of Lucid."
        )
    )

    add_h3(doc, "Hook Cycle Strength — 14/25")
    add_body(doc, (
        "Lucid's trigger is ambient and invisible — the phone call you're already on, the meeting already "
        "in progress. That's a genuinely clever distribution of trigger load onto existing behaviors. "
        "The problem is the reward. After a meeting, users get a score update. That's it. Instagram "
        "Stories delivers variable reward via social response — likes, replies, the question of 'who "
        "watched?' Lucid's reward is deterministic and self-referential. There's no 'what will I find?' "
        "moment, which is what drives compulsive checking. Grove and Waypoints are the most promising "
        "elements here — Grove's weekly reveal has genuine variable-reward DNA, but it fires weekly "
        "instead of daily, limiting its habit-formation ceiling."
    ))

    add_h3(doc, "Notification & Re-engagement Architecture — 12/25")
    add_body(doc, (
        "The notification strategy feels underdeveloped. I see a quiet hours feature, which is smart for "
        "a mental health product. But there's no evidence of push notification A/B testing, no "
        "personalized notification timing (sending notifications when the user has historically engaged, "
        "not at a fixed hour), and no re-engagement campaign for users who've gone dark for 3+ days. "
        "Instagram has learned that notification fatigue compounds fast — the answer isn't fewer "
        "notifications, it's smarter, better-timed ones. Lucid's notification architecture reads as "
        "MVP-complete but not retention-complete."
    ))

    add_h3(doc, "Social Proof & Viral Mechanics — 4/25")
    add_body(doc, (
        "This is the biggest gap. Lucid is entirely solitary. No sharing mechanics, no leaderboards, "
        "no community features, no social proof. I understand the mental health sensitivity — you don't "
        "broadcast your anxiety score. But even the most sensitive health apps have found ways to create "
        "social surface area. A 'Weekly Wrapped' share card (like Spotify's) showing streaks, Grove "
        "completions, and Waypoints reached (not scores) would give Lucid its first viral vector. "
        "Without social mechanics, Lucid relies entirely on intrinsic motivation, which is the hardest "
        "and slowest growth path available."
    ))

    add_h3(doc, "Habit Formation & Retention Depth — 16/25")
    add_body(doc, (
        "This is where Lucid actually shines. The passive capture mechanic is a master stroke for habit "
        "formation — you can't forget to use it. Waypoints and Grove create the 'investment' leg of "
        "the hook cycle (Nir Eyal's framework). The weekly Wrapped narrative is a good retention anchor. "
        "Echoes is underrated — remembering patterns across sessions ('Last time you had 3 consecutive "
        "high-stress Mondays, this is your fourth') creates the kind of persistent narrative that keeps "
        "users feeling seen. I'd give this higher, but the absence of streak recovery mechanics "
        "(what happens when I miss a week?) is a real gap. One missed week feels like a reset, not a "
        "stumble. That's a retention cliff."
    ))

    add_h3(doc, "Key Recommendations")
    add_bullet(doc, "Add a 'Weekly Wrapped' share card with Grove completions, Waypoints, and streak count (no raw scores). Give it a custom background that reflects the user's voice trend.", bold_prefix="Share Mechanic")
    add_bullet(doc, "Introduce streak-recovery (one free skip per 30 days). Frame it as a 'resilience day' not a failure. Netflix does this with downloads; Duolingo does it with streak freezes.", bold_prefix="Streak Recovery")
    add_bullet(doc, "Build notification A/B testing infrastructure. Test 8 AM vs. post-meeting vs. end-of-day. Lucid has the data to know when users are most receptive — it should use it.", bold_prefix="Notification Optimization")

    add_path_to_80(
        doc,
        auditor_name="Alex Chen",
        intro_text=(
            "The gap to 80 is achievable through four targeted interventions that directly address "
            "the reward architecture, notification intelligence, social surface, and streak resilience "
            "gaps identified above — each with a clear behavioral mechanism and measurable retention impact."
        ),
        current_scores=[14, 12, 4, 16],
        improvements=[
            {"text": "Rebuild Echoes as daily micro-reveal engine, replacing weekly Wrapped as primary variable reward delivery", "criterion_idx": 0, "points": 6},
            {"text": "Ship adaptive notification timing (ML model on voice activity windows) + 7-day lapse re-engagement campaign", "criterion_idx": 1, "points": 8},
            {"text": "Weekly Wrapped share card (no raw scores) + intimate trusted-buddy sharing with buddy-download incentive", "criterion_idx": 2, "points": 14},
            {"text": "Streak insurance (1 free skip/week as 'resilience day') + 21-day milestone delayed reveal unlock", "criterion_idx": 3, "points": 6},
        ],
        criterion_labels=["Hook Cycle Strength", "Notification & Re-engagement", "Social Proof & Viral Mechanics", "Habit Formation & Retention"],
    )

    doc.add_paragraph()

    # ── Round 2: Maya Rodriguez ──
    add_h2(doc, "Round 2 — Maya Rodriguez, Consumer Engagement Lead")

    add_score_card(
        doc,
        auditor="Maya Rodriguez",
        title="Consumer Engagement Lead, Instagram",
        company="Instagram",
        round_num=2,
        scores=[15, 14, 7, 17],
        labels=instagram_labels,
        commentary=(
            "Maya Rodriguez has a behavioral science background from her time at BJ Fogg's lab at "
            "Stanford before joining Instagram, where she now owns the 'time well spent' metrics "
            "and manages the team responsible for long-term engagement health. She has read Alex's audit "
            "and takes specific exception to his social-sharing framing."
        )
    )

    add_h3(doc, "A Note on Alex's Social Mechanics Recommendation")
    add_body(doc, (
        "Alex's instinct to add viral sharing mechanics is understandable from a pure growth lens — "
        "it's how Instagram grew. But it's the wrong instinct for a mental health app. Broadcast "
        "sharing of wellness metrics creates social comparison dynamics that are demonstrably harmful "
        "in this category (see: the fitness tracking research on Strava and eating disorders). Lucid "
        "should not try to become the Instagram of anxiety scores. The social surface area should be "
        "intimate, not public. One trusted person, not a feed."
    ), italic=True)

    add_h3(doc, "Hook Cycle Strength — 15/25")
    add_body(doc, (
        "Revising Alex's score slightly upward. The Echoes feature deserves more credit than he gave "
        "it — pattern memory is what transforms a data app into a 'knowing' app, and 'knowing' apps "
        "create deeper emotional investment than pure data apps. The trigger mechanism (passive capture "
        "during existing speech) is genuinely novel. My concern is the action layer: users have very "
        "few things they can DO in response to a high-stress reading. The app notifies them. Then what? "
        "The hook cycle is truncated because the response space is thin. Adding micro-interventions "
        "accessible from the notification ('Try a 2-minute breathing exercise' or 'Schedule a walk') "
        "would close the loop."
    ))

    add_h3(doc, "Notification & Re-engagement Architecture — 14/25")
    add_body(doc, (
        "I'm giving this higher than Alex because I think the quiet-hours feature represents genuine "
        "thoughtfulness about notification ethics, which is actually a competitive differentiator in "
        "2026. Users are increasingly notification-hostile; apps that demonstrate restraint earn trust. "
        "That said, the re-engagement strategy for lapsed users is still missing. I'd want to see a "
        "'comeback moment' — a notification that surfaces a meaningful memory after a 7-day absence. "
        "Something like: 'You had your calmest week on record three weeks ago. Want to see what was "
        "different about it?' That's memory-triggered re-engagement, not algorithmic spam."
    ))

    add_h3(doc, "Social Proof & Viral Mechanics — 7/25")
    add_body(doc, (
        "I'm giving this more than Alex's 4 because I think the solution exists — it's just not "
        "implemented yet, and the potential is real. The 'intimate sharing' model: allow users to share "
        "their Weekly Wrapped with one designated 'Lucid buddy' — a partner, therapist, or friend. "
        "The share is a narrative card, not raw data. It might say: 'This week I had 3 calm mornings "
        "in a row for the first time this month.' That creates connection without broadcasting. It also "
        "creates a two-player dynamic — the buddy gets a notification when you share, which drives "
        "app downloads. This is how Calm and Headspace's 'share a meditation' feature works."
    ))

    add_h3(doc, "Habit Formation & Retention Depth — 17/25")
    add_body(doc, (
        "Slightly above Alex's score. The morning ritual lock-in opportunity is underexplored. Lucid "
        "captures ambient speech — if a user tends to have a morning standup call, Lucid should "
        "recognize that pattern and anchor the morning check-in to that moment. 'Your day has started' "
        "becomes a trigger, not just a passive background process. The voice journal framing (Grove) "
        "is also an underutilized anchor habit. Weekly is too infrequent for ritual formation; "
        "users who check in daily with a 30-second voice note during their morning coffee would have "
        "dramatically better retention curves."
    ))

    add_h3(doc, "Key Recommendations")
    add_bullet(doc, "One trusted buddy who receives your Weekly Wrapped narrative. Drives word-of-mouth without public broadcasting or harmful social comparison.", bold_prefix="Intimate Sharing Model")
    add_bullet(doc, "Add a micro-intervention menu accessible from stress notifications: breathing exercise, walk scheduler, or a 60-second reflection prompt. Closes the hook cycle's open action gap.", bold_prefix="Micro-Intervention Response Layer")
    add_bullet(doc, "Introduce an optional 30-second daily voice note during a user-selected anchor time (morning coffee, commute). This creates a daily ritual touchpoint without requiring passive monitoring to be the sole trigger.", bold_prefix="Daily Ritual Voice Note")

    add_path_to_80(
        doc,
        auditor_name="Maya Rodriguez",
        intro_text=(
            "Closing the gap to 80 requires addressing the action layer, re-engagement mechanics, "
            "intimate sharing infrastructure, and daily ritual anchoring — all without compromising "
            "the product's mental health sensitivity or the 'time well spent' principles that "
            "distinguish Lucid from attention-harvesting competitors."
        ),
        current_scores=[15, 14, 7, 17],
        improvements=[
            {"text": "Add micro-intervention menu (breathing, walk, 60-sec reflection) accessible from every stress notification", "criterion_idx": 0, "points": 5},
            {"text": "Memory-triggered re-engagement: 'Your calmest week was 3 weeks ago — see what was different'", "criterion_idx": 1, "points": 6},
            {"text": "Full intimate sharing: narrative card to 1 trusted buddy + buddy-download notification on share", "criterion_idx": 2, "points": 10},
            {"text": "Optional 30-sec daily anchor voice note (morning coffee slot) + habit-stack detection for ritual lock-in", "criterion_idx": 3, "points": 6},
        ],
        criterion_labels=["Hook Cycle Strength", "Notification & Re-engagement", "Social Proof & Viral Mechanics", "Habit Formation & Retention"],
    )

    doc.add_paragraph()

    # ── Round 3: Jordan Kim ──
    add_h2(doc, "Round 3 — Jordan Kim, Behavioral Science Lead")

    add_score_card(
        doc,
        auditor="Jordan Kim",
        title="Behavioral Science Lead, Instagram (PhD Behavioral Economics, ex-Google DeepMind)",
        company="Instagram",
        round_num=3,
        scores=[16, 15, 9, 19],
        labels=instagram_labels,
        commentary=(
            "Jordan Kim holds a PhD in behavioral economics and joined Instagram from Google DeepMind, "
            "where they worked on reinforcement learning systems for recommendation engines. They read "
            "both Alex's and Maya's audits before writing this review, and challenge both on specific "
            "analytical points."
        )
    )

    add_h3(doc, "Building on Alex and Maya — A Synthesis")
    add_body(doc, (
        "Alex correctly identified the variable-reward gap but proposed the wrong solution (public "
        "social sharing). Maya correctly redirected toward intimate sharing but undervalued the deeper "
        "issue: Lucid's reward cadence is too predictable. From a behavioral economics standpoint, "
        "variable reward only drives compulsion when the reward is genuinely uncertain. Lucid's "
        "current design makes rewards entirely predictable: complete a week of data, get a Wrapped. "
        "The Echoes feature breaks this pattern — it fires at unpredictable moments when a meaningful "
        "pattern is detected. That unpredictability is precisely why it has the highest behavioral "
        "pull of any feature in the app. It's under-leveraged as a result."
    ), italic=True)

    add_h3(doc, "Hook Cycle Strength — 16/25")
    add_body(doc, (
        "The core hook architecture is sound but operating at roughly 60% of its potential. The "
        "trigger layer is excellent — passive capture removes friction entirely. The variable reward "
        "layer is the bottleneck. Echoes is a partial solution; the problem is it's reactive (fires "
        "only when a pattern is detected) rather than architecturally central. I'd redesign the app "
        "around Echoes as the primary reward delivery mechanism — daily micro-reveals rather than "
        "weekly macro-wraps. 'Yesterday you had a 20% calmer afternoon than your 30-day average. "
        "Here's what was different.' Short. Surprising. Behavior-connected."
    ))

    add_h3(doc, "Notification & Re-engagement Architecture — 15/25")
    add_body(doc, (
        "The single highest-ROI intervention here is adaptive notification timing. Lucid captures "
        "speech patterns — which means it knows when users are typically in meetings, when they're "
        "quiet, and when they're active. A machine learning layer (even a simple one) on top of "
        "this data would allow Lucid to identify each user's peak-engagement window — the 15-minute "
        "period each day when they're most likely to open a notification and act on it. Instagram "
        "uses this internally for push optimization; it reduced notification opt-outs by 31% in the "
        "2022 cohort study. Lucid has the behavioral data to do this — the question is whether it "
        "builds the model."
    ))

    add_h3(doc, "Social Proof & Viral Mechanics — 9/25")
    add_body(doc, (
        "I agree with Maya's reframing away from public sharing. But I want to push both of us "
        "further. The most powerful social mechanic isn't sharing data — it's shared narrative. "
        "What if Lucid allowed users who've both shared with each other to see a 'resonance score' — "
        "not individual stress levels, but a measure of how synchronized your wellbeing patterns are "
        "with someone you care about? Couples, close friends, therapy dyads. This is deeply intimate, "
        "not broadcast, and creates a genuinely new category of social feature that no wellness app "
        "currently offers. The behavioral pull is strong: humans are deeply curious about whether "
        "their emotional states are in sync with people they trust."
    ))

    add_h3(doc, "Habit Formation & Retention Depth — 19/25")
    add_body(doc, (
        "This is Lucid's strongest dimension and deserves the highest score I've given any criterion "
        "in this audit. The passive capture architecture is the single most powerful habit formation "
        "mechanic I've evaluated this year — it converts an existing behavior (speaking) into an "
        "app touchpoint without adding any new behavior. Waypoints and Grove add the investment layer "
        "correctly. Streak insurance (one free skip per week, no penalty) would make this nearly "
        "perfect by removing the loss-aversion cliff that kills retention when streaks break. I also "
        "want to call out the milestone reveal mechanic as underutilized: delayed reveals (you've "
        "been on a streak for 21 days — here's your unlock) create anticipatory reward that extends "
        "engagement cadence significantly."
    ))

    add_h3(doc, "Key Recommendations")
    add_bullet(doc, "Redesign Echoes as the primary daily reward delivery mechanism. Shift from weekly Wrapped to daily micro-reveals that surface surprising, behavior-connected insights from voice patterns.", bold_prefix="Echoes-First Architecture")
    add_bullet(doc, "Build a personal peak-engagement model from voice activity data to time notifications at each user's optimal engagement window. Low ML lift, high retention ROI.", bold_prefix="Adaptive Notification Timing")
    add_bullet(doc, "Introduce streak insurance (one free skip per 7 days) framed as a 'resilience day' — removing the loss-aversion cliff that causes retention collapse after a first missed week.", bold_prefix="Streak Insurance")
    add_bullet(doc, "Explore a 'resonance score' for shared Lucid connections — not individual stress data, but a synchronized wellbeing pattern metric for trusted pairs.", bold_prefix="Resonance Sharing")

    add_path_to_80(
        doc,
        auditor_name="Jordan Kim",
        intro_text=(
            "From a behavioral economics standpoint, the path to 80 is defined by four interventions "
            "with well-understood mechanisms: elevating Echoes to primary reward delivery, implementing "
            "personalized notification timing, creating intimate resonance-based social mechanics, "
            "and adding streak protection — each measurably closing the compulsion architecture gap."
        ),
        current_scores=[16, 15, 9, 19],
        improvements=[
            {"text": "Promote Echoes to primary daily reward delivery; daily micro-reveals over weekly Wraps", "criterion_idx": 0, "points": 3},
            {"text": "Personal peak-engagement model — notify at each user's ML-detected optimal window", "criterion_idx": 1, "points": 6},
            {"text": "Resonance score for trusted pairs: synchronized wellbeing patterns, no raw data sharing", "criterion_idx": 2, "points": 8},
            {"text": "Streak insurance + milestone cascade (30/60/90-day delayed reveals as compounding reward)", "criterion_idx": 3, "points": 4},
        ],
        criterion_labels=["Hook Cycle Strength", "Notification & Re-engagement", "Social Proof & Viral Mechanics", "Habit Formation & Retention"],
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # PART 2: FUNCTION HEALTH
    # ═══════════════════════════════════════════════════════
    add_h1(doc, "Part 2: Function Health Perspective")

    add_body(doc, (
        "Function Health evaluates Lucid through a biomarker science and clinical utility lens. "
        "The three auditors — Head of Biomarker Innovation, Chief Science Officer, and Director "
        "of Member Health Outcomes — represent the spectrum from pure science skepticism to "
        "pragmatic behavior-change focus."
    ))

    add_h2(doc, "Scoring Rubric — Function Health")
    function_rubric = [
        (1, "Biomarker Depth & Scientific Rigor", "Number/quality of markers, clinical validation, evidence standards"),
        (2, "Actionable Health Intelligence", "Prescriptiveness, personalization of recommendations, follow-through"),
        (3, "Data Ecosystem & Integration", "Health app integrations, multi-modal correlation, API richness"),
        (4, "User Health Literacy & Clinical Translation", "Translating raw scores to meaningful narratives, physician-ready reporting"),
    ]
    add_rubric_table(doc, function_rubric)
    function_labels = ["Biomarker Rigor", "Actionable Intel", "Data Ecosystem", "Clinical Translation"]

    # ── Round 1: Dr. Sarah Liu ──
    add_h2(doc, "Round 1 — Dr. Sarah Liu, Head of Biomarker Innovation")

    add_score_card(
        doc,
        auditor="Dr. Sarah Liu",
        title="Head of Biomarker Innovation, Function Health (MD/PhD)",
        company="Function Health",
        round_num=1,
        scores=[13, 9, 6, 11],
        labels=function_labels,
        commentary=(
            "Dr. Sarah Liu is an MD/PhD with eight years of clinical research experience, specializing "
            "in longitudinal biomarker validation. She joined Function Health to democratize access to "
            "lab medicine and personally oversees the scientific review process for any new biomarker "
            "Function considers adding to its panel. This is her first exposure to Lucid."
        )
    )

    add_h3(doc, "Biomarker Depth & Scientific Rigor — 13/25")
    add_body(doc, (
        "I have to be direct: voice biomarker science is nascent, not established. The literature on "
        "acoustic features as proxies for psychological stress is real — there are credible studies "
        "from DARPA, MIT Media Lab, and the CMU SSPNET consortium — but gold-standard clinical "
        "validation (randomized controlled trials correlating acoustic features with gold-standard "
        "stress biomarkers like salivary cortisol or HPA axis activity) is sparse. Lucid's DAM "
        "(Dynamic Acoustic Model) model appears to have clinical pedigree in dysphonia and cognitive "
        "assessment research, which I respect. But 'validated for speech pathology detection' is not "
        "the same as 'validated for daily stress monitoring in healthy adults.' I'd need to see Lucid's "
        "own validation study — not just citations to the field — before raising this score."
    ))

    add_h3(doc, "Actionable Health Intelligence — 9/25")
    add_body(doc, (
        "The Stress Index and Readiness Score are directionally useful, but the actionability gap is "
        "significant. Function Health's standard is: every biomarker must map to at least one "
        "concrete, evidence-based action the member can take to move it. What does a user do with "
        "a Stress Index of 73? The app tells them they're stressed — which they already know — "
        "but doesn't tell them what specifically in their voice pattern is elevated, what historical "
        "patterns predict improvement, or what behavioral interventions have worked for their specific "
        "profile. The recommendations layer feels generic. 'Consider taking a break' is not "
        "personalized health intelligence — it's a platitude."
    ))

    add_h3(doc, "Data Ecosystem & Integration — 6/25")
    add_body(doc, (
        "This is the critical gap from a clinical science perspective. Lucid sits in complete isolation. "
        "No HealthKit integration. No Oura or WHOOP sync. No cortisol correlation. No sleep quality "
        "cross-reference. This means Lucid's biomarkers can never be validated against ground truth, "
        "and Lucid's recommendations can never account for confounding variables. A high stress score "
        "on a night of poor sleep means something different than a high stress score after 8 hours "
        "of rest. Without sleep data, Lucid can't make that distinction — and is therefore offering "
        "decontextualized signals, not meaningful health intelligence."
    ))

    add_h3(doc, "User Health Literacy & Clinical Translation — 11/25")
    add_body(doc, (
        "The Weekly Wrapped narrative is a genuine strength here — translating raw metrics into a "
        "weekly story is good clinical communication practice. But there's no physician-shareable "
        "output. No PDF report. No data export format that a therapist or psychiatrist could review. "
        "If a user tells their therapist 'my stress score has been 80+ for three weeks,' the therapist "
        "has no way to independently review that data. For Function Health to ever consider a "
        "partnership, Lucid would need a clinical-grade data export — structured, time-stamped, "
        "with raw acoustic features alongside computed indices."
    ))

    add_h3(doc, "Key Recommendations")
    add_bullet(doc, "Commission or partner on a validation study: correlate Lucid's Stress Index with salivary cortisol in a controlled cohort (n≥50). Publish results. Without this, Lucid is a wellness gadget, not a health tool.", bold_prefix="Validation Study")
    add_bullet(doc, "Integrate with Apple HealthKit and Oura Ring API. Minimum viable integration: import sleep duration and HRV. Use these to contextualize stress readings.", bold_prefix="HealthKit/Oura Integration")
    add_bullet(doc, "Build a physician-shareable PDF report: 30-day voice trend chart, Stress Index distribution, notable events flagged. Export on demand, encrypted, user-controlled sharing.", bold_prefix="Clinical PDF Report")
    add_bullet(doc, "Add HRV cross-validation: on days when Apple Health HRV is low, flag that as a corroborating signal. Builds multi-modal credibility for the stress assessment.", bold_prefix="HRV Cross-Validation")

    add_path_to_80(
        doc,
        auditor_name="Dr. Sarah Liu",
        intro_text=(
            "Reaching 80 requires four concrete scientific investments: peer-reviewed validation, "
            "evidence-based intervention mapping, multi-modal sensor integration, and clinical-grade "
            "data translation — the foundational requirements for any product claiming health intelligence "
            "rather than consumer wellness."
        ),
        current_scores=[13, 9, 6, 11],
        improvements=[
            {"text": "Commission cortisol-correlation validation study (n≥50); publish in peer-reviewed journal", "criterion_idx": 0, "points": 7},
            {"text": "Evidence-based per-pattern action protocols: high pitch variability → specific breathing protocol; slow speech rate → specific intervention", "criterion_idx": 1, "points": 9},
            {"text": "HealthKit + Oura Ring API integration (sleep, HRV import); HRV cross-validation flags on stress readings", "criterion_idx": 2, "points": 14},
            {"text": "Physician-ready PDF (90-day trend, annotated events); HIPAA-compliant structured JSON export for EHR", "criterion_idx": 3, "points": 11},
        ],
        criterion_labels=["Biomarker Depth & Scientific Rigor", "Actionable Health Intelligence", "Data Ecosystem & Integration", "User Health Literacy & Clinical Translation"],
    )

    doc.add_paragraph()

    # ── Round 2: Marcus Thompson ──
    add_h2(doc, "Round 2 — Marcus Thompson, Chief Science Officer")

    add_score_card(
        doc,
        auditor="Marcus Thompson",
        title="Chief Science Officer, Function Health (Former NIH Researcher)",
        company="Function Health",
        round_num=2,
        scores=[15, 11, 8, 13],
        labels=function_labels,
        commentary=(
            "Marcus Thompson spent twelve years at the NIH studying longitudinal biomarker trajectories "
            "in aging populations before joining Function Health as CSO. He believes the future of "
            "preventive medicine is passive, longitudinal monitoring — not annual lab snapshots. He has "
            "read Dr. Liu's audit and respectfully challenges several of her more skeptical conclusions."
        )
    )

    add_h3(doc, "Pushing Back on Dr. Liu's Pessimism")
    add_body(doc, (
        "Dr. Liu's scientific rigor is valid, but I think her framing undersells what Lucid represents "
        "strategically. Her benchmark — 'show me your RCT' — is correct for a diagnostic device. "
        "But Lucid isn't claiming to be a diagnostic device. It's a passive longitudinal monitor. "
        "And passive longitudinal monitoring is the category that clinical research has historically "
        "been worst at delivering. Wearables have gotten us HRV and sleep staging. Voice is the next "
        "frontier. The lack of gold-standard validation is a gap, yes — but it's also a first-mover "
        "opportunity. Function Health should be asking: 'How do we get Lucid's longitudinal data "
        "into our member outcomes research?' not 'Does this meet our existing biomarker bar?'"
    ), italic=True)

    add_h3(doc, "Biomarker Depth & Scientific Rigor — 15/25")
    add_body(doc, (
        "Raising Dr. Liu's score because I believe the longitudinal dimension changes the calculus. "
        "A single voice snapshot has limited clinical value. Three hundred days of daily voice data "
        "from the same person is clinically significant, because you can detect drift from personal "
        "baseline — which is how I believe voice biomarkers will ultimately be validated. Personal "
        "longitudinal deviation is a stronger signal than population-level norms for most behavioral "
        "health markers. Lucid's architecture, built around historical comparison rather than absolute "
        "scores, is correctly aligned with this thesis. The 90-day voice trend is the product's "
        "scientific core — it just isn't surfaced prominently enough."
    ))

    add_h3(doc, "Actionable Health Intelligence — 11/25")
    add_body(doc, (
        "I agree with Dr. Liu's critique of generic recommendations. The opportunity I see that she "
        "didn't focus on: 'readiness for intervention' flagging. In clinical research, we know that "
        "behavior change interventions are dramatically more effective when delivered at the right "
        "moment — when the patient is in a receptive state. Lucid has a unique ability to detect "
        "that moment. A three-day trend of improving stress scores with high morning energy might "
        "indicate a user is in a growth-receptive window. That's when to surface the harder "
        "recommendations: 'Your metrics suggest now is a good time to try the 4-week sleep protocol.' "
        "Most health apps send recommendations on a schedule. Lucid could send them based on "
        "readiness state."
    ))

    add_h3(doc, "Data Ecosystem & Integration — 8/25")
    add_body(doc, (
        "Slightly above Dr. Liu's score because the API architecture for integration likely exists — "
        "the question is prioritization. I want to specifically flag the correlation engine opportunity: "
        "if Lucid integrates Oura Ring HRV data, we could run monthly correlation reports for members: "
        "'On days your Oura HRV was below 35ms, your Lucid Stress Index averaged 67 — 18 points "
        "higher than baseline.' That's Function Health-quality health intelligence. It creates a "
        "narrative about the user's specific physiology, not population averages. This is the product "
        "feature that would get me to recommend Lucid to our member base."
    ))

    add_h3(doc, "User Health Literacy & Clinical Translation — 13/25")
    add_body(doc, (
        "The 90-day trend report is the highest-leverage clinical translation feature Lucid doesn't "
        "yet have. A monthly 'Voice Health Summary' — one page, four charts (Stress Index trend, "
        "Grove participation rate, Waypoints velocity, notable high-stress events), plain-language "
        "narrative — would be the first voice biomarker longitudinal report any consumer product "
        "has delivered. Function Health members would find this genuinely useful as a complement "
        "to their quarterly lab panels. The physician-ready PDF Dr. Liu recommended is the right "
        "call; I'd add a structured JSON export for EHR integration as a longer-term goal."
    ))

    add_h3(doc, "Key Recommendations")
    add_bullet(doc, "Build a monthly 'Voice Health Summary' — 90-day trend report, four charts, plain-language narrative, PDF export. The first clinical-grade voice biomarker longitudinal report in consumer health.", bold_prefix="90-Day Voice Trend Report")
    add_bullet(doc, "Develop a readiness-state detection layer: identify 3-day improvement windows and surface harder behavior-change recommendations only when the user is physiologically receptive.", bold_prefix="Readiness-for-Intervention Detection")
    add_bullet(doc, "Build an Oura HRV correlation engine as the first post-integration feature. Show users the specific relationship between their HRV and their stress patterns.", bold_prefix="HRV Correlation Engine")

    add_path_to_80(
        doc,
        auditor_name="Marcus Thompson",
        intro_text=(
            "The longitudinal monitoring thesis supports reaching 80 through four advances: transparent "
            "methodology publication, receptivity-state intervention timing, Oura HRV correlation, "
            "and clinical-grade reporting — all achievable within a 12-month roadmap and consistent "
            "with how passive biomarker research translates into actionable member health tools."
        ),
        current_scores=[15, 11, 8, 13],
        improvements=[
            {"text": "Publish longitudinal baseline methodology (window length, deviation calculation); position as 'personal biomarker drift' monitoring", "criterion_idx": 0, "points": 7},
            {"text": "Readiness-for-intervention detection layer; surface harder recommendations only during 3-day improvement windows", "criterion_idx": 1, "points": 8},
            {"text": "Oura HRV correlation engine (user-specific HRV↔stress relationship report); sleep quality context layer", "criterion_idx": 2, "points": 11},
            {"text": "Monthly Voice Health Summary PDF (4 charts, plain-language narrative); structured JSON export for EHR", "criterion_idx": 3, "points": 7},
        ],
        criterion_labels=["Biomarker Depth & Scientific Rigor", "Actionable Health Intelligence", "Data Ecosystem & Integration", "User Health Literacy & Clinical Translation"],
    )

    doc.add_paragraph()

    # ── Round 3: Priya Patel ──
    add_h2(doc, "Round 3 — Priya Patel, Director of Member Health Outcomes")

    add_score_card(
        doc,
        auditor="Priya Patel",
        title="Director of Member Health Outcomes, Function Health",
        company="Function Health",
        round_num=3,
        scores=[16, 13, 9, 15],
        labels=function_labels,
        commentary=(
            "Priya Patel leads Member Health Outcomes at Function Health, responsible for the programs "
            "and interventions that actually change member behavior after they receive their results. "
            "She comes from an operations and member success background and focuses on what moves the "
            "needle in real-world member journeys, not just clinical trials. She has read both "
            "Dr. Liu's and Marcus's audits."
        )
    )

    add_h3(doc, "From Science to Member Journey")
    add_body(doc, (
        "Both Sarah and Marcus are asking the right scientific questions. I want to add the third "
        "dimension: what actually happens when a member's Lucid score is elevated for three weeks "
        "in a row? Right now, nothing. The app shows them a chart. Maybe it sends a push notification. "
        "But there's no care pathway. At Function Health, we've learned that the moment of insight — "
        "seeing a concerning biomarker result — is simultaneously the highest-leverage and "
        "highest-risk moment in the member journey. Highest-leverage because the member is motivated. "
        "Highest-risk because if we don't immediately give them somewhere to go, the motivation "
        "dissipates within 48 hours. Lucid needs a care escalation pathway."
    ), italic=True)

    add_h3(doc, "Biomarker Depth & Scientific Rigor — 16/25")
    add_body(doc, (
        "Raising further based on a point neither Sarah nor Marcus made: the behavioral biomarker "
        "dimension. Lucid isn't just measuring physiological stress markers — it's measuring "
        "behavioral consistency (did the user complete Grove this week?), social engagement patterns "
        "(meeting frequency derived from call volume), and cognitive load proxies. These are "
        "behavioral biomarkers that clinical research consistently shows are leading indicators "
        "of depression and burnout. The fact that Lucid captures them passively, without self-report "
        "bias, is a genuine scientific contribution. The scientific story should be reframed: not "
        "'we measure stress' but 'we are the first passive behavioral biomarker monitor for "
        "psychological health.'"
    ))

    add_h3(doc, "Actionable Health Intelligence — 13/25")
    add_body(doc, (
        "Marcus's readiness-for-intervention concept is correct and I want to operationalize it. "
        "At Function Health we call this 'care navigation' — when a member's result crosses a "
        "clinical threshold, a care navigator proactively reaches out. Lucid should build a tiered "
        "care escalation protocol: at Stress Index >70 for 5+ consecutive days, surface EAP (Employee "
        "Assistance Program) resources and mental health app recommendations. At depression risk "
        "score >70%, prompt a therapist-matching flow (integration with Headway, Alma, or BetterHelp). "
        "This is how you turn a wellness app into a mental health tool — without overstating clinical "
        "efficacy."
    ))

    add_h3(doc, "Data Ecosystem & Integration — 9/25")
    add_body(doc, (
        "The employer/HR integration opportunity is underexplored. Lucid's aggregate anonymized data "
        "has significant value for employers tracking workforce wellbeing — not individual scores "
        "(never that), but team-level trends. 'Your engineering org's average Stress Index increased "
        "23% in the two weeks following the product launch' is the kind of insight that CHROs pay "
        "for. An anonymous aggregate dashboard for HR leaders — opt-in at the employee level, "
        "aggregate-only at the employer level — would open a B2B revenue channel and a distribution "
        "flywheel: employers pay for Lucid licenses, employees get free access. This is exactly "
        "how Function Health's enterprise program scales."
    ))

    add_h3(doc, "User Health Literacy & Clinical Translation — 15/25")
    add_body(doc, (
        "The therapist data export Marcus and Sarah both recommended is correct. I want to add "
        "the specifics of what a HIPAA-compliant export should contain: (1) a 90-day Stress Index "
        "chart with notable events annotated, (2) Grove participation and theme summary (no raw "
        "transcripts — privacy), (3) a plain-language clinical summary paragraph generated by "
        "the app (not AI-generated text, but structured from the data), and (4) a 'flags' section "
        "highlighting any elevated-risk periods. This gives a therapist something to work with in "
        "a 45-minute session. The format should mirror what wearable companies (Oura, WHOOP) have "
        "built for clinician reports."
    ))

    add_h3(doc, "Key Recommendations")
    add_bullet(doc, "Build a tiered care escalation protocol: surface EAP resources at Stress Index >70 for 5+ days; prompt therapist-matching at depression risk >70%. Turns wellness app into mental health gateway.", bold_prefix="Care Escalation Pathway")
    add_bullet(doc, "Build an anonymous aggregate employer dashboard (opt-in per employee, aggregate-only). Opens B2B revenue channel and enterprise distribution flywheel.", bold_prefix="Enterprise Aggregate Dashboard")
    add_bullet(doc, "Create a HIPAA-compliant therapist export: 90-day trend chart, Grove theme summary (no transcripts), clinical summary paragraph, elevated-risk flags. Standard format therapists can use in sessions.", bold_prefix="HIPAA-Compliant Therapist Export")
    add_bullet(doc, "Reframe Lucid's scientific positioning: not 'voice stress monitor' but 'passive behavioral biomarker monitor for psychological health.' This framing is more defensible, more compelling to clinicians, and more differentiated from competitors.", bold_prefix="Scientific Positioning Reframe")

    add_path_to_80(
        doc,
        auditor_name="Priya Patel",
        intro_text=(
            "From a member journey perspective, reaching 80 means operationalizing four pathways: "
            "clearer scientific positioning, tiered care escalation, enterprise aggregate analytics, "
            "and HIPAA-compliant clinical export — each directly extending Lucid's value within a "
            "care ecosystem and opening distribution channels that no current consumer voice app has built."
        ),
        current_scores=[16, 13, 9, 15],
        improvements=[
            {"text": "Reframe public positioning as 'passive behavioral biomarker monitor for psychological health' with supporting data brief", "criterion_idx": 0, "points": 4},
            {"text": "Tiered care escalation: EAP resources at Stress >70 for 5+ days; therapist-matching at depression risk >70%", "criterion_idx": 1, "points": 7},
            {"text": "Anonymous aggregate employer dashboard (opt-in per employee); EHR JSON export pathway", "criterion_idx": 2, "points": 10},
            {"text": "HIPAA therapist export (90-day chart, Grove themes, clinical summary, elevated-risk flags); care navigator integration", "criterion_idx": 3, "points": 6},
        ],
        criterion_labels=["Biomarker Depth & Scientific Rigor", "Actionable Health Intelligence", "Data Ecosystem & Integration", "User Health Literacy & Clinical Translation"],
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # PART 3: OURA RING
    # ═══════════════════════════════════════════════════════
    add_h1(doc, "Part 3: Oura Ring Perspective")

    add_body(doc, (
        "Oura Ring evaluates Lucid as a potential complement to physical recovery monitoring — "
        "the 'mental HRV' to Oura's physiological HRV. The three auditors bring product, science, "
        "and member experience perspectives to an evaluation that starts with genuine respect for "
        "Lucid's passive architecture and builds toward a partnership vision."
    ))

    add_h2(doc, "Scoring Rubric — Oura Ring")
    oura_rubric = [
        (1, "Passive Monitoring Fidelity", "How truly passive the experience is, friction, coverage across contexts"),
        (2, "Recovery & Readiness Narrative Quality", "How well the app explains what the data means for today's performance"),
        (3, "Data Ecosystem & Cross-sensor Correlation", "Breadth of integrations, ability to combine with other health signals"),
        (4, "Long-term Engagement & Behavior Change", "90-day retention, behavioral outcomes, daily ritual design"),
    ]
    add_rubric_table(doc, oura_rubric)
    oura_labels = ["Passive Fidelity", "Readiness Narrative", "Cross-Sensor", "Long-term Engage"]

    # ── Round 1: Thomas Berg ──
    add_h2(doc, "Round 1 — Thomas Berg, Head of Product")

    add_score_card(
        doc,
        auditor="Thomas Berg",
        title="Head of Product, Oura Ring (Ex-Polar Sports Tracking)",
        company="Oura Ring",
        round_num=1,
        scores=[17, 13, 5, 14],
        labels=oura_labels,
        commentary=(
            "Thomas Berg joined Oura in 2021 after six years at Polar building sports performance "
            "monitoring products. He owns the Oura Ring app product and has a deep appreciation for "
            "passive monitoring done right — having spent years fighting the 'remember to log' "
            "problem in sports tracking. Lucid is his first exposure to voice-based wellness monitoring."
        )
    )

    add_h3(doc, "Passive Monitoring Fidelity — 17/25")
    add_body(doc, (
        "Lucid's passive architecture is, frankly, impressive. We've spent years at Oura explaining "
        "to users why wearing a ring to bed is worth the friction — because the data you get from "
        "continuous passive monitoring is qualitatively different from anything you can log manually. "
        "Lucid has taken that thesis and applied it to psychological health. No charging. No "
        "wearing. No remembering. It just captures what you're already doing. I give it high marks "
        "here but not full marks because of the coverage gaps: evenings and weekends when users "
        "speak less are monitored less. A user who works from home and makes few calls on weekends "
        "has a 30% data gap compared to their workweek baseline. Oura captures HRV 24/7; "
        "Lucid captures voice 60-70% of hours. That asymmetry limits some clinical claims."
    ))

    add_h3(doc, "Recovery & Readiness Narrative Quality — 13/25")
    add_body(doc, (
        "Oura's Readiness Score is our most-loved feature — not because it's scientifically complex, "
        "but because it answers the question users actually wake up asking: 'How am I today?' Lucid's "
        "equivalent is the Stress Index and the daily Grove entry, but neither quite answers "
        "that morning question. The Stress Index is real-time and retrospective, not forward-looking. "
        "What I'd want is a 'Mental Readiness Score' for the day ahead — synthesizing the prior "
        "day's voice patterns, Grove reflection, and (with integration) sleep quality to give a "
        "simple, actionable morning number. That's what Oura would build if it were extending into "
        "the psychological dimension."
    ))

    add_h3(doc, "Data Ecosystem & Cross-sensor Correlation — 5/25")
    add_body(doc, (
        "This is the product's largest gap from our perspective. Oura Ring is the most obvious "
        "integration partner for Lucid: we measure the physiological stress response (HRV depression, "
        "elevated resting heart rate, reduced sleep stage quality); Lucid measures the behavioral "
        "stress expression (voice pattern changes). Together, we'd have the most comprehensive "
        "non-invasive psychological stress picture available outside a clinical lab. Without the "
        "Oura integration, Lucid is measuring one dimension of a multi-dimensional phenomenon. "
        "The absence of any health ecosystem integration is the single most important product gap "
        "to close in the next 12 months."
    ))

    add_h3(doc, "Long-term Engagement & Behavior Change — 14/25")
    add_body(doc, (
        "The Grove and Waypoints mechanics are thoughtfully designed — they create meaningful "
        "engagement touchpoints beyond passive data collection, which is critical for long-term "
        "retention. Oura struggled early on with users who wore the ring for 6 weeks and then "
        "stopped because they felt they'd 'learned what they needed to know.' We solved that with "
        "ongoing narrative evolution — your Readiness Score story changes as your baseline shifts "
        "and your habits evolve. Lucid will face the same 6-week cliff. The Waypoints system "
        "addresses this somewhat, but I don't see enough evidence of ongoing narrative evolution — "
        "content and insights that get more sophisticated the longer you use the product."
    ))

    add_h3(doc, "Key Recommendations")
    add_bullet(doc, "Partner with Oura Ring for bidirectional API integration. Show correlation cards: 'On days your Oura sleep score was <70, your Lucid Stress Index peaked 34% higher.' This is the product demo that would close enterprise deals.", bold_prefix="Oura Integration Partnership")
    add_bullet(doc, "Build a morning 'Mental Readiness Score' synthesizing prior-day voice patterns, Grove reflection sentiment, and (optionally) Oura sleep data. Answer the morning question: 'How am I today?'", bold_prefix="Morning Mental Readiness Score")
    add_bullet(doc, "Design a 'narrative evolution system' where Lucid's insights increase in sophistication as tenure grows — unlock new pattern types, more granular historical comparisons, and advanced Echoes after 30/60/90 days.", bold_prefix="Narrative Evolution System")

    add_path_to_80(
        doc,
        auditor_name="Thomas Berg",
        intro_text=(
            "The path to 80 is defined by four product decisions that address the silence-coverage gap, "
            "narrative quality, Oura integration, and longitudinal engagement architecture — all "
            "consistent with how Oura solved its own 6-week retention cliff through narrative evolution "
            "and cross-sensor correlation."
        ),
        current_scores=[17, 13, 5, 14],
        improvements=[
            {"text": "Reframe silence as signal ('Unusually quiet weekend — your first genuine rest period this month'); evening speech check-in option", "criterion_idx": 0, "points": 4},
            {"text": "Morning Mental Readiness Score synthesizing prior-day voice patterns + Oura sleep data (optional)", "criterion_idx": 1, "points": 8},
            {"text": "Oura bidirectional API integration; weekly Correlation Cards (one personalized physical↔mental insight)", "criterion_idx": 2, "points": 14},
            {"text": "Narrative evolution system — insights unlock new pattern types at 30/60/90-day milestones", "criterion_idx": 3, "points": 5},
        ],
        criterion_labels=["Passive Monitoring Fidelity", "Recovery & Readiness Narrative", "Data Ecosystem & Cross-sensor", "Long-term Engagement & Behavior Change"],
    )

    doc.add_paragraph()

    # ── Round 2: Elena Virtanen ──
    add_h2(doc, "Round 2 — Elena Virtanen, VP of Science & Research")

    add_score_card(
        doc,
        auditor="Elena Virtanen",
        title="VP of Science & Research, Oura Ring (Neuroscientist, co-authored Oura validation studies)",
        company="Oura Ring",
        round_num=2,
        scores=[18, 14, 7, 15],
        labels=oura_labels,
        commentary=(
            "Dr. Elena Virtanen is a Finnish neuroscientist who co-authored Oura Ring's primary "
            "validation studies published in Sleep Medicine and Frontiers in Physiology. She evaluates "
            "products through a measurement science lens and has deep expertise in longitudinal "
            "biomarker normalization. She has read Thomas's audit and focuses on the scientific "
            "validity and measurement architecture dimensions he touched only briefly."
        )
    )

    add_h3(doc, "Passive Monitoring Fidelity — 18/25")
    add_body(doc, (
        "Thomas gave this 17; I'd push to 18 with nuance. The fidelity issue he raised — evenings "
        "and weekends — is real, but there's a more fundamental measurement architecture question "
        "he didn't address: how does Lucid handle baseline normalization? At Oura, we spent two "
        "years developing the 2-week rolling baseline window for Readiness scoring. Before that "
        "window was established, we had constant complaints that the score 'felt wrong.' The "
        "complaints stopped when we moved from absolute scoring (your HRV is 38ms) to deviation "
        "from personal baseline scoring (your HRV is 12% below your 14-day average). Lucid's "
        "architecture appears to use historical comparison — good — but the specific window length "
        "and normalization methodology should be explicitly documented and user-explainable."
    ))

    add_h3(doc, "Recovery & Readiness Narrative Quality — 14/25")
    add_body(doc, (
        "Thomas's Mental Readiness Score proposal is right. I want to add the scientific framing "
        "for how to build it credibly. The key insight from our Readiness research: the most "
        "predictive readiness signal is not the single best-performing metric — it's the "
        "pattern of co-variation across multiple metrics. A low HRV combined with elevated RHR "
        "and shortened deep sleep is a stronger readiness signal than any single metric alone. "
        "For Lucid, the equivalent would be: elevated stress index combined with reduced speech "
        "rate and compressed pitch range is a stronger signal than stress index alone. The "
        "narrative quality improves dramatically when you explain the multi-signal basis: "
        "'Three of your four voice markers are elevated today, including...' not just 'Stress is high.'"
    ))

    add_h3(doc, "Data Ecosystem & Cross-sensor Correlation — 7/25")
    add_body(doc, (
        "Higher than Thomas's 5 because I want to acknowledge what Lucid has done right in the "
        "ecosystem dimension: the calendar integration (if present) for meeting context is the "
        "right first integration choice. Contextualizing voice patterns relative to scheduled "
        "activities is a meaningful enrichment even without external health data. That said, "
        "Thomas's point about the Oura integration is correct and I'd extend it: population-level "
        "percentile comparisons become possible once you have a large enough dataset. '78th percentile "
        "for weekday stress recovery' is more motivating than an absolute score, and Lucid should "
        "be building toward that population-level normative dataset actively."
    ))

    add_h3(doc, "Long-term Engagement & Behavior Change — 15/25")
    add_body(doc, (
        "The longitudinal voice biomarker validation study is the highest-leverage long-term action "
        "Lucid could take. Not just for scientific credibility, but for user engagement. At Oura, "
        "the publications of our validation studies drove 40%+ spikes in app engagement — users "
        "who hadn't opened the app in months came back to see 'how their data compared.' A "
        "peer-reviewed study demonstrating that Lucid's 90-day voice trend predicts self-reported "
        "burnout scores 3 weeks in advance would be the most powerful retention and marketing asset "
        "the company could build. I'd propose a partnership: Oura + Lucid co-validation study, "
        "combining HRV trajectories with voice pattern trajectories in a burnout prediction model."
    ))

    add_h3(doc, "Key Recommendations")
    add_bullet(doc, "Document and surface the baseline normalization methodology (what window? what calculation?) to users. Transparency about how scores are computed increases trust in the product.", bold_prefix="Transparent Baseline Normalization")
    add_bullet(doc, "Upgrade the narrative to explain multi-signal basis: 'Three of four voice markers elevated' rather than 'Stress is high.' Mirrors Oura's multi-metric Readiness explanation and dramatically increases perceived scientific credibility.", bold_prefix="Multi-Signal Narrative")
    add_bullet(doc, "Pursue a co-validation study with Oura (or independently): 90-day voice pattern trajectory as a burnout predictor. Peer-reviewed publication as the most durable marketing and retention asset.", bold_prefix="Co-Validation Study")
    add_bullet(doc, "Build population-level percentile data infrastructure now, even if not surfaced to users immediately. This becomes a moat as dataset size grows.", bold_prefix="Population Percentile Database")

    add_path_to_80(
        doc,
        auditor_name="Elena Virtanen",
        intro_text=(
            "Scientific rigor and measurement transparency define the path to 80: surfacing normalization "
            "methodology, upgrading multi-signal narrative architecture, building population percentile "
            "infrastructure with Oura integration, and co-publishing validation research — each raising "
            "the product's credibility ceiling in ways that compound over time."
        ),
        current_scores=[18, 14, 7, 15],
        improvements=[
            {"text": "Surface baseline normalization methodology to users (window length, deviation logic); transparency as trust signal", "criterion_idx": 0, "points": 3},
            {"text": "Multi-signal narrative ('3 of 4 voice markers elevated today including...'); forward-looking readiness framing", "criterion_idx": 1, "points": 7},
            {"text": "Population percentile database infrastructure; Oura integration enabling HRV↔voice correlation", "criterion_idx": 2, "points": 11},
            {"text": "Co-publish validation study (Oura+Lucid burnout prediction model); publication drives re-engagement spike", "criterion_idx": 3, "points": 5},
        ],
        criterion_labels=["Passive Monitoring Fidelity", "Recovery & Readiness Narrative", "Data Ecosystem & Cross-sensor", "Long-term Engagement & Behavior Change"],
    )

    doc.add_paragraph()

    # ── Round 3: Kai Nakamura ──
    add_h2(doc, "Round 3 — Kai Nakamura, Director of Member Experience")

    add_score_card(
        doc,
        auditor="Kai Nakamura",
        title="Director of Member Experience, Oura Ring (Previously Headspace)",
        company="Oura Ring",
        round_num=3,
        scores=[19, 16, 8, 17],
        labels=oura_labels,
        commentary=(
            "Kai Nakamura joined Oura from Headspace, where they led content and member journey "
            "design. They are obsessed with the 'moment of insight' — the specific in-product "
            "experience where a user has their first genuine aha moment that transforms them from "
            "casual user to committed advocate. They have read both Thomas's and Elena's audits "
            "and synthesize the product and science perspectives through a member experience lens."
        )
    )

    add_h3(doc, "The Aha Moment Architecture")
    add_body(doc, (
        "Thomas focused on what Lucid measures and how it integrates. Elena focused on how it "
        "measures. I want to focus on the moment when a user first truly understands what Lucid "
        "is doing for them — the aha moment. At Headspace, we tracked this obsessively: most "
        "users' aha moment came at Day 10, when they noticed for the first time that they'd "
        "completed more consecutive days of meditation than they'd ever managed before. At Oura, "
        "the aha moment typically comes at Day 21, when a user sees their first HRV improvement "
        "trend correlated with a behavior change they made. Lucid's aha moment architecture "
        "is undefined. When does it happen? What is the insight? How is it designed? "
        "This is the product gap that underlies everything else both Thomas and Elena identified."
    ), italic=True)

    add_h3(doc, "Passive Monitoring Fidelity — 19/25")
    add_body(doc, (
        "Highest score in this audit for this criterion, because the passive architecture is "
        "genuinely exceptional. The coverage gap Thomas identified is real — but I frame it "
        "differently: evenings and weekends are an opportunity, not just a gap. 'Lucid detected "
        "you had an unusually quiet weekend — your voice patterns suggest genuine rest for the "
        "first time this month' is a valuable insight that turns the low-data period into a "
        "positive signal. Silence is data. Absence of speech in a usually-verbose user is "
        "meaningful. The fidelity score is high because Lucid has designed out the most "
        "common passive monitoring failure mode: user forgetting."
    ))

    add_h3(doc, "Recovery & Readiness Narrative Quality — 16/25")
    add_body(doc, (
        "Thomas and Elena both proposed versions of a forward-looking Readiness Score; I want "
        "to propose the full narrative arc concept that would make Lucid genuinely distinct. "
        "I call it the 'Voice Season' — a 3-month narrative arc that gives meaning to the "
        "user's data trajectory. At month 1, you're establishing baseline. At month 2, Lucid "
        "starts identifying your patterns. At month 3, Lucid begins predicting your weekly "
        "rhythm and can say 'Tuesday afternoons are historically your most creative window — "
        "here's what your voice markers show about this Tuesday.' The Season framing replaces "
        "linear stress scores with an arc of discovery and growth, which is far more compelling "
        "over a 90-day engagement window."
    ))

    add_h3(doc, "Data Ecosystem & Cross-sensor Correlation — 8/25")
    add_body(doc, (
        "Slightly above Elena's 7 because I want to advocate for a simple correlation card "
        "feature that would require minimal integration work but maximum emotional impact. "
        "The correlation card: after 30 days of Oura + Lucid data, serve a single card per week "
        "that surfaces the single most interesting correlation between physical and mental "
        "biomarkers for that specific user. 'On your lowest-HRV days, your voice stress peaks "
        "4 hours later — not immediately. Your body knows before your voice does.' That insight "
        "is personalized, surprising, and behavior-connected. It's the kind of thing users "
        "screenshot and share. It's the aha moment architecture made concrete."
    ))

    add_h3(doc, "Long-term Engagement & Behavior Change — 17/25")
    add_body(doc, (
        "The 'resilience journey' view is the long-term engagement design I'd recommend. Instead "
        "of showing users a linear stress metric over time (which is inherently depressing when "
        "the trend is bad), show them a 'resilience arc' — the ratio of stress events to recovery "
        "speed. A user who had a terrible week but recovered in 2 days is more resilient than one "
        "who had a moderate week but took 5 days to recover. Framing the metric this way rewards "
        "process over outcome, which is the behavior science gold standard for sustainable engagement. "
        "Additionally: celebration of invisible improvement — detecting subtle positive trends before "
        "the user consciously notices them ('Your baseline stress has decreased 11% over 60 days, "
        "quietly') creates moments of delight that dramatically increase long-term retention."
    ))

    add_h3(doc, "Key Recommendations")
    add_bullet(doc, "Design the Lucid aha moment explicitly: target Day 14, tied to the first Echoes insight that reveals a pattern the user hadn't consciously recognized. Make this moment central to onboarding design.", bold_prefix="Aha Moment Architecture")
    add_bullet(doc, "Introduce the Voice Season: a 3-month narrative arc (Discovery → Patterns → Prediction) that reframes data as a journey rather than a score. Makes 90-day retention feel like a destination, not a chore.", bold_prefix="Voice Season (3-Month Arc)")
    add_bullet(doc, "Build weekly Correlation Cards after Oura integration: one surprising, personalized physical-mental correlation per week. The screenshot-worthy insight that drives word-of-mouth.", bold_prefix="Correlation Cards")
    add_bullet(doc, "Shift long-term engagement framing from 'stress score trend' to 'resilience arc' — reward recovery speed, not just stress level. Celebrate invisible improvements before users consciously notice them.", bold_prefix="Resilience Arc View")

    add_path_to_80(
        doc,
        auditor_name="Kai Nakamura",
        intro_text=(
            "The aha moment architecture drives the path to 80: silence-as-signal recognition, "
            "full Voice Season implementation, weekly Correlation Cards post-Oura integration, "
            "and a resilience arc framing — all designed to create the moments of delight that "
            "convert passive users into committed long-term advocates."
        ),
        current_scores=[19, 16, 8, 17],
        improvements=[
            {"text": "Silence-as-data detection for evenings/weekends; rest period recognition as positive insight", "criterion_idx": 0, "points": 3},
            {"text": "Full Voice Season implementation (Discovery → Patterns → Prediction arc across 90 days)", "criterion_idx": 1, "points": 5},
            {"text": "Weekly Correlation Cards post-Oura integration; the screenshot-worthy personalized insight", "criterion_idx": 2, "points": 10},
            {"text": "Resilience Arc view (recovery speed ratio) + invisible improvement celebrations ('Your baseline stress dropped 11% quietly')", "criterion_idx": 3, "points": 2},
        ],
        criterion_labels=["Passive Monitoring Fidelity", "Recovery & Readiness Narrative", "Data Ecosystem & Cross-sensor", "Long-term Engagement & Behavior Change"],
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════
    # APPENDIX
    # ═══════════════════════════════════════════════════════
    add_h1(doc, "Appendix: Cross-Perspective Top 10 Recommendations")

    add_body(doc, (
        "The following recommendations represent convergence across all nine auditors. Items are ranked "
        "by the number of independent auditors who raised the theme, weighted by the urgency of their "
        "framing. Items mentioned by all three companies appear first."
    ))

    recommendations = [
        (
            "1. Calendar & Meeting Integration",
            "Instagram · Function Health · Oura Ring",
            "Contextualizing voice patterns relative to scheduled activities is the foundational "
            "enrichment layer all three companies referenced. Knowing that a stress spike occurred "
            "during a recurring 1:1 or immediately before a board presentation transforms raw "
            "data into meaningful life context. This is also the lowest-lift integration (read-only "
            "calendar access) with disproportionate impact on insight quality."
        ),
        (
            "2. Oura Ring + Apple Health Correlation",
            "Function Health · Oura Ring (primary) · Instagram (secondary)",
            "The single most frequently cited gap across all audits. Correlating voice stress with "
            "HRV, sleep quality, and resting heart rate creates the multi-modal health picture that "
            "transforms Lucid from a wellness gadget into a serious health monitoring tool. Marcus "
            "Thompson's HRV correlation engine and Thomas Berg's 'mental HRV' framing both point "
            "to this as the 12-month priority."
        ),
        (
            "3. Intimate / Trusted-Circle Sharing",
            "Instagram (primary) · Oura Ring (secondary)",
            "Maya Rodriguez and Jordan Kim both converged on this after Alex Chen's public-sharing "
            "recommendation was challenged. Sharing with one trusted person (partner, therapist, "
            "close friend) — as a narrative card, not raw data — addresses the social engagement "
            "gap without the mental health risks of broadcast sharing. Drives app downloads through "
            "the trusted-person notification mechanic."
        ),
        (
            "4. Weekly Wrapped Social Share Card",
            "Instagram (primary)",
            "A shareable card surfacing streaks, Grove completions, and Waypoints reached (no raw "
            "stress scores) gives Lucid its first viral distribution vector. Alex Chen identified "
            "this as the highest-ROI social feature; Maya and Jordan refined the format to protect "
            "mental health context. Spotify's Wrapped model is the obvious reference."
        ),
        (
            "5. Therapist / Physician-Ready PDF Export",
            "Function Health (primary) · Oura Ring (secondary)",
            "Dr. Sarah Liu, Marcus Thompson, and Priya Patel all cited the clinical export gap. "
            "A HIPAA-compliant PDF containing 90-day trend chart, Grove theme summary, notable "
            "elevated-risk periods, and a plain-language clinical summary gives Lucid a credibility "
            "pathway with the clinical community and creates a 'therapist referral' growth channel."
        ),
        (
            "6. Smart Quiet Hours (Meeting & Sleep Aware)",
            "Instagram · Function Health · Oura Ring",
            "All three companies referenced notification quality as a differentiator. Smart quiet "
            "hours — automatically suppressing notifications during detected meetings, sleep windows "
            "(from Oura integration), and user-defined focus periods — demonstrates the restraint "
            "that earns long-term user trust in a notification-hostile environment."
        ),
        (
            "7. 90-Day Longitudinal Trend Narrative (\"Voice Season\")",
            "Function Health (primary) · Oura Ring (primary)",
            "Marcus Thompson's 90-day voice trend report and Kai Nakamura's Voice Season concept "
            "converge on the same insight: Lucid's scientific and narrative differentiation is in "
            "the longitudinal dimension, not the point-in-time snapshot. A 3-month arc narrative "
            "(Discovery → Patterns → Prediction) transforms the user journey from data collection "
            "to self-discovery."
        ),
        (
            "8. Depression Risk Triage Pathway",
            "Function Health (primary)",
            "Priya Patel's care escalation protocol — surfacing EAP resources at sustained elevated "
            "stress and therapist-matching flows at high depression risk scores — is the highest-"
            "responsibility feature on this list. It also has the potential to be Lucid's most "
            "differentiating: no consumer voice app currently provides a care pathway. "
            "Implementation requires clinical advisory oversight and careful UX design to avoid "
            "false positives and user alarm."
        ),
        (
            "9. Streak Insurance / Recovery Mechanic",
            "Instagram (primary) · Oura Ring (secondary)",
            "Jordan Kim's streak insurance concept (one free skip per 7 days, framed as a "
            "'resilience day') directly addresses the retention cliff that occurs when users break "
            "a streak. Duolingo's streak freeze and Netflix's download expiration grace period "
            "are both built on the same behavioral insight: removing the loss-aversion cliff is "
            "more effective than increasing reward for continuation."
        ),
        (
            "10. Adaptive Notification Timing",
            "Instagram (primary)",
            "Jordan Kim's personal peak-engagement model uses Lucid's own behavioral data to time "
            "notifications at the moment each user is most likely to engage — derived from voice "
            "activity patterns. Low ML complexity, high retention ROI. Instagram's 2022 implementation "
            "reduced notification opt-outs by 31% in the same cohort. This is the notification "
            "architecture that evolves from MVP-complete to retention-complete."
        ),
    ]

    for title, companies, body in recommendations:
        add_h3(doc, title)
        p_co = doc.add_paragraph()
        p_co.paragraph_format.space_after = Pt(3)
        r = p_co.add_run(f"Referenced by: {companies}")
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = STEEL_BLUE
        add_body(doc, body)
        doc.add_paragraph()

    # Final note
    final_p = doc.add_paragraph()
    final_p.paragraph_format.space_before = Pt(12)
    r = final_p.add_run(
        "This report was prepared as an internal strategic document for Lucid development planning. "
        "All auditor personas are fictional and representative. Scores and recommendations reflect "
        "independent analytical frameworks applied to Lucid's publicly observable feature set. "
        "March 2026 · Confidential"
    )
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = BODY_TEXT

    doc.save(OUTPUT_PATH)
    print(f"✓ Document saved: {OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(f"Success: {path}")
