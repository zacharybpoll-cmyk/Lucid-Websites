#!/usr/bin/env python3
"""
Generate the Attune Wearable Extension Technical Research Report.
Output: Attune_Wearable_Extension_Technical_Report.docx
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Styling Constants ──────────────────────────────────────────────────────────
FONT_BODY = "Calibri"
FONT_MONO = "Courier New"
COLOR_PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)  # Dark navy
COLOR_BODY = RGBColor(0x2D, 0x2D, 0x2D)
COLOR_SECONDARY = RGBColor(0x55, 0x55, 0x55)
COLOR_HEADING2 = RGBColor(0x2C, 0x2C, 0x2C)
COLOR_HEADING3 = RGBColor(0x44, 0x44, 0x44)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_ACCENT = RGBColor(0x21, 0x96, 0xF3)
COLOR_GREEN = RGBColor(0x2E, 0x7D, 0x32)
COLOR_RED = RGBColor(0xCC, 0x00, 0x00)
COLOR_ORANGE = RGBColor(0xE6, 0x7E, 0x00)

HEADER_BG = "1B3A5C"
ALT_ROW_BG = "F2F6FA"
CALLOUT_BG = "E8F4FD"
CALLOUT_BORDER = "2196F3"

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Attune_Wearable_Extension_Technical_Report.docx")


# ── Helper Functions ───────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, font_size=10.5, bold=False, italic=False,
                         color=None, alignment=None, space_after=Pt(6),
                         space_before=Pt(0), font_name=None):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name or FONT_BODY
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if alignment:
        p.alignment = alignment
    p.paragraph_format.line_spacing = 1.15
    return p


def add_body(doc, text, space_after=Pt(6)):
    """Add body text paragraph."""
    return add_styled_paragraph(doc, text, font_size=10.5, color=COLOR_BODY,
                                space_after=space_after)


def add_body_rich(doc, segments, space_after=Pt(6)):
    """Add body paragraph with mixed formatting. segments = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = FONT_BODY
        run.font.color.rgb = COLOR_BODY
        run.bold = bold
        run.italic = italic
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text, bold_prefix="", indent_level=0):
    """Add a bullet point with optional bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = FONT_BODY
        run.font.color.rgb = COLOR_BODY
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = FONT_BODY
    run.font.color.rgb = COLOR_BODY
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * indent_level)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_section_header(doc, number, title):
    """Add a section header with page break."""
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {title}")
    run.font.size = Pt(22)
    run.font.name = FONT_BODY
    run.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    # Add thin separator
    sep = doc.add_paragraph()
    run = sep.add_run("─" * 72)
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_SECONDARY
    sep.paragraph_format.space_after = Pt(12)
    return p


def add_subsection(doc, number, title):
    """Add a subsection header."""
    p = doc.add_paragraph()
    run = p.add_run(f"{number}  {title}")
    run.font.size = Pt(14)
    run.font.name = FONT_BODY
    run.bold = True
    run.font.color.rgb = COLOR_HEADING3
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(16)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header row and alternating shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.size = Pt(9.5)
        run.font.name = FONT_BODY
        run.bold = True
        run.font.color.rgb = COLOR_WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, HEADER_BG)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = FONT_BODY
            run.font.color.rgb = COLOR_BODY
            if r % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)

    # Set column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacing after table
    return table


def add_callout_box(doc, title, text):
    """Add a callout box with light background."""
    # Use a single-cell table for the callout
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT_BG)

    # Title
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.size = Pt(10.5)
    run.font.name = FONT_BODY
    run.bold = True
    run.font.color.rgb = COLOR_PRIMARY

    # Body
    p2 = cell.add_paragraph()
    run2 = p2.add_run(text)
    run2.font.size = Pt(10)
    run2.font.name = FONT_BODY
    run2.font.color.rgb = COLOR_BODY

    # Left border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:color="{CALLOUT_BORDER}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    doc.add_paragraph()  # spacing
    return table


def add_code_block(doc, text):
    """Add a monospaced code block."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F5F5F5")
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.name = FONT_MONO
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph()
    return table


def add_architecture_diagram(doc, title, lines):
    """Add a text-based architecture diagram."""
    add_styled_paragraph(doc, title, font_size=11, bold=True, italic=True,
                         color=COLOR_HEADING3, space_after=Pt(4))
    diagram_text = "\n".join(lines)
    add_code_block(doc, diagram_text)


# ── Document Construction ──────────────────────────────────────────────────────

def build_report():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # TITLE PAGE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    for _ in range(6):
        doc.add_paragraph()

    add_styled_paragraph(doc, "ATTUNE STEEL", font_size=16, bold=True,
                         color=COLOR_SECONDARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(4))

    add_styled_paragraph(doc, "Wearable Extension", font_size=36, bold=True,
                         color=COLOR_PRIMARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(2))

    add_styled_paragraph(doc, "Technical Research Report", font_size=20,
                         color=COLOR_SECONDARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(24))

    # Divider
    add_styled_paragraph(doc, "━" * 50, font_size=10,
                         color=COLOR_SECONDARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(24))

    add_styled_paragraph(doc, f"March 2026", font_size=12,
                         color=COLOR_SECONDARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(6))
    add_styled_paragraph(doc, "Version 1.0", font_size=11,
                         color=COLOR_SECONDARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(6))
    add_styled_paragraph(doc, "CONFIDENTIAL", font_size=11, bold=True,
                         color=COLOR_RED,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(6))

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # TABLE OF CONTENTS
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    doc.add_page_break()
    add_styled_paragraph(doc, "Table of Contents", font_size=22, bold=True,
                         color=COLOR_PRIMARY, space_after=Pt(16))

    toc_items = [
        ("1", "Executive Summary"),
        ("2", "The Coverage Gap — Why Desktop-Only Falls Short"),
        ("   2.1", "Current Lucid Pipeline Technical Recap"),
        ("   2.2", "What Is Missed: Analysis of Typical User Day Patterns"),
        ("   2.3", "Design Requirements Derived from Codebase"),
        ("3", "Wearable Device Landscape"),
        ("   3.1", "Device Comparison Matrix"),
        ("   3.2", "Best-Fit Analysis Against Lucid Requirements"),
        ("   3.3", "Audio Quality Deep Dive: MEMS vs Laptop Microphones"),
        ("   3.4", "Edge Processing: What Runs Where"),
        ("4", "Architecture Options"),
        ("   4.1", 'Architecture A: "Store-and-Sync" (Recommended Phase 1)'),
        ("   4.2", 'Architecture B: "Smart Pendant with On-Device VAD"'),
        ("   4.3", 'Architecture C: "Direct Mac Sync"'),
        ("5", "Speaker Verification Across Environments"),
        ("6", "Privacy, Legal, and Ethical Framework"),
        ("7", "Technical Risks and Mitigations"),
        ("8", "Implementation Roadmap"),
        ("9", "Recommendation"),
        ("A", "Appendix A: Data Format Specifications"),
        ("B", "Appendix B: BLE Audio Streaming Protocol"),
        ("C", "Appendix C: Competitive Landscape"),
    ]
    for num, title in toc_items:
        is_sub = num.startswith("   ")
        p = doc.add_paragraph()
        run = p.add_run(f"{num.strip()}    {title}")
        run.font.size = Pt(10 if is_sub else 11)
        run.font.name = FONT_BODY
        run.font.color.rgb = COLOR_SECONDARY if is_sub else COLOR_BODY
        if not is_sub:
            run.bold = True
        p.paragraph_format.space_after = Pt(2)
        if is_sub:
            p.paragraph_format.left_indent = Cm(1.0)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 1. EXECUTIVE SUMMARY
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_section_header(doc, "1", "Executive Summary")

    add_body(doc, (
        "Attune Steel captures vocal biomarkers through the Lucid desktop application, analyzing "
        "speech patterns via the DAM (Depression/Anxiety Model) to deliver continuous mental wellness "
        "monitoring. However, desktop-only capture creates a fundamental coverage gap: users are at "
        "their computers for only a fraction of their waking hours, meaning 40–60% of daily voice "
        "production — social conversations, phone calls, meetings away from the desk, commutes, "
        "and leisure interactions — generates zero wellness data."
    ))

    add_body(doc, (
        "This report investigates whether wearable devices can fill this gap by capturing voice "
        "throughout the day and syncing audio back to the desktop for analysis. The central finding "
        "is that wearable voice capture is technically feasible, but all DAM analysis must remain "
        "on the desktop. The DAM model (702 MB, Whisper-Small EN backbone with LoRA fine-tuning) "
        "requires 500ms–2s of inference time on laptop CPU — far exceeding the capabilities of any "
        "wearable processor and most smartphone SoCs for real-time use."
    ))

    add_callout_box(doc, "Core Recommendation",
        "Begin with the Omi AI Dev Kit 2 ($70, open-source, BLE audio streaming) paired with a "
        "minimal phone companion app that performs Voice Activity Detection and speaker verification. "
        "Audio segments confirmed as belonging to the user are synced via WiFi to the desktop Lucid "
        "application, where the full DAM pipeline runs unchanged. This approach maximizes coverage "
        "while preserving analysis accuracy and privacy guarantees."
    )

    add_body(doc, (
        "The recommended phased approach begins with a proof-of-concept (Months 1–3, ~60–80 hours) "
        "to validate that wearable-captured audio produces DAM scores within ±15% of desktop-captured "
        "scores. If validated, a companion app MVP (Months 4–6, ~120–160 hours) adds automated "
        "speaker verification and WiFi sync. The production phase (Months 7–12, ~200–300 hours) "
        "brings on-device VAD, batch processing UI, and privacy review."
    ))

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 2. THE COVERAGE GAP
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_section_header(doc, "2", "The Coverage Gap — Why Desktop-Only Falls Short")

    add_body(doc, (
        "Attune Steel's current architecture is tightly coupled to the desktop environment. "
        "While this provides excellent audio quality and computational resources for DAM inference, "
        "it inherently limits data collection to periods when the user is at their computer with "
        "the application running. This section quantifies the gap and derives requirements for "
        "a wearable extension."
    ))

    # 2.1
    add_subsection(doc, "2.1", "Current Lucid Pipeline Technical Recap")

    add_body(doc, (
        "The Lucid application runs as a macOS menubar app (Electron + Python backend via FastAPI). "
        "Audio flows through a multi-stage pipeline before producing wellness scores:"
    ))

    add_architecture_diagram(doc, "Figure 2.1 — Current Lucid Pipeline", [
        "┌─────────────────────────────────────────────────────────────────────┐",
        "│                    LUCID DESKTOP PIPELINE                          │",
        "├─────────────────────────────────────────────────────────────────────┤",
        "│                                                                    │",
        "│  Microphone ──► sounddevice (16kHz, mono, 32ms chunks, float32)    │",
        "│       │                                                            │",
        "│       ▼                                                            │",
        "│  Silero VAD ──► Speech probability per chunk (threshold: 0.40)     │",
        "│       │                                                            │",
        "│       ▼                                                            │",
        "│  Speaker Gate ──► ECAPA-TDNN (192-dim embeddings)                  │",
        "│       │           Cosine similarity vs enrolled centroid            │",
        "│       │           Base threshold: 0.28, momentum: 0.24             │",
        "│       │           Sandwich recovery: 0.24 (borderline accepts)     │",
        "│       ▼                                                            │",
        "│  Speech Buffer ──► Accumulate verified speech                      │",
        "│       │            Soft trigger: 30s, Preferred: 60s, Max: 90s     │",
        "│       ▼                                                            │",
        "│  DAM 3.1 ──► Whisper-Small EN encoder + LoRA (702 MB)              │",
        "│       │      Depression/anxiety raw scores                         │",
        "│       │      Inference: 500ms–2s on CPU                            │",
        "│       ▼                                                            │",
        "│  Score Mapper ──► PHQ-9 / GAD-7 clinical equivalents               │",
        "│       │          Piecewise-linear isotonic regression               │",
        "│       │          Confidence intervals + indeterminate zones         │",
        "│       ▼                                                            │",
        "│  Acoustic Features ──► F0, jitter, shimmer, voice breaks           │",
        "│       │                speech rate, spectral entropy, ZCR           │",
        "│       ▼                                                            │",
        "│  Score Engine ──► Stress, mood, energy, calm scores                │",
        "│       │          EMA smoothing (alpha=0.4)                          │",
        "│       ▼                                                            │",
        "│  SQLite DB ──► readings, daily_summaries                           │",
        "│                burnout.db (WAL mode)                               │",
        "└─────────────────────────────────────────────────────────────────────┘",
    ])

    add_body(doc, (
        "Key technical parameters: Audio is captured at 16,000 Hz sample rate, mono channel, "
        "in 32ms chunks (512 samples). The VAD requires this exact chunk size (Silero v6.2.0 "
        "minimum). Speaker verification uses SpeechBrain's ECAPA-TDNN model producing 192-dimensional "
        "L2-normalized embeddings, with cosine similarity scoring against an enrolled centroid. "
        "The speaker gate operates on 2-second mini-buffers of accumulated speech."
    ))

    add_body(doc, (
        "Enrollment involves three 10-second samples (neutral, animated, calm moods), each augmented "
        "via Opus codec roundtrip at 32kbps and 16kbps to improve robustness against audio quality "
        "variation. This codec augmentation — already present in the codebase — provides built-in "
        "tolerance for the frequency response differences a wearable microphone would introduce."
    ))

    # 2.2
    add_subsection(doc, "2.2", "What Is Missed: Analysis of Typical User Day Patterns")

    add_body(doc, (
        "Consider a typical knowledge worker's day. Desktop capture windows are shown in gray; "
        "uncaptured voice activity appears in white:"
    ))

    add_table(doc,
        ["Time Window", "Activity", "Voice Type", "Captured?", "Wellness Signal"],
        [
            ["6:30–8:00 AM", "Morning routine, family interaction", "Casual/emotional", "No", "Baseline mood, family stress"],
            ["8:00–8:30 AM", "Commute (phone calls, podcast)", "Phone voice", "No", "Morning anxiety, social tone"],
            ["8:30–10:00 AM", "Desk work, Slack huddles", "Professional", "Yes", "Work stress onset"],
            ["10:00–10:30 AM", "Coffee break, hallway chat", "Social/relaxed", "No", "Social engagement quality"],
            ["10:30–12:00 PM", "Desk work, video calls", "Professional", "Yes", "Sustained work stress"],
            ["12:00–1:00 PM", "Lunch, social conversation", "Casual", "No", "Midday recovery quality"],
            ["1:00–3:00 PM", "Desk work", "Professional", "Yes", "Afternoon patterns"],
            ["3:00–3:15 PM", "Walk, phone call", "Mobile", "No", "Afternoon energy level"],
            ["3:15–5:30 PM", "Desk work, meetings", "Professional", "Yes", "Late-day fatigue"],
            ["5:30–6:00 PM", "Commute", "Phone/silent", "No", "End-of-day state"],
            ["6:00–10:00 PM", "Evening, family, social", "Emotional/casual", "No", "Evening recovery, social health"],
        ],
        col_widths=[3.0, 3.5, 2.5, 1.8, 4.0]
    )

    add_body(doc, (
        "In this scenario, desktop capture covers roughly 6.75 hours of an approximately 15.5-hour "
        "waking day — 43% coverage. More critically, all captured voice is professional context. "
        "The system has zero visibility into emotional conversations, social interactions, phone "
        "calls, and evening recovery patterns. This creates a systematic bias: wellness scores "
        "reflect only work-self vocal patterns, not the user's full vocal biomarker profile."
    ))

    add_callout_box(doc, "The Bias Problem",
        "Depression and anxiety manifest differently across social contexts. A user may present "
        "controlled, professional speech at their desk while exhibiting significant vocal biomarker "
        "changes during personal conversations. Desktop-only capture systematically misses the "
        "contexts where mental health signals may be most pronounced — creating a false sense of "
        "stability or missing early warning signs entirely."
    )

    # 2.3
    add_subsection(doc, "2.3", "Design Requirements Derived from Codebase")

    add_body(doc, (
        "Any wearable extension must satisfy requirements derived directly from the existing "
        "Lucid codebase. These are non-negotiable constraints that preserve analysis accuracy:"
    ))

    add_table(doc,
        ["Requirement", "Source", "Value", "Rationale"],
        [
            ["Sample rate", "app_config.py", "16,000 Hz", "DAM Whisper encoder expects 16kHz input"],
            ["Audio format", "audio_capture.py", "PCM float32 / 16-bit WAV", "Pipeline normalizes to [-1, 1] float32"],
            ["Minimum duration", "Speech buffer triggers", "30 seconds of speech", "Soft trigger; 60s preferred for reliability"],
            ["Speaker verification", "speaker_gate.py", "ECAPA-TDNN cosine > 0.28", "Must verify user identity before analysis"],
            ["Local-only processing", "Architecture constraint", "No cloud audio", "Privacy: raw audio never leaves user devices"],
            ["Temporal ordering", "EMA smoothing (alpha=0.4)", "Chronological sequence", "Readings must interleave correctly for EMA"],
            ["Mono channel", "audio_capture.py", "1 channel", "DAM trained on mono; stereo must be downmixed"],
        ],
        col_widths=[3.0, 3.0, 3.5, 5.5]
    )

    add_body(doc, (
        "Additionally, the existing codec augmentation during enrollment (Opus at 32kbps and 16kbps) "
        "provides partial tolerance for audio quality degradation — but was designed for VoIP scenarios, "
        "not the specific frequency response characteristics of wearable MEMS microphones. A calibration "
        "study will be needed to quantify the domain gap."
    ))

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 3. WEARABLE DEVICE LANDSCAPE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_section_header(doc, "3", "Wearable Device Landscape")

    add_body(doc, (
        "The AI wearable market has seen rapid evolution through 2024–2026, with several devices "
        "entering and exiting the market. This section provides an honest assessment of eight "
        "devices evaluated against Attune Steel's specific requirements: raw audio access at 16kHz, "
        "developer API availability, speaker diarization capability, and privacy-preserving "
        "architecture."
    ))

    # 3.1
    add_subsection(doc, "3.1", "Device Comparison Matrix")

    add_table(doc,
        ["Device", "Price", "Battery", "Storage", "Open Source", "API/SDK", "Diarization", "Status"],
        [
            ["Omi AI\n(Dev Kit 2)", "$70", "10–14 hr", "Streams\nto phone", "Full\n(MIT)", "Full API\n+ MCP", "Yes", "Active"],
            ["Plaud\nNotePin", "$149", "20 hr", "64 GB\non-device", "No", "Developer\nPlatform\n(Oct 2025)", "Yes", "Active"],
            ["Plaud\nNotePin S", "$119", "24 hr", "32 GB\non-device", "No", "Developer\nPlatform", "Yes", "Active"],
            ["Friend AI\n(Schiffmann)", "$99–129", "~15 hr", "Streams\nto phone", "GPL\n(partial)", "Limited", "Via Gemini", "Active"],
            ["Bee\n(Amazon)", "$50 +\n$19/mo", "N/A", "No audio\nstored", "No", "Limited", "Partial\n(manual)", "Active"],
            ["Soundcore\nWork", "$160", "8 hr\n(32 w/case)", "8 GB", "No", "No", "Yes\n(up to 4)", "Active"],
            ["DIY\nnRF52840", "$18–22\n(+ BOM)", "Custom", "BLE\nstream", "Full", "Full", "DIY", "N/A"],
            ["Limitless", "$99\n(was)", "~100 hr", "Cloud", "No", "N/A", "Yes", "DEAD\n(Meta, Dec 2025)"],
            ["Humane\nAI Pin", "$699\n(was)", "N/A", "N/A", "No", "No", "N/A", "DEAD\n(HP, Feb 2025)"],
        ],
        col_widths=[2.2, 1.5, 1.5, 1.8, 1.5, 1.8, 1.8, 1.5]
    )

    add_body_rich(doc, [
        ("Note: ", True, False),
        ("Omi AI was formerly called 'Friend' by Based Hardware (Nik Shevchenko), distinct from ", False, False),
        ("the other ", False, False),
        ("Friend AI pendant by Avi Schiffmann. Based Hardware rebranded to Omi after the naming dispute. ", False, False),
        ("Limitless was acquired by Meta (December 2025) and discontinued; pendant sales stopped. ", False, False),
        ("Humane AI Pin assets were purchased by HP for $116M (February 2025); devices shut down February 28, 2025. ", False, False),
        ("Bee was acquired by Amazon (July 2025) and continues as an Amazon product.", False, False),
    ])

    # 3.2
    add_subsection(doc, "3.2", "Best-Fit Analysis Against Lucid Requirements")

    add_body(doc, "Each device is evaluated against the five critical requirements for Attune integration:")

    add_table(doc,
        ["Requirement", "Omi AI", "Plaud NotePin", "Friend AI", "Bee", "DIY nRF52840"],
        [
            ["Raw 16kHz audio access", "Yes (BLE stream)", "Yes (USB export)", "Partial (API)", "No (deleted after\ntranscription)", "Yes (full control)"],
            ["Developer API/SDK", "Full open-source\napp + firmware", "Developer Platform\n(closed source)", "Limited", "Limited API,\nno raw audio", "Full (bare metal)"],
            ["On-device storage", "No (phone relay)", "64 GB internal", "No (phone relay)", "No (cloud only)", "External flash\n(add-on)"],
            ["Speaker isolation", "Via app diarization", "Built-in labels", "Limited", "Built-in but\naudio deleted", "Manual\nimplementation"],
            ["Privacy compatible\n(no cloud audio)", "Yes (local\nprocessing option)", "Yes (offline\nrecording)", "No (streams\nto cloud by default)", "No (cloud\nprocessing)", "Yes (fully\noffline)"],
        ],
        col_widths=[2.8, 2.5, 2.5, 2.3, 2.5, 2.5]
    )

    add_callout_box(doc, "Verdict: Omi AI as Phase 1, Plaud NotePin as Phase 2",
        "Omi AI offers the best Phase 1 fit: full open-source access (firmware, app, API), BLE audio "
        "streaming at 16kHz, and an active developer community with 250+ third-party apps. The Dev Kit 2 "
        "at $70 minimizes PoC risk. Plaud NotePin is the strongest Phase 2 candidate for a "
        "no-phone workflow: 64 GB on-device storage, 20-hour battery, and USB export — enabling "
        "end-of-day dock-and-sync without requiring a companion phone app. Plaud's Developer Platform "
        "(launched October 2025) provides SOC 2, HIPAA, and GDPR-compliant APIs for integration."
    )

    # 3.3
    add_subsection(doc, "3.3", "Audio Quality Deep Dive: MEMS vs Laptop Microphones")

    add_body(doc, (
        "Wearable devices universally use MEMS (Micro-Electro-Mechanical Systems) microphones, "
        "which differ from laptop microphones in ways that directly affect DAM model accuracy:"
    ))

    add_table(doc,
        ["Characteristic", "Laptop Mic (current)", "Wearable MEMS", "Impact on DAM"],
        [
            ["Frequency response", "80 Hz – 15 kHz\n(relatively flat)", "100 Hz – 10 kHz\n(high-frequency rolloff)", "Reduced spectral detail\nabove 10 kHz"],
            ["Signal-to-noise ratio", "58–65 dB", "42–50 dB", "More background noise\nin features"],
            ["Proximity to mouth", "40–60 cm (desk)", "15–30 cm (chest/ear)", "Stronger direct signal,\nless room reverb"],
            ["Environment", "Quiet office\n(typically)", "Variable (street,\ncafe, outdoors)", "Domain shift in\nacoustic features"],
            ["Wind/body noise", "Minimal", "Clothing rustle,\nwind, body movement", "Spurious low-frequency\nenergy"],
            ["Codec compression", "None (direct PCM)", "BLE: SBC/LC3\nStored: PCM/Opus", "Potential quality loss\nduring BLE streaming"],
        ],
        col_widths=[3.0, 3.5, 3.5, 3.5]
    )

    add_body(doc, (
        "The DAM model was trained on audio from controlled recording environments. Wearable audio "
        "introduces a domain gap that could shift model outputs. The Whisper backbone provides some "
        "robustness (trained on 680,000 hours of diverse audio), but the LoRA fine-tuning for "
        "depression/anxiety detection was performed on cleaner audio. A calibration study comparing "
        "simultaneous wearable and desktop recordings is essential before production deployment."
    ))

    # 3.4
    add_subsection(doc, "3.4", "Edge Processing: What Runs Where")

    add_body(doc, (
        "A critical architectural question is which pipeline components can run on mobile/wearable "
        "hardware versus requiring the desktop. The following analysis maps each component to its "
        "minimum viable hardware:"
    ))

    add_table(doc,
        ["Component", "Model Size", "Inference Time\n(Desktop CPU)", "Phone Feasible?", "Wearable Feasible?"],
        [
            ["Silero VAD", "~2 MB", "2–5 ms / chunk", "Yes (real-time)", "Possible (nRF52840\nhas limited RAM)"],
            ["Picovoice Cobra\n(alternative VAD)", "~100 KB", "<1 ms / chunk", "Yes", "Yes (designed for\nmicrocontrollers)"],
            ["ECAPA-TDNN\n(speaker verify)", "~25 MB", "50–100 ms /\n2s segment", "Yes (batch, not\nreal-time)", "No (needs PyTorch\nor ONNX runtime)"],
            ["DAM 3.1\n(Whisper + LoRA)", "702 MB", "500ms–2s /\n30–60s audio", "No (too slow,\ntoo large)", "No (impossible)"],
            ["Acoustic features\n(librosa)", "~15 MB\n(with deps)", "50–200 ms", "Possible but\nnot needed", "No"],
            ["Score mapping\n(isotonic regression)", "<1 MB", "<1 ms", "Yes (trivial)", "Possible"],
        ],
        col_widths=[2.8, 2.0, 2.5, 2.8, 3.0]
    )

    add_callout_box(doc, "The Fundamental Constraint",
        "The DAM model (702 MB) cannot run on any wearable device and would drain phone battery "
        "unacceptably for real-time use. However, VAD and speaker verification CAN run on a phone, "
        "enabling a 'capture-and-filter' architecture where the phone acts as an intelligent relay — "
        "filtering non-speech and non-user audio before syncing only relevant segments to the desktop."
    )

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 4. ARCHITECTURE OPTIONS
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_section_header(doc, "4", "Architecture Options")

    add_body(doc, (
        "Three progressively ambitious architectures are proposed, each building on the previous. "
        "All three share a core principle: the DAM model runs only on the desktop. They differ in "
        "where filtering (VAD + speaker verification) occurs and how audio reaches the desktop."
    ))

    # Architecture A
    add_subsection(doc, "4.1", 'Architecture A: "Store-and-Sync" (Recommended Phase 1)')

    add_body(doc, (
        "The simplest viable architecture. The wearable streams raw audio via BLE to a companion "
        "phone app. The phone performs VAD and speaker verification, then stores verified speech "
        "segments locally. When the phone is on the same WiFi network as the desktop, segments "
        "sync automatically to a watched directory, where Lucid's existing pipeline processes them."
    ))

    add_architecture_diagram(doc, "Figure 4.1 — Architecture A Data Flow", [
        "┌──────────┐     BLE      ┌──────────────────┐     WiFi     ┌──────────────┐",
        "│          │   16kHz PCM   │                  │    Sync      │              │",
        "│  Omi AI  │ ───────────► │   Phone App      │ ──────────► │  Desktop     │",
        "│ Wearable │              │                  │              │  Lucid       │",
        "└──────────┘              │  ┌─────────────┐ │              │              │",
        "                          │  │ Silero VAD  │ │              │ ┌──────────┐ │",
        "                          │  │ (2 MB)      │ │              │ │ wearable │ │",
        "                          │  └──────┬──────┘ │              │ │ _ingest  │ │",
        "                          │         ▼        │              │ │ .py      │ │",
        "                          │  ┌─────────────┐ │              │ └────┬─────┘ │",
        "                          │  │ ECAPA-TDNN  │ │              │      ▼       │",
        "                          │  │ Speaker     │ │              │ ┌──────────┐ │",
        "                          │  │ Verify      │ │              │ │ Existing │ │",
        "                          │  │ (25 MB)     │ │              │ │ DAM      │ │",
        "                          │  └──────┬──────┘ │              │ │ Pipeline │ │",
        "                          │         ▼        │              │ └──────────┘ │",
        "                          │  ┌─────────────┐ │              │              │",
        "                          │  │ Local WAV   │ │              │              │",
        "                          │  │ Storage     │ │              │              │",
        "                          │  └─────────────┘ │              │              │",
        "                          └──────────────────┘              └──────────────┘",
    ])

    add_body(doc, "Storage requirements for the phone app:")

    add_table(doc,
        ["Scenario", "Speech / Day", "Raw Audio Size", "After VAD Filter", "After Speaker Gate"],
        [
            ["Light vocal day", "2 hr speech", "230 MB (16kHz PCM)", "~115 MB (50% speech)", "~58 MB (50% user)"],
            ["Moderate vocal day", "4 hr speech", "460 MB", "~230 MB", "~115 MB"],
            ["Heavy vocal day", "6 hr speech", "690 MB", "~345 MB", "~173 MB"],
        ],
        col_widths=[3.0, 2.5, 3.0, 3.0, 3.0]
    )

    add_body(doc, "Codebase modifications required:")

    add_bullet(doc, "New module in python/backend/ that watches a sync directory for incoming "
               "WAV files, validates format (16kHz, mono), and feeds them into the existing "
               "analysis pipeline with a source='wearable' tag.", bold_prefix="wearable_ingest.py: ")
    add_bullet(doc, "Add 'source' column (TEXT, default 'desktop') to readings table. Add "
               "'wearable_sync_timestamp' column for deduplication.", bold_prefix="Database migration: ")
    add_bullet(doc, "Interleave wearable readings chronologically with desktop readings before "
               "applying EMA smoothing (alpha=0.4). Currently EMA uses insertion order.",
               bold_prefix="EMA ordering: ")
    add_bullet(doc, "Show wearable vs desktop readings with distinct visual treatment (e.g., "
               "different icon or subtle background).", bold_prefix="Dashboard UI: ")

    add_body_rich(doc, [
        ("Pros: ", True, False),
        ("Minimal desktop codebase changes. Proven DAM analysis unchanged. Phone does heavy "
         "filtering, reducing sync volume by ~75%. Omi's open-source app provides reference "
         "implementation for BLE streaming.", False, False),
    ])

    add_body_rich(doc, [
        ("Cons: ", True, False),
        ("Requires phone to be nearby (BLE range ~10m). Phone battery drain from continuous "
         "BLE reception and VAD processing. No readings until WiFi sync completes. User must "
         "carry both wearable and phone.", False, False),
    ])

    # Architecture B
    add_subsection(doc, "4.2", 'Architecture B: "Smart Pendant with On-Device VAD" (Phase 2)')

    add_body(doc, (
        "Extends Architecture A by moving VAD to the wearable itself. The wearable runs a "
        "lightweight edge VAD (Picovoice Cobra, ~100 KB model) and only streams speech-positive "
        "segments to the phone via BLE. This dramatically reduces BLE bandwidth, phone battery "
        "drain, and phone storage requirements."
    ))

    add_architecture_diagram(doc, "Figure 4.2 — Architecture B Data Flow", [
        "┌────────────────────┐     BLE       ┌──────────────────┐    WiFi    ┌────────┐",
        "│    Omi AI          │  Speech-only   │   Phone App      │   Sync    │Desktop │",
        "│  ┌──────────────┐  │  ──────────►  │  ┌────────────┐  │ ───────► │ Lucid  │",
        "│  │ Picovoice    │  │               │  │ ECAPA-TDNN │  │          │        │",
        "│  │ Cobra VAD    │  │               │  │ Speaker    │  │          │  DAM   │",
        "│  │ (~100 KB)    │  │               │  │ Verify     │  │          │Pipeline│",
        "│  └──────────────┘  │               │  └────────────┘  │          │        │",
        "└────────────────────┘               └──────────────────┘          └────────┘",
    ])

    add_body(doc, "Impact on battery and bandwidth:")

    add_table(doc,
        ["Metric", "Architecture A\n(no on-device VAD)", "Architecture B\n(on-device VAD)", "Improvement"],
        [
            ["BLE streaming time", "10–14 hr\n(continuous)", "3–5 hr\n(speech-only)", "3–5x reduction"],
            ["Phone battery drain", "~15–20% / day\n(BLE + VAD)", "~5–8% / day\n(BLE receive only)", "2–3x reduction"],
            ["Phone storage / day", "~115 MB\n(after phone VAD)", "~115 MB\n(same, but pre-filtered)", "Similar total"],
            ["Wearable battery", "10–14 hr", "14–18 hr\n(less BLE tx)", "~30% improvement"],
        ],
        col_widths=[3.5, 3.5, 3.5, 3.0]
    )

    add_body_rich(doc, [
        ("Pros: ", True, False),
        ("Significantly better battery life on both wearable and phone. Lower BLE bandwidth "
         "reduces audio dropout risk. Wearable becomes useful even without phone nearby (buffer "
         "speech segments on-device for later sync).", False, False),
    ])

    add_body_rich(doc, [
        ("Cons: ", True, False),
        ("Requires custom firmware development for the wearable (or fork of Omi firmware). "
         "Picovoice Cobra requires a license for commercial use ($variable/device). Edge VAD "
         "has lower accuracy than Silero (more false positives/negatives). Testing and debugging "
         "firmware is significantly harder than app development.", False, False),
    ])

    # Architecture C
    add_subsection(doc, "4.3", 'Architecture C: "Direct Mac Sync" (Long-Term)')

    add_body(doc, (
        "The most privacy-preserving and simplest user experience: the wearable records all day "
        "to on-device storage, and the user docks/connects it to their Mac at the end of the day "
        "for batch processing. No phone required. This maps naturally to the Plaud NotePin model "
        "(64 GB storage, 20-hour battery, USB-C export)."
    ))

    add_architecture_diagram(doc, "Figure 4.3 — Architecture C Data Flow", [
        "┌────────────────────┐                              ┌──────────────────────┐",
        "│    Plaud NotePin   │        USB-C / Dock          │   Desktop Lucid      │",
        "│                    │                              │                      │",
        "│  ┌──────────────┐  │     End-of-day transfer      │  ┌────────────────┐  │",
        "│  │ 64 GB Flash  │  │  ────────────────────────►  │  │ wearable_      │  │",
        "│  │ 20 hr battery│  │                              │  │ ingest.py      │  │",
        "│  │ PCM recording│  │                              │  │                │  │",
        "│  └──────────────┘  │                              │  │  ┌───────────┐ │  │",
        "│                    │                              │  │  │ VAD       │ │  │",
        "│  No processing     │                              │  │  │ Speaker   │ │  │",
        "│  Just record       │                              │  │  │ Gate      │ │  │",
        "│                    │                              │  │  │ DAM       │ │  │",
        "│                    │                              │  │  └───────────┘ │  │",
        "│                    │                              │  └────────────────┘  │",
        "└────────────────────┘                              └──────────────────────┘",
    ])

    add_body(doc, "Storage math for all-day recording:")

    add_table(doc,
        ["Recording Duration", "16kHz 16-bit PCM", "With Opus 32kbps\ncompression", "Plaud 64 GB\ncapacity"],
        [
            ["8 hours", "922 MB", "~115 MB", "~556 days"],
            ["12 hours", "1.38 GB", "~173 MB", "~370 days"],
            ["16 hours (full day)", "1.84 GB", "~230 MB", "~278 days"],
        ],
        col_widths=[3.5, 3.5, 3.5, 3.5]
    )

    add_body_rich(doc, [
        ("Pros: ", True, False),
        ("No phone required — simplest user experience. Most private — no BLE streaming, "
         "no cloud, just a file transfer. Plaud NotePin has proven hardware with 64 GB storage. "
         "Batch processing can run overnight without user interaction.", False, False),
    ])

    add_body_rich(doc, [
        ("Cons: ", True, False),
        ("Delayed-only readings — no real-time wearable wellness scores until dock sync. "
         "Desktop must run full VAD + speaker gate on hours of raw audio (computationally "
         "intensive batch job). Plaud NotePin is closed-source; API access is through their "
         "Developer Platform with limited raw audio export guarantees. Higher device cost ($149).",
         False, False),
    ])

    # Architecture comparison
    add_subsection(doc, "4.4", "Architecture Comparison Summary")

    add_table(doc,
        ["Factor", "A: Store-and-Sync", "B: Smart Pendant", "C: Direct Mac Sync"],
        [
            ["Complexity", "Low", "High", "Medium"],
            ["Phone required?", "Yes", "Yes (reduced role)", "No"],
            ["Real-time scores?", "Near-real-time\n(WiFi sync delay)", "Near-real-time\n(WiFi sync delay)", "No (end-of-day\nbatch only)"],
            ["Privacy level", "High\n(local phone + desktop)", "High\n(local processing)", "Highest\n(no wireless audio)"],
            ["Battery impact", "Moderate\n(phone + wearable)", "Low\n(on-device VAD)", "None\n(phone not used)"],
            ["Dev effort (PoC)", "60–80 hrs", "120–160 hrs", "40–60 hrs"],
            ["Best device fit", "Omi AI ($70)", "Omi AI (custom FW)", "Plaud NotePin ($149)"],
            ["Recommended phase", "Phase 1", "Phase 2", "Phase 2 alternative"],
        ],
        col_widths=[3.0, 3.5, 3.5, 3.5]
    )

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 5. SPEAKER VERIFICATION ACROSS ENVIRONMENTS
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_section_header(doc, "5", "Speaker Verification Across Environments")

    add_body(doc, (
        "The speaker gate is Lucid's first line of defense against analyzing non-user audio. "
        "Currently calibrated for desktop microphone characteristics, it faces new challenges "
        "when verifying speech captured through wearable MEMS microphones in varied acoustic "
        "environments."
    ))

    add_subsection(doc, "5.1", "The Cross-Device Domain Gap")

    add_body(doc, (
        "ECAPA-TDNN produces 192-dimensional speaker embeddings. When the enrollment centroid "
        "is computed from desktop microphone samples, cosine similarity scores for the same "
        "speaker recorded through a wearable microphone will be systematically lower due to:"
    ))

    add_bullet(doc, "Different frequency response curves between microphone types, "
               "causing embedding vectors to shift in the 192-dimensional space.",
               bold_prefix="Frequency response mismatch: ")
    add_bullet(doc, "Wearable placement (chest, collar) captures more body-conducted "
               "sound and less air-conducted sound than a desk microphone.",
               bold_prefix="Proximity and body conduction: ")
    add_bullet(doc, "Street noise, wind, other speakers, and room acoustics affect "
               "the spectral characteristics of the recorded speech.",
               bold_prefix="Environmental noise: ")
    add_bullet(doc, "BLE audio may be encoded with SBC or LC3 codecs before reaching "
               "the phone, introducing compression artifacts.",
               bold_prefix="BLE codec artifacts: ")

    add_body(doc, (
        "Estimated impact: cosine similarity scores from wearable audio will be 0.05–0.15 lower "
        "than equivalent desktop scores. Given the base threshold of 0.28, this could push many "
        "legitimate segments below threshold, causing excessive false rejections."
    ))

    add_subsection(doc, "5.2", "Cross-Device Enrollment Strategy")

    add_body(doc, (
        "The solution is to extend the existing enrollment process to include wearable-sourced "
        "samples. The current enrollment already captures three moods (neutral, animated, calm) "
        "with codec augmentation (Opus at 32kbps and 16kbps). The proposed extension:"
    ))

    add_table(doc,
        ["Enrollment Sample", "Source", "Purpose"],
        [
            ["neutral", "Desktop mic", "Baseline embedding (existing)"],
            ["animated", "Desktop mic", "Excited speech variant (existing)"],
            ["calm", "Desktop mic", "Quiet speech variant (existing)"],
            ["neutral_codec32", "Desktop mic + Opus 32k", "VoIP robustness (existing)"],
            ["neutral_codec16", "Desktop mic + Opus 16k", "VoIP robustness (existing)"],
            ["wearable_neutral", "Wearable mic", "Wearable baseline (new)"],
            ["wearable_animated", "Wearable mic", "Wearable excited variant (new)"],
            ["wearable_calm", "Wearable mic", "Wearable quiet variant (new)"],
            ["wearable_outdoor", "Wearable mic (outside)", "Environmental noise variant (new)"],
        ],
        col_widths=[3.5, 3.5, 5.0]
    )

    add_body(doc, (
        "The centroid computation would use all enrollment samples (desktop + wearable), producing "
        "a blended embedding that captures the user's voice across microphone types. Alternatively, "
        "separate centroids could be maintained per source, with the speaker gate selecting the "
        "appropriate centroid based on the audio source."
    ))

    add_subsection(doc, "5.3", "Adaptive Centroid Considerations for Wearable Segments")

    add_body(doc, (
        "The current adaptive centroid system maintains a sliding window of 50 recent high-confidence "
        "embeddings (similarity > 0.75) and blends 75% adapted + 25% enrollment centroid. For wearable "
        "segments, this poses a risk: if the centroid drifts toward wearable characteristics, desktop "
        "verification accuracy could degrade, and vice versa."
    ))

    add_body(doc, "Recommended approach:")

    add_bullet(doc, "Maintain separate adaptive windows for desktop and wearable sources.",
               bold_prefix="Dual-window adaptation: ")
    add_bullet(doc, "Do not update adaptive centroid from wearable segments until cross-device "
               "enrollment is completed and validated.",
               bold_prefix="Gated adaptation: ")
    add_bullet(doc, "Desktop similarity > 0.28, Wearable similarity > 0.22 "
               "(accounting for the systematic offset). This is above the existing absolute "
               "floor of 0.22.",
               bold_prefix="Source-specific thresholds: ")

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 6. PRIVACY, LEGAL, AND ETHICAL FRAMEWORK
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_section_header(doc, "6", "Privacy, Legal, and Ethical Framework")

    add_body(doc, (
        "Always-on audio capture in public and private spaces raises significant privacy, legal, "
        "and ethical concerns. This section outlines the regulatory landscape and recommends a "
        "privacy-preserving architecture."
    ))

    add_subsection(doc, "6.1", "Recording Consent Laws")

    add_body(doc, (
        "The United States has a patchwork of recording consent laws. The distinction between "
        "one-party and all-party (two-party) consent states is critical for a wearable that "
        "captures audio in social settings:"
    ))

    add_table(doc,
        ["Consent Type", "States", "Requirement"],
        [
            ["All-party consent\n(11 core states)", "California, Delaware,\nFlorida, Illinois,\n"
             "Maryland, Massachusetts,\nMontana, Nevada,\nNew Hampshire,\n"
             "Pennsylvania, Washington", "All participants in a conversation must consent "
             "to being recorded. Recording without consent is a criminal offense in most "
             "of these states."],
            ["Mixed rules\n(3 states)", "Connecticut (all-party\nfor phone, one-party\nin-person),\n"
             "Missouri and Oregon\n(all-party in-person,\none-party for phone)",
             "Rules differ between phone calls and in-person "
             "conversations. Wearable captures are typically classified "
             "as in-person recording."],
            ["One-party consent", "All other states\n(~36) + federal law", "Only one participant "
             "(the recorder) needs to know about the recording. The user wearing the device "
             "satisfies this requirement."],
        ],
        col_widths=[3.0, 3.5, 7.0]
    )

    add_body(doc, (
        "In all-party consent states (including California, where many tech workers are based), "
        "the user wearing the device must inform all conversation participants that recording "
        "is occurring. This is a significant UX and social barrier."
    ))

    add_subsection(doc, "6.2", "GDPR and International Privacy Law")

    add_body(doc, (
        "Under the GDPR (applicable to EU/EEA users), voice recordings constitute biometric data "
        "— a 'special category' of personal data under Article 9. Processing biometric data "
        "requires explicit consent from each data subject. Additionally:"
    ))

    add_bullet(doc, "The user must be able to request deletion of all voice data "
               "associated with them.", bold_prefix="Right to erasure (Art. 17): ")
    add_bullet(doc, "Processing must be limited to the stated purpose "
               "(wellness monitoring). Voice data cannot be repurposed.",
               bold_prefix="Purpose limitation (Art. 5(1)(b)): ")
    add_bullet(doc, "Only the minimum necessary voice data should be retained. "
               "Raw audio should be deleted after feature extraction.",
               bold_prefix="Data minimization (Art. 5(1)(c)): ")
    add_bullet(doc, "Required for processing that involves systematic "
               "monitoring of publicly accessible areas.",
               bold_prefix="Data Protection Impact Assessment: ")

    add_subsection(doc, "6.3", "Recommended Privacy Architecture")

    add_body(doc, (
        "The proposed privacy architecture is designed to minimize legal risk while maximizing "
        "wellness data collection:"
    ))

    add_table(doc,
        ["Principle", "Implementation", "Legal Basis"],
        [
            ["User-only audio", "Speaker gate filters out non-user\naudio before storage",
             "Only user's own voice is analyzed;\nothers' voices are immediately discarded"],
            ["Triggered recording\n(not always-on)", "VAD activates recording only during\nspeech; "
             "silence is never stored",
             "Reduces scope of recording;\nminimizes bystander capture"],
            ["Immediate deletion\nof non-user audio", "Audio failing speaker verification\nis "
             "deleted within seconds, never\nsynced to desktop",
             "Non-user biometric data is never\nstored or processed"],
            ["On-device only", "Raw audio never leaves user's\npersonal devices (phone + laptop)",
             "No cloud processing eliminates\nthird-party data processor risks"],
            ["LED indicator", "Wearable LED illuminated during\nrecording (Omi AI: built-in)",
             "Satisfies notification requirements\nin some jurisdictions"],
            ["Derived features only\nin long-term storage", "Raw audio deleted after DAM\ninference; "
             "only scores and acoustic\nfeatures retained",
             "Minimizes stored biometric data;\nscores are not personally identifiable"],
            ["User-responsible\nconsent", "App prompts user to inform\nconversation participants;\n"
             "user bears legal responsibility",
             "Shifts consent burden to user\n(similar to phone recording apps)"],
        ],
        col_widths=[3.0, 4.5, 4.5]
    )

    add_callout_box(doc, "Privacy by Design",
        "The architecture ensures that at no point does the system store or transmit another person's "
        "voice. The speaker gate acts as a privacy firewall: audio that does not match the enrolled "
        "user is discarded within seconds on the phone, before any sync to the desktop occurs. "
        "This means even if the wearable captures a multi-person conversation, only the user's "
        "own voice segments are retained and analyzed."
    )

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 7. TECHNICAL RISKS AND MITIGATIONS
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_section_header(doc, "7", "Technical Risks and Mitigations")

    add_table(doc,
        ["Risk", "Severity", "Likelihood", "Mitigation"],
        [
            ["DAM accuracy\ndegradation on\nwearable audio",
             "HIGH", "MEDIUM",
             "Phase 1 calibration study: record 20+ hours with simultaneous "
             "wearable and desktop capture. Compare DAM scores. If delta > 15%, "
             "investigate fine-tuning DAM on augmented data or applying a "
             "calibration offset per source."],
            ["Speaker verification\nfalse rejections\n(cross-device)",
             "MEDIUM", "HIGH",
             "Cross-device enrollment (Section 5.2). Lower wearable threshold "
             "to 0.22 (above absolute floor). Dual-centroid approach with "
             "source-aware selection. Monitor rejection rates in production."],
            ["BLE audio\ndropouts / gaps",
             "MEDIUM", "MEDIUM",
             "Implement gap detection in wearable_ingest.py. Mark segments with "
             "gaps > 500ms. Use Omi's built-in BLE reconnection. Accept partial "
             "segments (DAM handles 30-90s variable input)."],
            ["Phone battery\ndrain unacceptable",
             "LOW", "MEDIUM",
             "Architecture A: ~15-20% per day (acceptable for most users). "
             "Architecture B reduces to ~5-8%. Offer 'scheduled capture' mode "
             "(e.g., only during work hours outside office)."],
            ["Sync reliability\n(WiFi drops, partial\ntransfers)",
             "MEDIUM", "LOW",
             "Idempotent sync protocol: each segment has unique hash + timestamp. "
             "wearable_ingest.py skips already-processed hashes. Retry on next "
             "WiFi connection. No data loss possible."],
            ["Temporal ordering\nof interleaved\nreadings",
             "LOW", "HIGH",
             "All wearable segments carry original recording timestamps. "
             "wearable_ingest.py inserts readings with correct timestamps. "
             "EMA recalculation uses chronological order, not insertion order."],
            ["User social\ndiscomfort wearing\nrecording device",
             "MEDIUM", "HIGH",
             "Omi AI is small and discreet (pendant form factor). LED indicator "
             "is subtle. User education: device only stores YOUR voice, not "
             "others'. Offer 'pause' gesture (double-tap to disable)."],
            ["Wearable market\ninstability\n(devices discontinued)",
             "LOW", "MEDIUM",
             "Architecture designed to be device-agnostic. wearable_ingest.py "
             "accepts standard WAV files from any source. Omi is open-source "
             "(firmware can be maintained independently). Plaud as fallback."],
        ],
        col_widths=[3.0, 1.5, 1.8, 7.0]
    )

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 8. IMPLEMENTATION ROADMAP
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_section_header(doc, "8", "Implementation Roadmap")

    add_subsection(doc, "8.1", "Phase 1: Proof of Concept (Months 1–3)")

    add_body_rich(doc, [
        ("Goal: ", True, False),
        ("Validate that wearable-captured audio produces DAM scores within ±15% of "
         "desktop-captured scores. Minimal code, maximum learning.", False, False),
    ])

    add_body_rich(doc, [
        ("Estimated effort: ", True, False),
        ("60–80 engineering hours", False, False),
    ])

    add_table(doc,
        ["Task", "Effort", "Deliverable"],
        [
            ["Acquire Omi AI Dev Kit 2", "1 hr", "Hardware in hand"],
            ["Set up Omi open-source app\n(iOS/Android) for raw audio export", "8–12 hr",
             "App capturing 16kHz PCM\nvia BLE from Omi"],
            ["Design comparison study protocol:\nsimultaneous desktop + wearable\nrecording sessions",
             "4–6 hr", "Study protocol document"],
            ["Conduct comparison recordings:\n20+ hours across varied environments\n(office, home, outdoors, social)",
             "10–15 hr\n(elapsed: 2 weeks)", "Raw audio dataset with\nmatched pairs"],
            ["Build wearable_ingest.py:\nwatch directory, validate WAV,\nfeed to existing pipeline",
             "12–16 hr", "Working ingest module\nwith source tagging"],
            ["Run DAM on both sources,\ncompare scores statistically",
             "8–10 hr", "Comparison report with\nscatter plots, Bland-Altman"],
            ["Database migration:\nadd source column to readings",
             "4–6 hr", "Updated schema, migration\nscript, backward compat"],
            ["Decision gate: proceed to\nPhase 2 or investigate\ncalibration offsets",
             "4 hr", "Go/no-go decision with\ndata-backed rationale"],
        ],
        col_widths=[4.5, 2.0, 4.5]
    )

    add_subsection(doc, "8.2", "Phase 2: Companion App MVP (Months 4–6)")

    add_body_rich(doc, [
        ("Goal: ", True, False),
        ("Automated wearable-to-desktop pipeline with speaker verification on phone "
         "and WiFi sync. No manual file transfer required.", False, False),
    ])

    add_body_rich(doc, [
        ("Estimated effort: ", True, False),
        ("120–160 engineering hours", False, False),
    ])

    add_table(doc,
        ["Task", "Effort", "Deliverable"],
        [
            ["Design companion app architecture\n(iOS: Swift, Android: Kotlin)",
             "8–12 hr", "Architecture doc, wireframes"],
            ["Implement BLE audio receiver\n(Omi protocol, 16kHz PCM stream)",
             "16–24 hr", "Working BLE connection\nwith audio capture"],
            ["Port Silero VAD to mobile\n(ONNX runtime, CoreML/TFLite)",
             "12–16 hr", "On-phone VAD with\n<5ms latency per chunk"],
            ["Port ECAPA-TDNN to mobile\n(ONNX, ~25 MB model)",
             "16–20 hr", "Speaker verification\non 2s segments"],
            ["Cross-device enrollment flow:\nwearable sample capture in Lucid\ndesktop, sync to phone",
             "12–16 hr", "Enrollment sharing\nbetween desktop and phone"],
            ["WiFi sync service:\nauto-discover desktop Lucid,\ntransfer verified WAV segments",
             "16–20 hr", "Automated sync when\non same network"],
            ["Desktop: wearable_ingest.py\nenhancements (auto-import,\ndeduplication, status API)",
             "12–16 hr", "Robust ingestion with\nstatus reporting"],
            ["Dashboard: wearable readings UI\n(distinct visual treatment,\ncoverage timeline)",
             "12–16 hr", "Updated dashboard showing\ndesktop + wearable readings"],
            ["Integration testing across\nfull pipeline", "8–12 hr",
             "End-to-end test suite"],
        ],
        col_widths=[4.5, 2.0, 4.0]
    )

    add_subsection(doc, "8.3", "Phase 3: Production Wearable Integration (Months 7–12)")

    add_body_rich(doc, [
        ("Goal: ", True, False),
        ("Production-quality wearable pipeline with on-device VAD, batch processing UI, "
         "privacy review, and support for multiple wearable devices.", False, False),
    ])

    add_body_rich(doc, [
        ("Estimated effort: ", True, False),
        ("200–300 engineering hours", False, False),
    ])

    add_table(doc,
        ["Task", "Effort", "Deliverable"],
        [
            ["On-device VAD (Architecture B):\nPicovoice Cobra on Omi firmware\nor custom nRF52840 firmware",
             "40–60 hr", "Custom firmware with\nedge VAD"],
            ["Batch processing UI:\nqueue wearable segments,\nprogress indicator, results view",
             "24–32 hr", "Batch processing screen\nin Lucid desktop"],
            ["Coverage analytics:\ndaily coverage percentage,\ngap detection, source breakdown",
             "16–24 hr", "Coverage dashboard with\ndesktop vs wearable timeline"],
            ["Plaud NotePin integration\n(Architecture C alternative):\nUSB import, format conversion",
             "20–30 hr", "Plaud import workflow"],
            ["Privacy review:\nlegal counsel, DPIA template,\nconsent flow design",
             "24–40 hr\n(includes legal)", "Privacy documentation,\nin-app consent flows"],
            ["Multi-device management:\nsupport multiple wearables,\ndevice settings, pairing UI",
             "20–30 hr", "Device management in\nLucid settings"],
            ["DAM calibration offsets\n(if Phase 1 showed >15% delta):\nper-source score adjustment",
             "16–24 hr", "Calibrated scoring with\nsource-aware offsets"],
            ["Performance optimization:\nbatch DAM inference,\nparallel processing",
             "16–24 hr", "Efficient batch processing\nfor full-day imports"],
            ["Documentation and user guides", "12–16 hr", "User-facing documentation"],
        ],
        col_widths=[4.5, 2.0, 4.0]
    )

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 9. RECOMMENDATION
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_section_header(doc, "9", "Recommendation")

    add_body(doc, (
        "Based on the analysis in this report, the recommended path forward is a phased approach "
        "starting with the Omi AI Dev Kit 2 and Architecture A (Store-and-Sync)."
    ))

    add_callout_box(doc, "Primary Recommendation: Omi AI + Architecture A",
        "Start with a $70 Omi AI Dev Kit 2 and a minimal companion app. The open-source firmware "
        "and app provide a reference implementation for BLE audio streaming at 16kHz. The companion "
        "app runs Silero VAD and ECAPA-TDNN speaker verification, syncing only verified user speech "
        "to the desktop via WiFi. The existing Lucid DAM pipeline runs unchanged on synced audio."
    )

    add_body(doc, "Success criteria for Phase 1 (PoC):")

    add_bullet(doc, "Wearable DAM depression scores correlate with desktop scores "
               "(Pearson r > 0.80) across 20+ hours of matched recordings.",
               bold_prefix="Score correlation: ")
    add_bullet(doc, "Absolute difference between wearable and desktop DAM scores "
               "averages less than 15% of the score range.",
               bold_prefix="Score accuracy: ")
    add_bullet(doc, "Speaker verification on wearable audio achieves > 90% true positive "
               "rate with < 5% false positive rate after cross-device enrollment.",
               bold_prefix="Speaker gate reliability: ")
    add_bullet(doc, "BLE audio streaming from Omi to phone maintains < 2% gap rate "
               "across 50+ hours of recording.",
               bold_prefix="BLE reliability: ")

    add_body(doc, "If Phase 1 criteria are met, proceed to Phase 2 (companion app MVP). "
             "If DAM scores show > 15% divergence, investigate:")

    add_bullet(doc, "Fine-tuning DAM on augmented data (wearable-simulated audio via "
               "frequency response transfer functions).")
    add_bullet(doc, "Per-source calibration offsets (additive or multiplicative correction "
               "applied to wearable DAM outputs before score mapping).")
    add_bullet(doc, "Audio preprocessing pipeline (noise reduction, frequency equalization) "
               "applied to wearable audio before DAM inference.")

    add_body(doc, "")
    add_body_rich(doc, [
        ("Alternative for no-phone workflow: ", True, False),
        ("Plaud NotePin ($149) with Architecture C (Direct Mac Sync). This provides the simplest "
         "user experience (wear all day, dock at night) but sacrifices real-time readings. "
         "Best suited for users who prefer a phone-free workflow and are comfortable with "
         "end-of-day batch processing.", False, False),
    ])

    add_body(doc, "")
    add_body_rich(doc, [
        ("Total estimated investment: ", True, False),
        ("380–540 engineering hours across 12 months, plus $70–149 in hardware. "
         "The phased approach ensures each investment is validated before proceeding, "
         "with clear decision gates between phases.", False, False),
    ])

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # APPENDIX A
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_section_header(doc, "A", "Appendix A: Data Format Specifications")

    add_subsection(doc, "A.1", "WAV File Format for Wearable Audio")

    add_code_block(doc, "\n".join([
        "Format:          WAV (RIFF)",
        "Sample Rate:     16,000 Hz",
        "Channels:        1 (mono)",
        "Bit Depth:       16-bit signed integer (PCM)",
        "Byte Rate:       32,000 bytes/sec",
        "File Size:       ~1.92 MB per minute",
        "",
        "Alternative (if BLE codec used):",
        "  Opus at 32 kbps → ~240 KB per minute",
        "  Must be decoded to PCM before DAM processing",
    ]))

    add_subsection(doc, "A.2", "Metadata JSON per Segment")

    add_code_block(doc, "\n".join([
        '{',
        '  "segment_id": "uuid-v4",',
        '  "source": "wearable",',
        '  "device_type": "omi_ai_v2",',
        '  "device_id": "AA:BB:CC:DD:EE:FF",',
        '  "recording_start": "2026-03-15T14:23:01.000Z",',
        '  "recording_end": "2026-03-15T14:24:07.500Z",',
        '  "duration_sec": 66.5,',
        '  "speech_duration_sec": 42.3,',
        '  "sample_rate": 16000,',
        '  "channels": 1,',
        '  "bit_depth": 16,',
        '  "vad_model": "silero_v6.2.0",',
        '  "speaker_verified": true,',
        '  "speaker_similarity": 0.34,',
        '  "ble_codec": "none",',
        '  "file_hash_sha256": "a1b2c3d4...",',
        '  "sync_timestamp": "2026-03-15T18:45:00.000Z"',
        '}',
    ]))

    add_subsection(doc, "A.3", "Sync Directory Structure")

    add_code_block(doc, "\n".join([
        "~/Library/Application Support/lucid/wearable_sync/",
        "├── incoming/              # Phone drops files here via WiFi",
        "│   ├── seg_2026-03-15_142301_omi.wav",
        "│   ├── seg_2026-03-15_142301_omi.json",
        "│   ├── seg_2026-03-15_150445_omi.wav",
        "│   └── seg_2026-03-15_150445_omi.json",
        "├── processing/            # Moved here during DAM inference",
        "│   └── (empty when idle)",
        "├── completed/             # Successfully processed",
        "│   ├── seg_2026-03-15_091200_omi.wav",
        "│   └── seg_2026-03-15_091200_omi.json",
        "├── failed/                # Processing errors (for debugging)",
        "│   └── (error logs)",
        "└── sync_state.json        # Tracks processed hashes for dedup",
    ]))

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # APPENDIX B
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_section_header(doc, "B", "Appendix B: BLE Audio Streaming Protocol")

    add_body(doc, (
        "Bluetooth Low Energy (BLE) audio streaming from the Omi AI Dev Kit 2 uses a custom "
        "GATT service for raw audio data. The Omi open-source firmware (nRF52840-based) exposes:"
    ))

    add_table(doc,
        ["GATT Characteristic", "UUID", "Description"],
        [
            ["Audio Data", "Custom UUID\n(Omi-specific)", "Raw PCM audio chunks, MTU-dependent\nsize (typically 244 bytes per packet)"],
            ["Audio Config", "Custom UUID", "Sample rate, codec, channels configuration"],
            ["Audio Control", "Custom UUID", "Start/stop/pause streaming commands"],
            ["Battery Level", "0x2A19", "Standard BLE battery service"],
            ["Device Info", "0x180A", "Model, firmware version, serial number"],
        ],
        col_widths=[3.0, 2.5, 6.0]
    )

    add_body(doc, "BLE audio streaming bandwidth analysis:")

    add_table(doc,
        ["Parameter", "Value", "Notes"],
        [
            ["BLE version", "5.0 (Omi Dev Kit 2)", "Supports 2 Mbps PHY"],
            ["MTU size", "244 bytes (typical)", "Negotiated at connection"],
            ["Connection interval", "7.5–15 ms", "Affects latency and power"],
            ["Raw PCM throughput needed", "32 KB/s (16kHz × 16-bit)", "Comfortable within BLE 5.0 limits"],
            ["Practical throughput", "~100–200 KB/s", "Well above requirement"],
            ["Latency", "15–30 ms (typical)", "Acceptable for non-real-time analysis"],
            ["Range", "~10 meters (indoor)", "Phone must be carried with wearable"],
        ],
        col_widths=[3.5, 3.5, 5.0]
    )

    add_body(doc, (
        "The BLE bandwidth requirement for 16kHz 16-bit mono audio (32 KB/s) is well within "
        "BLE 5.0 capabilities. The primary concern is connection stability, not throughput. "
        "The Omi firmware implements automatic reconnection with buffering during brief "
        "disconnections."
    ))

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # APPENDIX C
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_section_header(doc, "C", "Appendix C: Competitive Landscape")

    add_body(doc, (
        "Several companies operate in the voice-based wellness and mental health monitoring space. "
        "Understanding their approaches provides context for Attune Steel's wearable strategy."
    ))

    add_table(doc,
        ["Company", "Approach", "Model", "Wearable?", "Relevance to Attune"],
        [
            ["Kintsugi Health\n(SHUT DOWN\nFeb 2026)", "Voice biomarker detection\nfor depression/anxiety.\n"
             "Provided the DAM model\nused in Attune. Open-sourced\n$30M of R&D on shutdown.",
             "Was B2B API + clinical.\nModels now on HuggingFace\n(KintsugiHealth org).\n"
             "FDA pursuit was too\ncostly for VC model.",
             "No (was API-based)",
             "CRITICAL: DAM model origin.\nShutdown means no upstream\nsupport. Open-sourced models\n"
             "reduce licensing risk.\nWearable extension is now\nentirely self-maintained."],
            ["Hume AI", "Multimodal emotion AI.\nVoice prosody analysis\nfor sentiment, emotion.\n$50M Series B funded.",
             "Cloud API (EVI 3,\nlaunched May 2025).\nSDKs for React, Python,\nSwift, TypeScript.",
             "No (cloud API,\nrequires internet)",
             "Potential future complement\nfor emotion features. Not\nsuitable for offline/privacy\narchitecture."],
            ["Beyond Verbal\n(now Vocalis Health)", "Voice-based health\nbiomarkers. Chronic\ndisease detection via\nvocal features.",
             "Clinical validation\nwith Mayo Clinic.\nAlzheimer's + Parkinson's\ndetection (2025).",
             "No (clinical\nsettings only)",
             "Validates the voice-biomarker\napproach. Their clinical\nresearch supports the\nscientific basis."],
            ["Sonde Health", "Voice-based health\nmonitoring. Respiratory,\ncognitive, mental health\nbiomarkers.",
             "B2B SDK. Partners\nintegrate into existing\napps and devices.",
             "Exploring (SDK\ncould integrate\nwith wearables)",
             "Competitive intelligence.\nTheir SDK model could\nenable competitor wearable\nintegrations."],
            ["Ellipsis Health", "Clinical voice analysis\nfor behavioral health.\nPHQ-9 and GAD-7 scoring\nfrom voice.",
             "Telehealth integration.\nClinical-grade analysis\nvia phone calls.",
             "No (phone-based,\nclinical setting)",
             "Similar clinical scoring\n(PHQ-9/GAD-7). Validates\nAttune's scoring approach.\nDoes not target consumer\nwearable market."],
        ],
        col_widths=[2.2, 3.0, 2.8, 2.2, 3.5]
    )

    add_callout_box(doc, "Attune's Differentiator",
        "None of these competitors offer a consumer wearable solution with local-only voice "
        "biomarker analysis. Kintsugi's shutdown (February 2026) — despite validated technology "
        "and $30M in R&D — underscores the difficulty of the FDA regulatory path and validates "
        "Attune's consumer-wellness positioning over clinical diagnostics. Attune Steel's "
        "desktop-first architecture with wearable extension occupies a unique position: continuous, "
        "private, wellness-grade voice monitoring that never sends audio to the cloud. A successful "
        "wearable extension would further differentiate Attune by providing all-day coverage — "
        "something no competitor currently offers in a privacy-preserving manner."
    )

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # FINAL PAGE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()

    add_styled_paragraph(doc, "━" * 50, font_size=10, color=COLOR_SECONDARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16))
    add_styled_paragraph(doc, "End of Report", font_size=14, bold=True,
                         color=COLOR_PRIMARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    add_styled_paragraph(doc, "Attune Steel — Wearable Extension Technical Research Report",
                         font_size=10, color=COLOR_SECONDARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_styled_paragraph(doc, "March 2026 · Version 1.0 · CONFIDENTIAL",
                         font_size=10, color=COLOR_SECONDARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Save
    doc.save(OUTPUT_FILE)
    print(f"Report saved to: {OUTPUT_FILE}")

    # Count approximate pages (rough estimate: ~40 paragraphs per page)
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    est_pages = (para_count + table_count * 5) // 35
    print(f"Paragraphs: {para_count}, Tables: {table_count}, Estimated pages: {est_pages}")


if __name__ == "__main__":
    build_report()
