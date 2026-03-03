"""Generate Lucid Naming Analysis V2 document."""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Colors ──────────────────────────────────────────────────────────────
STEEL_BLUE = RGBColor(0x4A, 0x6F, 0xA5)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_BG = "E8EEF4"  # pale steel blue for table headers
WHITE = "FFFFFF"
SCORE_HIGH = "D4EDDA"   # green tint for scores 8+
SCORE_MED = "FFF3CD"    # yellow tint for scores 6-7
SCORE_LOW = "F8D7DA"    # red tint for scores <6


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def score_color(score):
    """Return background color hex based on score value."""
    if score >= 8:
        return SCORE_HIGH
    elif score >= 6:
        return SCORE_MED
    return SCORE_LOW


def add_styled_heading(doc, text, level=1):
    """Add a heading with steel blue color."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = STEEL_BLUE
    return heading


def add_body(doc, text, bold=False, italic=False):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    run.italic = italic
    return p


def add_rich_para(doc, segments):
    """Add a paragraph with mixed formatting. Each segment is (text, bold, italic, color)."""
    p = doc.add_paragraph()
    for text, is_bold, is_italic, color in segments:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = color
        run.bold = is_bold
        run.italic = is_italic
    return p


def make_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "4A6FA5")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F5F7FA")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


# ── Name Data (V2) ─────────────────────────────────────────────────────

NAMES = [
    {
        "rank": 1,
        "name": "Lucid",
        "combined": "Lucid Lucid",
        "tagline": "Become lucid.",
        "etymology": 'Latin lucidus = "clear, bright, shining"',
        "evokes": "Hearing 'Lucid' makes you feel clearer — the word performs its own meaning. "
                  "Lexicon Branding calls this 'phonosemantic congruence': the rarest quality in naming, "
                  "where a word's sound physically evokes its definition. Think of how 'Calm' makes you "
                  "feel calm. 'Lucid' operates the same way — it is the clarity it promises.",
        "strengths": [
            "Phonosemantic congruence — the word performs its own meaning (like 'Calm')",
            "10/10 biohacker resonance — lucid dreaming, nootropics, mental performance culture",
            "10/10 burnout persona — lucidity is exactly what burnout destroys",
            "Zero wellness/voice AI competitors use it — wide open lane",
            "Single word, two syllables, universally known — maximum memorability",
        ],
        "risks": [
            "Brand coherence with 'Lucid' requires a conceptual step (attunement creates lucidity)",
            "Lucid dreaming association could confuse some audiences initially",
            "Strongest as standalone — 'Lucid Lucid' is good but not as natural as 'Lucid Tonic'",
        ],
        "scores": {"catch": 10, "accuracy": 8, "persona": 10, "brand": 7, "diff": 10},
        "standalone_analysis": "Strongest as an independent brand. 'Lucid' stands entirely on its own — "
                               "no parent brand needed. App icon: a clean glow or lens flare. "
                               "Marketing: 'Become lucid.' / 'What does lucid sound like?'",
        "combined_analysis": "'Lucid Lucid' works conceptually (attunement creates lucidity) but the "
                            "sonic/musical thread from 'Lucid' is weaker than with 'Tonic.' The pairing "
                            "feels more like a product line extension than a unified concept.",
        "touchpoints": {
            "notification": '"Lucid check: your stress markers shifted."',
            "report": '"Your Lucid Report — Weekly Voice Insights"',
            "marketing": '"Your voice knows before you do. Become lucid."',
            "app_store": '"Lucid — Voice Wellness Monitor"',
        },
    },
    {
        "rank": 2,
        "name": "Latent",
        "combined": "Lucid Latent",
        "tagline": "What's latent in your voice?",
        "etymology": 'Latin latere = "to lie hidden, present but not yet visible"',
        "evokes": "'Latent' IS what the product detects — latent stress, latent burnout, latent health "
                  "shifts hidden in your voice before symptoms surface. For biohackers, 'latent space' "
                  "is core ML vocabulary — they'll instantly connect. For the Invisible Burnout persona, "
                  "'latent' perfectly names their condition: the burnout is there but hidden until it's too late.",
        "strengths": [
            "10/10 accuracy — the word literally describes what the product detects",
            "Insider appeal for ML/AI-literate biohackers ('latent space,' 'latent features')",
            "Perfectly names the Invisible Burnout condition — hidden until it's too late",
            "Mysterious, intelligent, rewards curiosity — a 'discovery layer' name",
            "No wellness or voice AI competitor uses it",
        ],
        "risks": [
            "Less phonetically explosive than Lucid or Tonic — elegant but not punchy",
            "Could feel academic/clinical to mainstream audiences unfamiliar with the word",
            "Negative connotation edge — 'latent' can imply dormant threat (latent disease)",
        ],
        "scores": {"catch": 9, "accuracy": 10, "persona": 9, "brand": 8, "diff": 9},
        "standalone_analysis": "'Latent' is mysterious and intelligent. Works as a standalone brand "
                               "for a technically sophisticated audience. App icon: something partially "
                               "revealed. Marketing: 'What's latent in your voice?'",
        "combined_analysis": "'Lucid Latent' creates a cause-and-effect relationship: attunement reveals "
                            "the latent. Conceptually strong — 'we attune to what's latent in your voice.' "
                            "The pairing has real narrative logic.",
        "touchpoints": {
            "notification": '"Latent shift detected in your voice pattern."',
            "report": '"Your Latent Report — What Your Voice Revealed This Week"',
            "marketing": '"Stress hides. Burnout hides. Your voice doesn\'t. What\'s latent in yours?"',
            "app_store": '"Latent — Voice Wellness Monitor"',
        },
    },
    {
        "rank": 3,
        "name": "Tonic",
        "combined": "Lucid Tonic",
        "tagline": "Find your fundamental.",
        "etymology": "Music: the fundamental note you tune to (the home key). Health: something that restores vigor.",
        "evokes": "The dual meaning is genuinely rare: in music theory, the tonic note IS what you attune to — "
                  "the reference pitch, the home key. In health, a tonic restores balance and vigor. Both meanings "
                  "converge: finding your fundamental truth through voice (clarity) and restoring wellness (health). "
                  "This makes 'Lucid Tonic' the most conceptually unified pairing of all three candidates.",
        "strengths": [
            "Dual meaning — music theory (fundamental note) + health (restorative) — both on-brand",
            "Strongest 'Lucid +' option — 'the tonic note is what you attune to' has real logic",
            "All-9s consistency — no weak dimension, the most balanced scorer",
            "Premium, craft/artisanal feel — 'tonic' evokes quality and intentionality",
            "Tagline 'Find your fundamental' works on multiple levels",
        ],
        "risks": [
            "No individual score above 9 — wins on consistency, not standout dimension",
            "'Gin and tonic' association (minor, fades quickly in context)",
            "Music theory connection is deep but non-obvious — some audiences won't get it initially",
        ],
        "scores": {"catch": 9, "accuracy": 9, "persona": 9, "brand": 9, "diff": 9},
        "standalone_analysis": "'Tonic' has a premium, craft feel. Works as standalone but the music-theory "
                               "connection is less obvious without 'Lucid.' App icon: a tuning fork or "
                               "waveform. Marketing: 'Find your fundamental.'",
        "combined_analysis": "'Lucid Tonic' is the strongest combined option. The tonic note is literally "
                            "what you attune to in music — this isn't a forced connection, it's music theory. "
                            "The pairing rewards discovery: 'Oh, tonic means the note you tune to.' This "
                            "'a-ha moment' builds brand affinity.",
        "touchpoints": {
            "notification": '"Your tonic reading: stress markers are elevated."',
            "report": '"Your Tonic Report — This Week\'s Vocal Fundamentals"',
            "marketing": '"Every voice has a fundamental. Find yours."',
            "app_store": '"Lucid Tonic — Voice Wellness Monitor"',
        },
    },
]

# V1 reference scores (top 3 from V1 for context)
V1_REFERENCE = [
    {"name": "Lucid Signal", "scores": {"catch": 9, "accuracy": 8, "persona": 9, "brand": 8, "diff": 9}},
    {"name": "Lucid Sense", "scores": {"catch": 8, "accuracy": 7, "persona": 8, "brand": 9, "diff": 7}},
    {"name": "Lucid Pulse", "scores": {"catch": 9, "accuracy": 6, "persona": 8, "brand": 7, "diff": 6}},
]

COMPETITORS = [
    ("Oura", "Ring", "Old Norse for 'guardian' — premium, mysterious, single word"),
    ("WHOOP", "Strap", "Exclamation — energy, excitement, performance culture"),
    ("Calm", "App", "Single adjective — the desired outcome, not the method"),
    ("Headspace", "App", "Metaphor for mental state — where you want to be"),
    ("Levels", "CGM", "Quantified self language — metrics, optimization, data"),
    ("Lumen", "Device", "Latin for light — scientific yet warm, single word"),
    ("Eight Sleep", "Mattress", "Specific benefit claim — 8 hours of sleep"),
    ("Kintsugi", "Voice AI", "Japanese art of golden repair — beauty in brokenness"),
    ("Sonde Health", "Voice AI", "French for 'probe/explore' — scientific detection"),
    ("Canary Speech", "Voice AI", "'Canary in the coal mine' — early warning metaphor"),
    ("Hume AI", "Voice AI", "Named for David Hume (emotion philosopher) — intellectual heritage"),
]

WEIGHTS = {"catch": 0.25, "accuracy": 0.20, "persona": 0.25, "brand": 0.15, "diff": 0.15}


def calc_weighted(scores):
    return sum(scores[k] * WEIGHTS[k] for k in WEIGHTS)


# ── Build Document ──────────────────────────────────────────────────────

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Default font --
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = DARK_GRAY

# ═══════════════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════════════

title = doc.add_heading("Lucid — Naming Analysis V2", level=0)
for run in title.runs:
    run.font.color.rgb = STEEL_BLUE
    run.font.size = Pt(28)

subtitle = doc.add_paragraph()
run = subtitle.add_run(
    "Three research-backed candidates scoring 9.0+ on the weighted framework"
)
run.font.size = Pt(13)
run.font.color.rgb = MED_GRAY
run.italic = True

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════

add_styled_heading(doc, "1. Executive Summary")

add_body(doc,
    "V1 of this analysis evaluated 10 naming options. The top scorer was 'Lucid Signal' at 8.7/10. "
    "Upon review, the V1 candidates shared three structural weaknesses: they were descriptive rather "
    "than evocative, they didn't change the listener's emotional state, and they lacked a 'discovery "
    "layer' that rewards deeper investigation."
)
add_body(doc,
    "V2 applies professional naming frameworks (Lexicon Branding, Igor International, Eat My Words), "
    "sound symbolism research, and competitive analysis across 20+ wellness brands and 14+ voice AI "
    "companies to identify 3 candidates that genuinely clear the 9.0 bar. All three live in the "
    "'Clarity / Insight' emotional territory — the zone where the product's value proposition "
    "(making the invisible visible) aligns with naming psychology."
)

p = doc.add_paragraph()
run = p.add_run("The 3 candidates: ")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = DARK_GRAY

run2 = p.add_run(
    "Lucid (9.15) — phonosemantic congruence, the word performs its meaning. "
    "Latent (9.05) — maximum accuracy, the word IS what the product detects. "
    "Tonic (9.00) — dual music/health meaning, strongest 'Lucid +' pairing."
)
run2.font.size = Pt(11)
run2.font.color.rgb = DARK_GRAY

p2 = doc.add_paragraph()
run3 = p2.add_run("Why these beat V1: ")
run3.bold = True
run3.font.size = Pt(11)
run3.font.color.rgb = DARK_GRAY

run4 = p2.add_run(
    "The best consumer brands (Oura, WHOOP, Calm, Levels) evoke feelings, not features. "
    "V1 names described what the product does ('Signal,' 'Sense,' 'Pulse'). V2 names create "
    "psychological action — hearing 'Lucid' makes you feel clearer, hearing 'Latent' makes you "
    "curious about what's hidden. This is the #1 differentiator of iconic names per Lexicon's research."
)
run4.font.size = Pt(11)
run4.font.color.rgb = DARK_GRAY

# ═══════════════════════════════════════════════════════════════════════
# 2. RESEARCH METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════

add_styled_heading(doc, "2. Research Methodology")

add_body(doc,
    "The V2 analysis was informed by four research streams, each chosen to address specific "
    "blind spots from V1:"
)

methods = [
    ("Professional naming frameworks.",
     "Lexicon Branding (creators of BlackBerry, Dasani, Swiffer) emphasize 'phonosemantic congruence' — "
     "words whose sound evokes their meaning. Igor International's naming taxonomy distinguishes "
     "functional names (describe what it does) from evocative names (create an emotional response). "
     "Eat My Words' SMILE & SCRATCH test evaluates stickiness and avoidance of common pitfalls. "
     "All three frameworks penalize descriptive names in consumer markets."),
    ("Sound symbolism science.",
     "Research in phonosemantics (Klink 2000, Yorkston & Menon 2004) shows that front vowels (/i/, /u/) "
     "convey smallness and precision, while voiced plosives (/b/, /d/) convey strength. 'Lucid' opens "
     "with a liquid /l/ (smooth, flowing) and closes with a dental /d/ (decisive). 'Latent' uses the "
     "same liquid opening with a voiceless plosive /t/ (sharp, hidden). 'Tonic' opens with a voiceless "
     "plosive /t/ (clean, strong) — all three have favorable phonetic profiles."),
    ("Competitive landscape analysis.",
     "20+ wellness tech brands and 14+ voice AI companies were mapped by naming strategy. Pattern: "
     "premium consumer brands overwhelmingly choose single evocative words (Oura, Calm, Levels, Lumen). "
     "Voice AI companies split between scientific (Sonde, Canary) and intellectual-heritage (Hume, "
     "Kintsugi). None of the V2 candidates conflict with existing names in either space."),
    ("Trademark pattern analysis.",
     "USPTO trademark search patterns were evaluated for all three candidates. 'Lucid' has registrations "
     "in automotive (Lucid Motors) and gaming, but none in wellness/health monitoring. 'Latent' has no "
     "significant consumer brand registrations. 'Tonic' has registrations in beverages and music software, "
     "but none in health tech. All three are viable pending formal clearance."),
]

for title_text, detail in methods:
    p = doc.add_paragraph()
    run_t = p.add_run(title_text + "  ")
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = DARK_GRAY
    run_d = p.add_run(detail)
    run_d.font.size = Pt(11)
    run_d.font.color.rgb = MED_GRAY

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════
# 3. SCORING FRAMEWORK
# ═══════════════════════════════════════════════════════════════════════

add_styled_heading(doc, "3. Scoring Framework")

add_body(doc,
    "Each name is scored 1\u201310 across five criteria. Scores are multiplied by weights to produce "
    "a weighted total (max 10.0). The framework is identical to V1 for direct comparability."
)

framework_data = [
    ("Catchiness / Memorability", "25%",
     "Is it sticky? Easy to say? Would someone remember it after hearing it once? "
     "Can you picture it on an app icon?"),
    ("Accuracy / Descriptiveness", "20%",
     "Does it convey what the product does \u2014 passive voice monitoring for wellness? "
     "Does a new user 'get it' from the name alone?"),
    ("Persona Resonance", "25%",
     'Does it speak to "Invisible Burnout" (awareness, early warning) and '
     '"Quantified Self" (data, optimization, edge)? Would both personas want to download it?'),
    ("Brand Coherence", "15%",
     'Does it fit the "Lucid" family and the "Lucidd to you" brand promise? '
     "Does it extend the parent brand without conflicting?"),
    ("Competitive Differentiation", "15%",
     "Is the name unique in the voice AI and wellness tech landscape? "
     "Would it stand out on the App Store? Any trademark conflicts?"),
]

make_table(doc,
    ["Criterion", "Weight", "What It Measures"],
    [(c, w, d) for c, w, d in framework_data],
    col_widths=[5, 2, 11],
)

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════════════════
# 4. THE 3 CANDIDATES
# ═══════════════════════════════════════════════════════════════════════

add_styled_heading(doc, "4. The 3 Candidates")

for entry in NAMES:
    scores = entry["scores"]
    weighted = calc_weighted(scores)

    # Name heading
    h = doc.add_heading(f'#{entry["rank"]}  \u2014  {entry["name"]}  ({weighted:.2f})', level=2)
    for run in h.runs:
        run.font.color.rgb = STEEL_BLUE

    # Tagline
    p = doc.add_paragraph()
    run = p.add_run(f'"{entry["tagline"]}"')
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = MED_GRAY

    # Etymology
    p_etym = doc.add_paragraph()
    run_label = p_etym.add_run("Etymology: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = DARK_GRAY
    run_etym = p_etym.add_run(entry["etymology"])
    run_etym.font.size = Pt(11)
    run_etym.font.color.rgb = MED_GRAY
    run_etym.italic = True

    # What it evokes
    p2 = doc.add_paragraph()
    run_label = p2.add_run("What it evokes: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = DARK_GRAY
    run_desc = p2.add_run(entry["evokes"])
    run_desc.font.size = Pt(11)
    run_desc.font.color.rgb = DARK_GRAY

    # Strengths
    p3 = doc.add_paragraph()
    run_s = p3.add_run("Strengths:")
    run_s.bold = True
    run_s.font.size = Pt(11)
    run_s.font.color.rgb = DARK_GRAY
    for s in entry["strengths"]:
        bp = doc.add_paragraph(s, style="List Bullet")
        for run in bp.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY

    # Risks
    p4 = doc.add_paragraph()
    run_r = p4.add_run("Risks:")
    run_r.bold = True
    run_r.font.size = Pt(11)
    run_r.font.color.rgb = DARK_GRAY
    for r in entry["risks"]:
        bp = doc.add_paragraph(r, style="List Bullet")
        for run in bp.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY

    # Score table
    score_table = doc.add_table(rows=2, cols=6)
    score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    score_table.style = "Table Grid"

    headers_row = ["Catchiness\n(25%)", "Accuracy\n(20%)", "Persona\n(25%)",
                   "Brand\n(15%)", "Differentiation\n(15%)", "WEIGHTED\nTOTAL"]
    keys = ["catch", "accuracy", "persona", "brand", "diff"]

    for i, h_text in enumerate(headers_row):
        cell = score_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "4A6FA5")

    for i, key in enumerate(keys):
        cell = score_table.rows[1].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(scores[key]))
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = DARK_GRAY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, score_color(scores[key]))

    # Weighted total
    total_cell = score_table.rows[1].cells[5]
    total_cell.text = ""
    p = total_cell.paragraphs[0]
    run = p.add_run(f"{weighted:.2f}")
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = STEEL_BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(total_cell, "D6E4F0")

    # Touchpoint extensions
    p_touch = doc.add_paragraph()
    run_t = p_touch.add_run("Touchpoint extensions:")
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = DARK_GRAY

    for label, example in entry["touchpoints"].items():
        label_display = label.replace("_", " ").title()
        bp = doc.add_paragraph(style="List Bullet")
        run_l = bp.add_run(f"{label_display}: ")
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_l.font.color.rgb = DARK_GRAY
        run_e = bp.add_run(example)
        run_e.font.size = Pt(10)
        run_e.font.color.rgb = MED_GRAY
        run_e.italic = True

    doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════════════════
# 5. STANDALONE VS. "LUCID +" ANALYSIS
# ═══════════════════════════════════════════════════════════════════════

add_styled_heading(doc, "5. Standalone vs. 'Lucid +' Analysis")

add_body(doc,
    "Each candidate works in two formats: as a standalone brand or paired with 'Lucid.' "
    "The right choice depends on long-term brand architecture — whether Lucid remains the "
    "parent brand or the product eventually stands alone."
)

# Build comparison table
standalone_headers = ["Name", "Standalone Format", "Combined Format", "Stronger As"]
standalone_rows = []
for entry in NAMES:
    stronger = "Standalone" if entry["name"] == "Lucid" else (
        "Either" if entry["name"] == "Latent" else "Combined"
    )
    standalone_rows.append((
        entry["name"],
        entry["standalone_analysis"],
        entry["combined_analysis"],
        stronger,
    ))

# Use a custom table for better formatting
sa_table = doc.add_table(rows=1 + len(standalone_rows), cols=4)
sa_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sa_table.style = "Table Grid"

for i, h in enumerate(standalone_headers):
    cell = sa_table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, "4A6FA5")

for r_idx, row_data in enumerate(standalone_rows):
    for c_idx, val in enumerate(row_data):
        cell = sa_table.rows[r_idx + 1].cells[c_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_GRAY
        if c_idx == 3:  # "Stronger As" column
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if r_idx % 2 == 1:
            set_cell_shading(cell, "F5F7FA")

# Set column widths
widths = [2.5, 5.5, 5.5, 2.5]
for i, w in enumerate(widths):
    for row in sa_table.rows:
        row.cells[i].width = Cm(w)

doc.add_paragraph()

add_body(doc,
    "Key insight: If you plan to keep the Lucid parent brand long-term, Tonic is the strongest "
    "pairing. If the product may eventually stand alone (like how 'Instagram' dropped 'Burbn'), "
    "Lucid has the most independent brand equity.",
    italic=True,
)

# ═══════════════════════════════════════════════════════════════════════
# 6. SUMMARY RANKING TABLE
# ═══════════════════════════════════════════════════════════════════════

add_styled_heading(doc, "6. Summary Ranking")

# V2 candidates
summary_rows = []
for entry in NAMES:
    scores = entry["scores"]
    weighted = calc_weighted(scores)
    summary_rows.append((
        str(entry["rank"]),
        entry["name"],
        entry["tagline"],
        str(scores["catch"]),
        str(scores["accuracy"]),
        str(scores["persona"]),
        str(scores["brand"]),
        str(scores["diff"]),
        f"{weighted:.2f}",
    ))

summary_headers = ["#", "Name", "Tagline", "Cat.", "Acc.", "Per.", "Brd.", "Dif.", "Total"]
num_v2 = len(summary_rows)

# Add V1 reference rows
v1_rows = []
for i, ref in enumerate(V1_REFERENCE):
    scores = ref["scores"]
    weighted = calc_weighted(scores)
    v1_rows.append((
        f"V1-{i+1}",
        ref["name"],
        "(V1 reference)",
        str(scores["catch"]),
        str(scores["accuracy"]),
        str(scores["persona"]),
        str(scores["brand"]),
        str(scores["diff"]),
        f"{weighted:.2f}",
    ))

all_rows = summary_rows + v1_rows

summary_table = doc.add_table(rows=1 + len(all_rows), cols=len(summary_headers))
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
summary_table.style = "Table Grid"

for i, h_text in enumerate(summary_headers):
    cell = summary_table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(h_text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, "4A6FA5")

for r_idx, row_data in enumerate(all_rows):
    is_v1 = r_idx >= num_v2
    for c_idx, val in enumerate(row_data):
        cell = summary_table.rows[r_idx + 1].cells[c_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if is_v1:
            # Grayed out V1 reference rows
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.italic = True
            set_cell_shading(cell, "F0F0F0")
        else:
            run.font.color.rgb = DARK_GRAY
            # Highlight V2 rows in green
            set_cell_shading(cell, "E8F5E9")

doc.add_paragraph()

add_body(doc,
    "V1 top-3 included for reference (grayed). All three V2 candidates outscore the previous "
    "best (Lucid Signal at 8.65) by 0.35\u20130.50 points.",
    italic=True,
)

# ═══════════════════════════════════════════════════════════════════════
# 7. COMPETITIVE NAMING LANDSCAPE
# ═══════════════════════════════════════════════════════════════════════

add_styled_heading(doc, "7. Competitive Naming Landscape")

add_body(doc,
    "The table below shows how leading wellness tech and voice AI companies name their products, "
    "with positioning notes for where the V2 candidates would sit in this landscape."
)

make_table(doc,
    ["Brand", "Category", "Naming Strategy"],
    COMPETITORS,
    col_widths=[3.5, 3, 11.5],
)

doc.add_paragraph()

# Positioning annotations
add_body(doc,
    "Where V2 candidates fit in this landscape:",
    bold=True,
)

positions = [
    ("Lucid", "Sits alongside Calm and Headspace in the 'desired mental state' category. "
              "Same naming strategy as the two most successful wellness apps ever built. "
              "No direct collision \u2014 'lucid' occupies the clarity/awareness niche, not relaxation."),
    ("Latent", "Occupies a unique position: technical-evocative. Closest analog is Hume AI "
               "(intellectual heritage) but with more mystery. Appeals to the same audience as "
               "Levels (data-literate) but with a more poetic register."),
    ("Tonic", "Bridges the Lumen/Oura space (scientific-premium) with the Calm/Headspace space "
              "(wellness-aspirational). The music-theory connection gives it cultural depth that "
              "pure wellness names lack."),
]

for name, pos in positions:
    p = doc.add_paragraph()
    run_n = p.add_run(f"{name}: ")
    run_n.bold = True
    run_n.font.size = Pt(11)
    run_n.font.color.rgb = STEEL_BLUE
    run_p = p.add_run(pos)
    run_p.font.size = Pt(11)
    run_p.font.color.rgb = DARK_GRAY

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════
# 8. RECOMMENDATION (Decision Tree)
# ═══════════════════════════════════════════════════════════════════════

add_styled_heading(doc, "8. Recommendation")

add_body(doc,
    "All three candidates clear the 9.0 bar \u2014 any would be a strong choice. "
    "The right pick depends on your strategic priorities:"
)

doc.add_paragraph()

# Decision tree entries
decisions = [
    ("Choose LUCID if...",
     "you want maximum standalone brand power and the strongest emotional trigger. "
     "Lucid is the name that makes people feel something just by hearing it. Best for: "
     "an independent brand identity that doesn't need 'Lucid' as a crutch. "
     "Risk tolerance: you're comfortable with the lucid-dreaming association. "
     "Brand architecture: the product may eventually outgrow the Lucid family.",
     "9.15"),
    ("Choose LATENT if...",
     "accuracy and intellectual appeal are your top priorities. 'Latent' is the most "
     "technically precise name \u2014 it IS what the product detects. Best for: a technically "
     "sophisticated audience (ML engineers, biohackers, data scientists). "
     "Risk tolerance: you're comfortable with a name that some mainstream users won't "
     "immediately understand. Brand architecture: works well in either format.",
     "9.05"),
    ("Choose TONIC if...",
     "you want the strongest 'Lucid +' pairing and the most balanced name with no weak "
     "dimensions. The music-theory connection makes 'Lucid Tonic' feel like it was always "
     "meant to be. Best for: keeping Lucid as the parent brand long-term. "
     "Risk tolerance: lowest \u2014 Tonic has no controversial associations and all-9s consistency. "
     "Brand architecture: committed to the Lucid family.",
     "9.00"),
]

for title_text, detail, score in decisions:
    p_title = doc.add_paragraph()
    run_t = p_title.add_run(title_text)
    run_t.bold = True
    run_t.font.size = Pt(13)
    run_t.font.color.rgb = STEEL_BLUE

    p_score = doc.add_paragraph()
    run_sc = p_score.add_run(f"Score: {score}")
    run_sc.font.size = Pt(11)
    run_sc.font.color.rgb = MED_GRAY
    run_sc.italic = True

    p_detail = doc.add_paragraph()
    run_d = p_detail.add_run(detail)
    run_d.font.size = Pt(11)
    run_d.font.color.rgb = DARK_GRAY

    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════
# 9. VERIFICATION CHECKLIST
# ═══════════════════════════════════════════════════════════════════════

add_styled_heading(doc, "9. Verification Checklist")

add_body(doc,
    "Before committing to a name, complete the following due diligence for your chosen candidate:"
)

checklist_items = [
    ("USPTO Trademark Search",
     "Search TESS (tess2.uspto.gov) for the candidate name in International Classes 9 (software), "
     "10 (medical devices), and 42 (SaaS). Check for live registrations, pending applications, "
     "and abandoned marks that could be revived. Budget: $0 (self-service) or $500\u2013$1,500 "
     "for a professional comprehensive search."),
    ("Domain Availability",
     "Check .com, .app, .health, and .io TLDs. Exact-match .com is ideal but not required \u2014 "
     "Oura uses ouraring.com, Levels uses levelshealth.com. Fallback patterns: "
     "get[name].com, [name]app.com, try[name].com, [name]wellness.com."),
    ("App Store Namespace",
     "Search the iOS App Store and Google Play for the candidate name. Check if the name is "
     "taken as an app title (not just mentioned in descriptions). Verify the developer account "
     "can claim the name. Note: App Store allows duplicate names across different developers."),
    ("Social Media Handles",
     "Check availability on Instagram, Twitter/X, LinkedIn, and TikTok. Use namechk.com or "
     "similar tools for batch checking. Consistent handles across platforms are ideal but "
     "variations (e.g., @get_lucid, @lucid_app) are acceptable."),
    ("Informal User Testing",
     "Test the name with 5\u201310 people from each target persona. Ask: (1) What do you think "
     "this product does? (2) Would you download an app called [name]? (3) How does the name "
     "make you feel? (4) Can you spell it after hearing it once? Look for consistent positive "
     "signals across both personas, not unanimous agreement."),
    ("International Connotation Check",
     "Verify the candidate name doesn't have negative or embarrassing meanings in major "
     "languages (Spanish, French, German, Mandarin, Japanese). Use native speakers or a "
     "professional linguistic screening service. This is especially important for 'Latent' "
     "(check connotations of 'hidden/dormant' across cultures)."),
]

for i, (item_title, item_detail) in enumerate(checklist_items, 1):
    p = doc.add_paragraph()
    run_num = p.add_run(f"{i}. {item_title}  ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_num.font.color.rgb = DARK_GRAY
    run_det = p.add_run(item_detail)
    run_det.font.size = Pt(11)
    run_det.font.color.rgb = MED_GRAY

doc.add_paragraph()

add_body(doc,
    "Recommendation: Complete items 1\u20134 before announcing the name publicly. "
    "Items 5\u20136 can run in parallel and inform final selection if multiple candidates "
    "pass trademark clearance.",
    italic=True,
)

# ═══════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════

output_dir = "/Users/zacharypoll/Desktop/Documents/Claude Code/Lucid-Steel/Business Documents"
output_path = os.path.join(output_dir, "Lucid Naming Analysis V2.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")

# Verify scores match expected values
print("\nScore verification:")
for entry in NAMES:
    weighted = calc_weighted(entry["scores"])
    print(f"  {entry['name']}: {weighted:.2f}")
