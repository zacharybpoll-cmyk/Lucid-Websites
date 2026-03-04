#!/usr/bin/env python3
"""Generate Lucid YC CEO Advisory Document.

Advisory written as if from the CEO of Y Combinator to the founder of Lucid
after acceptance into the next batch.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Style Configuration ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for level, (size, color) in enumerate([
    (Pt(28), RGBColor(0x1B, 0x1B, 0x1B)),  # Heading 1
    (Pt(18), RGBColor(0x2C, 0x2C, 0x2C)),  # Heading 2
    (Pt(14), RGBColor(0x44, 0x44, 0x44)),  # Heading 3
], start=1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = size
    h.font.color.rgb = color
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(6)

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ── Helper Functions ─────────────────────────────────────────────────

def add_callout_box(doc, title, text, color_hex="E8F4FD", border_hex="2196F3"):
    """Add a shaded callout box using a single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    # Background shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

    # Border
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:color="{border_hex}"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(borders)

    # Title
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1B, 0x1B, 0x1B)

    # Body text
    for line in text.strip().split('\n'):
        p2 = cell.add_paragraph()
        run2 = p2.add_run(line.strip())
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()  # spacing


def add_metric_box(doc, metrics):
    """Add a metrics table with key/value pairs."""
    table = doc.add_table(rows=1, cols=len(metrics))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(metrics):
        cell = table.cell(0, i)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
        cell._tc.get_or_add_tcPr().append(shading)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_val = p.add_run(value + '\n')
        run_val.bold = True
        run_val.font.size = Pt(16)
        run_val.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
        run_lbl = p.add_run(label)
        run_lbl.font.size = Pt(9)
        run_lbl.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()


def add_bullet(doc, text, bold_prefix=None, indent_level=0):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + indent_level * 1.0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text)
    else:
        run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_section_header(doc, number, title):
    """Add a styled section header with page break."""
    doc.add_page_break()
    doc.add_heading(f'{number}. {title}', level=1)


def add_styled_table(doc, headers, rows, header_color="E0E0E0"):
    """Add a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx, col_idx).text = value

    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('LUCID')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1B, 0x1B, 0x1B)
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('YC S26 CEO Advisory')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'Calibri'

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('From Clinical Validation to Market Traction')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)

doc.add_paragraph()
doc.add_paragraph()

# Divider line
divider = doc.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = divider.add_run('\u2501' * 40)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
run.font.size = Pt(12)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Prepared for Batch S26\nConfidential \u2014 For Internal Use Only')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.italic = True

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════

doc.add_heading('Contents', level=1)
doc.add_paragraph()

toc_items = [
    ('1.', 'Executive Assessment'),
    ('2.', 'Mission-Critical Priorities (RIGHT NOW)'),
    ('3.', 'Getting to Market: The Launch Playbook'),
    ('4.', 'First Customers \u2014 Where They Are'),
    ('5.', 'Marketing Channels Ranked by ROI'),
    ('6.', 'Strengths & Weaknesses: The Honest Audit'),
    ('7.', 'D2C vs B2B vs Therapists: The Distribution Decision'),
    ('8.', 'Target Market Ranking & Quantification'),
    ('9.', 'Product Improvements Needed'),
    ('10.', 'Strategic Direction: 12-Month Vision'),
    ('11.', '90-Day Execution Timeline'),
    ('12.', 'Mistakes to Avoid'),
]

for num, title_text in toc_items:
    p = doc.add_paragraph()
    run_num = p.add_run(f'{num}  ')
    run_num.bold = True
    run_num.font.size = Pt(12)
    run_title = p.add_run(title_text)
    run_title.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE ASSESSMENT
# ══════════════════════════════════════════════════════════════════════

doc.add_heading('1. Executive Assessment', level=1)

doc.add_paragraph(
    'I\'m going to be direct with you, the way I would be with any founder who just '
    'got into the batch. You have built something extraordinary. A 244-million-parameter '
    'clinical model, validated in a 14,898-person peer-reviewed study, published in the '
    'Annals of Family Medicine, running entirely on-device on a consumer laptop. That is '
    'genuinely rare. Most companies in this batch will never achieve that level of '
    'scientific credibility.'
)

doc.add_paragraph(
    'And yet: you have zero customers. Zero revenue. An unsigned app that macOS '
    'Gatekeeper blocks on download. A 2.3GB binary that most people will never finish '
    'downloading. No Stripe integration, so even if someone wanted to pay you, they '
    'couldn\'t. No analytics, so you have no idea what happens after someone installs.'
)

doc.add_paragraph(
    'This is both the best and worst position to be in. Best, because the hard part \u2014 '
    'building technology worth paying for \u2014 is done. Worst, because the thing that kills '
    'startups isn\'t bad technology. It\'s the inability to get that technology into '
    'people\'s hands. You have a $30 million model sitting behind a door that no one '
    'can open.'
)

doc.add_paragraph(
    'Let\'s fix that.'
)

doc.add_paragraph()

add_metric_box(doc, [
    ('Paying Customers', '0'),
    ('Clinical Validation', 'Peer-Reviewed'),
    ('Test Coverage', '264 Tests'),
    ('Price Point', '$14.99/mo'),
])

doc.add_paragraph()

add_callout_box(doc,
    'The One Question That Matters',
    'Every decision for the next 90 days should be filtered through this:\n'
    '"Does this get the app into one more person\'s hands?"\n'
    '\n'
    'Not "does this make the product better." Not "does this impress investors."\n'
    'Does this get one more person to download, install, open, and use Lucid?\n'
    'If the answer is no, it can wait.',
)


# ══════════════════════════════════════════════════════════════════════
# SECTION 2: MISSION-CRITICAL PRIORITIES
# ══════════════════════════════════════════════════════════════════════

add_section_header(doc, 2, 'Mission-Critical Priorities (RIGHT NOW)')

doc.add_paragraph(
    'Before you write another line of feature code, before you think about marketing, '
    'before you book a single podcast appearance \u2014 these three things must happen. '
    'They are not suggestions. They are prerequisites for everything else in this document '
    'to work.'
)

doc.add_heading('Priority #1: Apple Developer Enrollment + Notarization', level=2)

doc.add_paragraph(
    'This is the single highest-leverage action you can take this week. Right now, '
    'when someone downloads Lucid, macOS throws up a warning that the app is '
    'from an "unidentified developer" and blocks it from opening. The user has to '
    'right-click, select Open, confirm a second dialog, and possibly go into System '
    'Settings to approve it. Every single step loses 40-60% of users. Your install '
    'completion rate is probably 5-10%.'
)

doc.add_paragraph(
    'This is a silent killer. You will never see the users you lose here. They won\'t '
    'email you. They won\'t file a bug. They\'ll just close the dialog and move on '
    'with their day. You are invisible to the 90% of people who tried and failed.'
)

notarize_steps = [
    ('Enroll in Apple Developer Program ($99/year). ', 'Do this today. Not tomorrow. Today. Enrollment can take 24-48 hours to process.'),
    ('Code-sign the app with your Developer ID certificate. ', 'This eliminates the Gatekeeper warning entirely. The app opens like any other macOS application.'),
    ('Submit for notarization via xcrun notarytool. ', 'Apple scans for malware and approves within minutes. This is automated, not a manual review.'),
    ('Staple the notarization ticket. ', 'Run xcrun stapler staple "Lucid.app". Now the app installs cleanly even offline.'),
]

for bold, rest in notarize_steps:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

add_callout_box(doc,
    'Timeline: 3-5 Days',
    'Day 1: Submit Apple Developer enrollment + payment\n'
    'Day 2-3: Wait for Apple approval\n'
    'Day 3-4: Set up code signing certificates, integrate into build pipeline\n'
    'Day 4-5: Submit for notarization, staple, test on a fresh Mac\n'
    '\n'
    'After this, your install completion rate goes from ~5% to ~70%.\n'
    'That is a 14x improvement in your entire funnel, for $99 and 5 days of work.',
    'E8F5E9', '4CAF50'
)

doc.add_heading('Priority #2: Stripe Payment Integration', level=2)

doc.add_paragraph(
    'You have a price on your website: $14.99/month, $149/year. But there is no way '
    'to actually pay you. No checkout flow, no payment link, no billing system. This '
    'means you could have a hundred people desperate to subscribe and you\'d have no '
    'mechanism to accept their money.'
)

stripe_steps = [
    ('Set up Stripe account + Checkout. ', 'Use Stripe Checkout (hosted payment page). Do not build a custom billing UI. Stripe handles tax, invoicing, receipts, and subscription management. Integration time: 4-6 hours.'),
    ('Create two price points. ', '$14.99/month and $149/year. Match what\'s on the website exactly. No additional tiers, no enterprise pricing yet.'),
    ('Add a subscription gate in the app. ', '30-day free trial, then prompt for payment. Use Stripe Customer Portal for users to manage their subscription. Keep it simple.'),
    ('Send the first invoice manually if you have to. ', 'If Stripe integration takes longer than expected, create a Stripe Payment Link (takes 5 minutes) and email it to your first users. Do not let "billing isn\'t ready" stop you from charging.'),
]

for bold, rest in stripe_steps:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

doc.add_heading('Priority #3: Analytics Instrumentation', level=2)

doc.add_paragraph(
    'You are flying blind. You have no idea how many people have installed the app, '
    'how often they open it, which features they use, or where they drop off. Without '
    'analytics, every decision you make is a guess.'
)

analytics_steps = [
    ('Install PostHog or Mixpanel. ', 'PostHog is free up to 1M events/month and self-hostable (good for privacy positioning). Mixpanel is easier to set up. Pick one and integrate it in a day.'),
    ('Track 5 critical events. ', 'App opened, first voice analysis completed, dashboard viewed, settings changed, app closed. That\'s it. Don\'t over-instrument.'),
    ('Respect your privacy promise. ', 'No voice data, no audio, no PII. Track behavioral events only. Be explicit in your privacy policy about what you track. Users who care about privacy will read it.'),
    ('Set up a daily check habit. ', 'Look at your analytics dashboard every morning before you write code. Let the data tell you what to build.'),
]

for bold, rest in analytics_steps:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

add_callout_box(doc,
    'Non-Negotiable',
    'These three priorities are not features. They are infrastructure.\n'
    'Without notarization, nobody can install.\n'
    'Without Stripe, nobody can pay.\n'
    'Without analytics, you can\'t learn.\n'
    '\n'
    'Estimated time: 10-14 days to complete all three.\n'
    'Nothing else in this document matters until these are done.',
    'FFF3E0', 'FF9800'
)


# ══════════════════════════════════════════════════════════════════════
# SECTION 3: GETTING TO MARKET
# ══════════════════════════════════════════════════════════════════════

add_section_header(doc, 3, 'Getting to Market: The Launch Playbook')

doc.add_paragraph(
    'Once the infrastructure is in place (notarized app, Stripe, analytics), '
    'here\'s how to sequence your launch. This is not a single event \u2014 it\'s a '
    'four-phase campaign designed to build momentum progressively.'
)

doc.add_heading('Phase 1: Soft Launch (Weeks 1-4)', level=2)
doc.add_paragraph('Goal: 25 hand-selected users, intense feedback loop')

soft_launch = [
    ('Hand-pick 25 users from your personal network. ', 'Not random people. Specifically: quantified-self enthusiasts who own Oura/Whoop, knowledge workers in 4+ hours of meetings/day, people who\'ve expressed interest in voice/mental health tech. You want users who will actually use it daily and give you real feedback.'),
    ('Co-install sessions. ', '15-minute Zoom calls where you walk each person through download, install, and first voice reading. Yes, for all 25 of them. This is where you learn what\'s confusing, what breaks, what delights.'),
    ('Weekly check-in cadence. ', 'Text or Slack each user every Monday: "How was Lucid last week? Anything surprising?" Ship fixes within 24 hours of any feedback. Speed of iteration is your competitive advantage.'),
    ('Conversion ask at Week 4. ', '"I\'m launching paid in two weeks. $14.99/month. Would you pay? Why or why not?" Their answers shape everything that follows.'),
]

for bold, rest in soft_launch:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

doc.add_heading('Phase 2: Community Launch (Weeks 4-8)', level=2)
doc.add_paragraph('Goal: 100-200 signups from targeted communities')

community_launch = [
    ('Show HN on Hacker News. ', 'Title: "Show HN: I put a $30M clinical voice AI model on your Mac for $15/month." Post on a Tuesday or Wednesday, 9-10am ET. Be in the comments responding to every question for 6 hours straight. HN loves privacy-first, on-device, scientifically-backed tools.'),
    ('Reddit deep dives. ', 'r/QuantifiedSelf (85K members), r/Biohackers (400K), r/productivity (2.5M). Don\'t just post a link \u2014 write a 500-word story about why you built this, what the clinical data shows, and what you\'ve learned from your first 25 users. Be genuinely helpful, not promotional.'),
    ('Twitter/X thread. ', 'Write a thread breaking down the Kintsugi DAM research: what it is, why it matters, how the 14,898-person study was designed. End with a soft CTA. Science threads outperform product threads 3:1 on engagement.'),
]

for bold, rest in community_launch:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

doc.add_heading('Phase 3: Content-Led Growth (Weeks 6-12)', level=2)
doc.add_paragraph('Goal: Own the "voice biomarkers" search category')

content_growth = [
    ('Publish 2 articles per week. ', 'Topics: "What Are Voice Biomarkers?", "How Your Voice Reveals Stress Before You Feel It", "PHQ-9 vs. Voice Analysis: What the Research Shows", "Why On-Device AI Matters for Mental Health Privacy." Target long-tail SEO terms that nobody is competing for yet.'),
    ('Launch a weekly newsletter. ', '"The Voice Biomarker Report" \u2014 5-minute read covering research papers, clinical studies, and what you\'re learning from your user data (aggregated and anonymized). Position yourself as the authority in this category.'),
    ('Guest on 3-5 podcasts. ', 'Target: health tech podcasts, quantified self shows, founder podcasts. Your story is inherently interesting: "$30M of clinical AI research, now running on a laptop for $15/month." Every podcast host will want that narrative.'),
]

for bold, rest in content_growth:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

doc.add_heading('Phase 4: Paid Acquisition Test (Weeks 8-12)', level=2)
doc.add_paragraph('Goal: Validate unit economics with small budget ($2,000-5,000)')

paid_acq = [
    ('Google Ads on exact-match keywords. ', '"voice stress analysis," "mental health monitoring app," "voice biomarker app." Start with $50/day and measure cost-per-install and cost-per-trial-start.'),
    ('Reddit promoted posts. ', 'Target r/QuantifiedSelf and r/Biohackers with promoted posts. Reddit ads are cheap ($2-5 CPM) and the audience is pre-qualified.'),
    ('Retargeting only after organic. ', 'Don\'t run cold traffic ads until you have at least 500 website visitors from organic sources. Retarget those visitors with a "come back and try it free for 30 days" message.'),
]

for bold, rest in paid_acq:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

add_callout_box(doc,
    'Phase Gates',
    'Do not advance to the next phase until the current one is working.\n'
    'Phase 1 to 2: At least 15 of 25 users active after Week 3, 3+ willing to pay.\n'
    'Phase 2 to 3: At least 50 organic signups, install completion rate above 40%.\n'
    'Phase 3 to 4: At least 100 signups, 10+ paying customers, LTV/CAC math pencils out.\n'
    '\n'
    'If you get stuck at a gate, the answer is not to skip ahead.\n'
    'The answer is to figure out why the current phase isn\'t working.',
    'E8F4FD', '2196F3'
)


# ══════════════════════════════════════════════════════════════════════
# SECTION 4: FIRST CUSTOMERS
# ══════════════════════════════════════════════════════════════════════

add_section_header(doc, 4, 'First Customers \u2014 Where They Are')

doc.add_paragraph(
    'Your first 10 paying customers will not come from a landing page. They will come '
    'from you, personally, finding them and convincing them one at a time. Here are the '
    'channels, ranked by how quickly they will produce your first user.'
)

doc.add_paragraph()

add_styled_table(doc,
    ['Rank', 'Channel', 'Reach', 'Difficulty', 'Conversion', 'Time to First User'],
    [
        ['1', 'Personal Network', 'Low (50-100)', 'Very Easy', 'High (20-30%)', '1-3 days'],
        ['2', 'r/QuantifiedSelf', 'Medium (85K)', 'Easy', 'Medium (3-5%)', '1-2 weeks'],
        ['3', 'Hacker News (Show HN)', 'High (500K+)', 'Medium', 'Low-Med (1-2%)', '2-4 weeks'],
        ['4', 'Podcaster Outreach', 'Variable', 'Medium', 'Medium (5-10%)', '3-6 weeks'],
        ['5', 'LinkedIn Content', 'Medium', 'Medium', 'Low (1-2%)', '2-4 weeks'],
        ['6', 'Therapist Referral', 'Low (local)', 'Hard', 'Medium (5-10%)', '6-10 weeks'],
        ['7', 'Paid Ads', 'High', 'Easy', 'Very Low (<1%)', '4-8 weeks'],
    ]
)

doc.add_heading('Channel Details', level=2)

doc.add_heading('1. Personal Network (Start Here)', level=3)
doc.add_paragraph(
    'You know people who are in meetings all day. You know people who track their '
    'sleep with Oura and their workouts with Whoop. You know people who are curious '
    'about health tech. Text them. Not a mass email \u2014 individual texts. "Hey, I '
    'built something I think you\'d find interesting. Can I show you in 15 minutes?" '
    'Your first 5 users should come from people who already trust you.'
)

doc.add_heading('2. r/QuantifiedSelf', level=3)
doc.add_paragraph(
    'This subreddit is your spiritual home. These are people who wear multiple '
    'trackers, export their data to spreadsheets, and argue about HRV algorithms. '
    'They are the exact early adopters who will forgive a rough onboarding experience '
    'in exchange for a genuinely novel data stream. Write a thoughtful post, not a '
    'product announcement. Share your data, your methodology, your clinical validation. '
    'They will respect the science and try the product.'
)

doc.add_heading('3. Hacker News (Show HN)', level=3)
doc.add_paragraph(
    'HN is your highest-leverage single event. A front-page Show HN can drive '
    '5,000-15,000 visits in 24 hours. But it needs to be genuine. Lead with the '
    'technology story: a 244M-parameter model, peer-reviewed in the Annals of Family '
    'Medicine, running on-device with zero cloud dependency. HN loves technical depth, '
    'privacy-first architecture, and founders who respond to every comment. If you '
    'nail this, it\'s 2-4 weeks of organic traffic.'
)

doc.add_heading('4. Podcaster Outreach', level=3)
doc.add_paragraph(
    'Your story is inherently podcast-friendly: "$30 million of clinical R&D, '
    'validated on 14,898 people, now running on any Mac for $15/month." Target health '
    'tech podcasts (Digital Health Today, Health Unchained), quantified self shows '
    '(Quantified Body, Body of Knowledge), and founder-journey podcasts. '
    'Cold email 20 hosts. Expect 3-5 to say yes. Each appearance brings 50-200 '
    'qualified visitors.'
)

doc.add_heading('5. LinkedIn Content', level=3)
doc.add_paragraph(
    'LinkedIn works for reaching knowledge workers and HR leaders, but only if you '
    'post consistently. Write about meeting fatigue, burnout research, and voice '
    'biomarker science. Don\'t pitch the product \u2014 build authority. When you do mention '
    'Lucid, it should feel like a natural extension of your expertise, not a sales pitch.'
)

doc.add_heading('6. Therapist Referral', level=3)
doc.add_paragraph(
    'This channel is high-conversion but slow to build. Therapists who use objective '
    'measures (PHQ-9, GAD-7) between sessions will immediately see the value of '
    'continuous voice biomarker data. But earning therapist trust takes time. Start '
    'with 3-5 local therapists, offer free accounts for their clients, and let the '
    'clinical data speak for itself. This becomes powerful at scale but is a Q3/Q4 play.'
)

doc.add_heading('7. Paid Ads', level=3)
doc.add_paragraph(
    'Last because it should be last. Paid acquisition before you have a working '
    'funnel, proven onboarding, and clear conversion metrics is burning money. Use '
    'paid ads to amplify what\'s already working, not to find product-market fit. '
    'Budget $2,000-5,000 for initial tests only after organic channels produce 100+ '
    'installs.'
)

doc.add_paragraph()

add_callout_box(doc,
    'Outreach Script Template',
    'Hey [Name],\n'
    '\n'
    'I\'m building a voice biomarker tool that detects stress and depression risk\n'
    'from the way you speak in meetings. It\'s based on a 14,898-person peer-reviewed\n'
    'study (published in the Annals of Family Medicine) and runs 100% on-device \u2014\n'
    'your voice never leaves your Mac.\n'
    '\n'
    'I\'m looking for 25 people to try it free for 30 days and give honest feedback.\n'
    'Would you be open to a 15-minute demo this week?\n'
    '\n'
    '\u2014 [Your name]',
    'F3E5F5', '9C27B0'
)


# ══════════════════════════════════════════════════════════════════════
# SECTION 5: MARKETING CHANNELS RANKED
# ══════════════════════════════════════════════════════════════════════

add_section_header(doc, 5, 'Marketing Channels Ranked by ROI')

doc.add_paragraph(
    'Not all channels are created equal for a pre-revenue, privacy-first voice '
    'biomarker tool. Here\'s how to allocate your time and money, ranked by expected '
    'return on investment for the first 6 months.'
)

doc.add_heading('Tier 1: High ROI, Low Cost (Do These First)', level=2)

doc.add_paragraph()

add_styled_table(doc,
    ['Channel', 'Cost', 'Expected Reach', 'Why It Works'],
    [
        ['Hacker News (Show HN)', '$0', '5,000-15,000 visits', 'Technical audience values privacy, on-device AI, clinical rigor. Perfect fit.'],
        ['Twitter/X Science Threads', '$0', '500-5,000 per thread', 'Voice biomarker research threads get high engagement. Build authority first, convert later.'],
        ['Reddit (QS, Biohackers)', '$0', '200-2,000 per post', 'Engaged niche communities. Authenticity > promotion. Share data, not ads.'],
    ]
)

doc.add_heading('Tier 2: Medium ROI, Medium Cost (Weeks 4-8)', level=2)

doc.add_paragraph()

add_styled_table(doc,
    ['Channel', 'Cost', 'Expected Reach', 'Why It Works'],
    [
        ['SEO / Blog Content', '$0 (time)', '100-500/mo growing', 'Own "voice biomarkers" category. No competition yet. Compounds over time.'],
        ['YouTube Explainers', '$0-500', '500-5,000 per video', 'Demo the product, explain the science. Video builds trust faster than text.'],
        ['Podcast Appearances', '$0', '500-2,000 per episode', 'Your story sells itself. "$30M model on a laptop." Every host wants this angle.'],
    ]
)

doc.add_heading('Tier 3: Lower ROI, Higher Cost (Weeks 8+)', level=2)

doc.add_paragraph()

add_styled_table(doc,
    ['Channel', 'Cost', 'Expected Reach', 'Why It Works (When Ready)'],
    [
        ['Google Ads (Exact Match)', '$2,000-5,000/mo', 'Variable', 'Only after organic validates demand. Target "voice stress app" and similar.'],
        ['Health Tech Influencers', '$500-2,000/post', '5,000-50,000', 'Only after product is notarized and onboarding is polished. Influencer traffic is harsh on rough edges.'],
        ['Product Hunt Launch', '$0', '2,000-10,000 visits', 'One-shot event. Save for when app is notarized, Stripe works, and onboarding is tight.'],
    ]
)

doc.add_paragraph()

add_callout_box(doc,
    'The 80/20 Rule for Marketing',
    'In the first 90 days, 80% of your marketing energy should go into Tier 1.\n'
    'These channels are free, targeted, and favor exactly what you have:\n'
    'deep technical substance and genuine scientific credibility.\n'
    '\n'
    'Don\'t spend money on ads until you know that a user who installs\n'
    'the app actually uses it for 7+ days. Until then, every dollar\n'
    'spent on acquisition is a dollar spent on people who churn.',
    'E8F4FD', '2196F3'
)


# ══════════════════════════════════════════════════════════════════════
# SECTION 6: STRENGTHS & WEAKNESSES
# ══════════════════════════════════════════════════════════════════════

add_section_header(doc, 6, 'Strengths & Weaknesses: The Honest Audit')

doc.add_paragraph(
    'Every YC partner will ask you: "What are your biggest weaknesses?" If you can '
    'answer this honestly and specifically, with a plan for each one, you demonstrate '
    'self-awareness \u2014 which is more important than having no weaknesses.'
)

doc.add_heading('Strengths (What\'s Working)', level=2)

strengths = [
    ('$30M in R&D you didn\'t have to fund. ',
     'The Kintsugi DAM model represents $30 million in clinical research, 7 years of development, '
     '863 hours of labeled clinical speech, and training on 35,000 patients across multiple clinical '
     'sites. You are distributing the output of this investment for $14.99/month. Your marginal cost '
     'of serving each user is effectively zero \u2014 the model runs on their hardware. This is one of '
     'the most asymmetric cost structures I\'ve seen.'),
    ('Peer-reviewed clinical validation. ',
     'Published in the Annals of Family Medicine (January 2025), one of the top primary care journals '
     'in the world. 14,898 adults in a double-blind, multi-site study. 71.3% sensitivity, 73.6% '
     'specificity for major depressive disorder from 25 seconds of natural speech. This isn\'t a '
     'preprint or a conference poster \u2014 it\'s the highest standard of evidence in medical research. '
     'No competitor in consumer mental health has anything close to this.'),
    ('100% on-device privacy architecture. ',
     'Every analysis runs locally on the user\'s Mac. Voice never leaves the device. Only numerical '
     'biomarker scores are stored. In a world where people are increasingly anxious about AI '
     'companies harvesting their data, "works with your internet disconnected" is verifiable and '
     'powerful. This is not just a marketing claim \u2014 it\'s an architectural decision that builds '
     'genuine trust.'),
    ('Polished, feature-rich product. ',
     'Wellness Score for daily wellness tracking, Correlation Explorer for connecting voice data with '
     'Oura/Whoop/Apple Health, Rhythm Rings for circadian patterns, Compass for directional trends, '
     'Echoes for historical voice signatures, Burnout Risk assessment, Grove for ecological metaphor '
     'tracking. 264 passing tests. This is not an MVP \u2014 it\'s a mature product that happens to '
     'have zero users.'),
    ('v10 website with strong clinical positioning. ',
     'Your website leads with the science, not the product. "$30 million in research. 14,898 adults. '
     'Published in the Annals of Family Medicine." This is exactly the right positioning for a tool '
     'that needs to overcome skepticism about "AI wellness apps." The trust strip ($30M, 35,000 '
     'patients, 71.3% sensitivity, Annals of Family Medicine) does more work than any feature list.'),
    ('Well-anchored pricing. ',
     '$14.99/month is brilliantly positioned: expensive enough to signal quality, cheap enough that '
     'it\'s a rounding error next to what your target customer spends on Oura ($6/mo), Whoop '
     '($30/mo), or therapy ($400-800/mo). The "save 97% vs. therapy" comparison is aggressive but '
     'defensible. No tiers, no upsells, full access \u2014 simple pricing reduces friction.'),
]

for bold, rest in strengths:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

doc.add_heading('Weaknesses (What Needs Fixing)', level=2)

weaknesses = [
    ('Zero customers, zero revenue. ',
     'This is the elephant in the room. All the clinical validation in the world means nothing until '
     'someone pays you $14.99. Until you have paying customers, you don\'t have a business \u2014 you '
     'have a research project. Every day without a customer is a day where "will people pay for this?" '
     'remains an open question.'),
    ('Unsigned app / Gatekeeper blocks installation. ',
     'This is actively preventing adoption. Non-technical users cannot install an unsigned macOS app. '
     'They don\'t know how to bypass Gatekeeper, and they shouldn\'t have to. Until the app is '
     'notarized, you are limited to users technical enough to right-click and override security '
     'warnings. That\'s maybe 20% of your target market.'),
    ('2.3GB download size. ',
     'The app is 2.3GB because it bundles a 244M-parameter ML model plus Python runtime. On a fast '
     'connection, that\'s a 2-minute download. On a coffee shop Wi-Fi, it\'s 10-15 minutes. Every '
     'minute of download time is a minute where the user can change their mind. For comparison, '
     'Spotify is 200MB. Slack is 300MB.'),
    ('macOS-only limits addressable market. ',
     'macOS has roughly 15-17% desktop market share globally, higher among your target personas '
     '(knowledge workers, creatives). But mobile is where people spend 4+ hours/day, and you don\'t '
     'have a mobile app. This is acceptable for the batch but becomes a hard constraint when you '
     'try to raise a Series A.'),
    ('No established acquisition channel. ',
     'You have a polished website but no funnel, no content engine, no referral mechanism, no '
     'outbound sales motion. Users don\'t find you yet. No one is searching for "voice biomarker '
     'app" (monthly search volume: basically zero). You need to create the demand, not capture '
     'existing demand.'),
    ('Regulatory gray zone. ',
     'Lucid is positioned as a "wellness tool, not a medical device," similar to how Oura provides '
     'HRV data. But you\'re detecting depression risk with a clinically validated model. If the FDA '
     'decides voice biomarker tools that reference clinical studies need clearance, you have a '
     'compliance problem. This isn\'t urgent, but it\'s a background risk that investors will ask about.'),
    ('Model dependency on Kintsugi. ',
     'The DAM model is the core of your product, and you didn\'t build it. If Kintsugi changes '
     'licensing terms, restricts usage, or goes out of business, your product is at risk. This is '
     'a strategic dependency that needs a long-term mitigation plan (either secure perpetual license '
     'terms or begin building proprietary models from your own user data).'),
]

for bold, rest in weaknesses:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

add_callout_box(doc,
    'Investor Framing',
    'When a YC partner asks about weaknesses, lead with:\n'
    '"Our technology is validated. Our product is built. Our only weakness is distribution,\n'
    'and that\'s what the next 90 days are about."\n'
    '\n'
    'This is honest, specific, and shows you know what matters.\n'
    'Don\'t hide the weaknesses \u2014 frame them as solvable problems with clear timelines.',
    'E8F5E9', '4CAF50'
)


# ══════════════════════════════════════════════════════════════════════
# SECTION 7: D2C VS B2B VS THERAPISTS
# ══════════════════════════════════════════════════════════════════════

add_section_header(doc, 7, 'D2C vs B2B vs Therapists: The Distribution Decision')

doc.add_paragraph(
    'This is the most important strategic decision you\'ll make this batch. Every '
    'mentor will have an opinion. Half will say "go enterprise." Half will say "go '
    'consumer." Here\'s the analysis, with numbers.'
)

doc.add_heading('Option A: Direct-to-Consumer (D2C)', level=2)

doc.add_paragraph()

add_styled_table(doc,
    ['Dimension', 'Assessment'],
    [
        ['Time to first dollar', '2-4 weeks (after Stripe integration)'],
        ['Price point', '$14.99/mo or $149/yr'],
        ['Sales cycle', 'Zero. Self-serve. Download, try, pay.'],
        ['CAC estimate', '$20-80 per paying user (blended organic + paid)'],
        ['LTV estimate', '$90-180 (6-12 month retention at $14.99/mo)'],
        ['Scalability', 'High. No sales team needed. Viral potential.'],
        ['Key risk', 'Consumer churn. Mental health apps average 30-day retention of 4-8%.'],
    ]
)

doc.add_paragraph(
    'Pros: Fastest path to revenue. Shortest feedback loops. You own the relationship. '
    'Every user teaches you something. You can iterate daily based on real usage data.'
)

doc.add_paragraph(
    'Cons: Consumer mental health has brutal retention curves. Users download wellness '
    'apps with good intentions and stop using them within a month. You need to fight this '
    'with habit loops, engagement features (you have many \u2014 Wellness Score, Grove, streaks), '
    'and genuine ongoing value delivery.'
)

doc.add_heading('Option B: Business-to-Business (B2B)', level=2)

doc.add_paragraph()

add_styled_table(doc,
    ['Dimension', 'Assessment'],
    [
        ['Time to first dollar', '3-6 months (pilot negotiation + procurement)'],
        ['Price point', '$15-30/user/mo (wellness budget allocation)'],
        ['Sales cycle', '4-12 weeks. Requires champion, budget approval, IT review.'],
        ['CAC estimate', '$2,000-10,000 per account (founder-led sales)'],
        ['LTV estimate', '$5,000-30,000 per account (20-100 seats, 12+ months)'],
        ['Scalability', 'Medium. Requires sales team eventually. Slower growth.'],
        ['Key risk', 'Procurement cycles. IT security reviews. "We need SOC 2."'],
    ]
)

doc.add_paragraph(
    'Pros: Higher contract values. Lower churn (annual contracts). IT can push the app '
    'to 50+ Macs at once (solves distribution). Wellness budgets at mid-market companies '
    'are $200-500/employee/year \u2014 Lucid at $180/year fits neatly.'
)

doc.add_paragraph(
    'Cons: Slow. You won\'t close your first B2B deal for 3-6 months. Every deal requires '
    'security review, legal review, and a champion who will push it through procurement. '
    'You need SOC 2 compliance (or at least a credible plan for it). And you\'ll need to '
    'build admin features: team dashboard, aggregated (anonymized) insights, billing per seat.'
)

doc.add_heading('Option C: Therapist / Clinician Channel', level=2)

doc.add_paragraph()

add_styled_table(doc,
    ['Dimension', 'Assessment'],
    [
        ['Time to first dollar', '6-12 weeks (trust-building phase)'],
        ['Price point', '$14.99/mo per client (therapist gets it free)'],
        ['Sales cycle', '2-6 weeks per therapist, ongoing client referrals.'],
        ['CAC estimate', '$100-300 per therapist (who refers 5-15 clients)'],
        ['LTV estimate', '$50-150 per client (shorter engagement windows)'],
        ['Scalability', 'Medium. Each therapist is a distribution node. Network effects.'],
        ['Key risk', 'Regulatory sensitivity. Therapists are risk-averse about tools that touch diagnosis.'],
    ]
)

doc.add_paragraph(
    'Pros: Therapists are trusted recommenders. A therapist who tells 10 clients "try this '
    'between sessions" converts at 40-60%. Therapist endorsement also gives you clinical '
    'credibility that\'s hard to buy. And the clinical validation (peer-reviewed, PHQ-9 '
    'comparison) is exactly what therapists need to feel comfortable recommending.'
)

doc.add_paragraph(
    'Cons: Therapists are conservative. They won\'t recommend an unsigned app. They need '
    'HIPAA-adjacent assurances (even though Lucid processes locally). And each therapist '
    'is a slow, relationship-based sale. This channel takes 3-6 months to build but compounds '
    'beautifully once you have 20-30 therapists recommending you.'
)

doc.add_heading('The Recommendation', level=2)

doc.add_paragraph()

add_callout_box(doc,
    'Strategic Recommendation: D2C First, B2B in Parallel, Therapists in Q3',
    'Start D2C immediately. It\'s the fastest path to your first dollar and your\n'
    'first real feedback. Your $14.99 price point, self-serve model, and free trial\n'
    'are already set up for D2C. You just need Stripe.\n'
    '\n'
    'Plant B2B seeds in parallel. Start conversations with 3-5 HR/People Ops leaders\n'
    'at Mac-heavy companies (tech startups, design agencies, media companies). Share the\n'
    'clinical research. Offer a free 10-seat pilot. These conversations take months to\n'
    'close, so start them now even though revenue won\'t come until Q3.\n'
    '\n'
    'Therapist channel is a Q3/Q4 play. Don\'t invest here until you have a notarized\n'
    'app, proven retention, and at least 100 consumer users. Therapists need to see that\n'
    'real people use and benefit from the tool before they\'ll recommend it to clients.',
    'FFF3E0', 'FF9800'
)


# ══════════════════════════════════════════════════════════════════════
# SECTION 8: TARGET MARKET RANKING
# ══════════════════════════════════════════════════════════════════════

add_section_header(doc, 8, 'Target Market Ranking & Quantification')

doc.add_paragraph(
    'Let\'s be precise about who is going to pay you $14.99/month, how many of them '
    'exist, and what that means for revenue. These are conservative estimates designed '
    'to be defensible in front of YC partners, not impressive on a pitch deck.'
)

doc.add_heading('Tier 1: Primary Markets (Pursue Now)', level=2)

doc.add_heading('Quantified Self Professionals', level=3)
doc.add_paragraph(
    'Who: Adults 27-45, already wearing Oura Ring and/or Whoop, actively tracking sleep, '
    'HRV, and activity data. They see a gap in their health stack \u2014 physical wellness is '
    'measured, mental wellness is not. They want the "mental health layer" for their existing '
    'data ecosystem.'
)

add_styled_table(doc,
    ['Metric', 'Estimate', 'Basis'],
    [
        ['TAM (Global QS Market)', '2.5M people', 'Oura has sold 2.5M+ rings; ~60% active trackers'],
        ['SAM (Mac Users, US/EU)', '200,000', '~40% Mac usage in QS demographic, US/EU focus'],
        ['SOM (Year 1 Realistic)', '2,000 users', '1% of SAM with organic + community marketing'],
        ['Revenue (Year 1)', '$360,000 ARR', '2,000 users x $14.99/mo x 12'],
    ]
)

doc.add_paragraph(
    'Why they convert: They already spend $30+/month on health tracking (Oura $6/mo, '
    'Whoop $30/mo, supplements, blood panels). $14.99 for the missing mental health data '
    'is a trivial add. They understand biomarkers. They trust peer-reviewed research. They '
    'will read the Annals of Family Medicine paper.'
)

doc.add_heading('Meeting-Burned Knowledge Workers', level=3)
doc.add_paragraph(
    'Who: Senior ICs, managers, and executives (28-45) at tech companies, spending 4-7 '
    'hours/day in video calls. They know meetings drain them but can\'t quantify it. They '
    'want objective data to advocate for schedule changes, meeting-free days, or role '
    'transitions.'
)

add_styled_table(doc,
    ['Metric', 'Estimate', 'Basis'],
    [
        ['TAM (Knowledge Workers)', '50M+ globally', 'Remote/hybrid workers in meetings 4+hrs/day'],
        ['SAM (Mac, Aware, WTP)', '500,000', 'Mac users who self-identify as meeting-burned, WTP $15/mo'],
        ['SOM (Year 1 Realistic)', '1,000 users', '0.2% of SAM. Harder to reach than QS segment.'],
        ['Revenue (Year 1)', '$180,000 ARR', '1,000 users x $14.99/mo x 12'],
    ]
)

doc.add_paragraph(
    'Why they convert: Lucid answers a question they ask every day: "Why am I so exhausted '
    'after that meeting?" The Wellness Score and stress/calm tracking provide objective evidence '
    'that was previously invisible. Calendar integration (when shipped) makes this a must-have.'
)

doc.add_heading('Tier 2: Growth Markets (Pursue in Months 4-6)', level=2)

tier2_items = [
    ('Remote workers with burnout anxiety. ',
     'Broader than "meeting-burned" \u2014 includes freelancers, contractors, remote employees who '
     'feel isolated and want to monitor their mental state. Larger market but lower pain '
     'intensity. SOM: 500-1,000 users in Year 1.'),
    ('Mental health-aware professionals. ',
     'People who\'ve been in therapy, use meditation apps (Calm, Headspace), or track mood '
     'manually in journals. They value mental health data but don\'t have an objective source. '
     'SOM: 500-1,000 users in Year 1.'),
]

for bold, rest in tier2_items:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

doc.add_heading('Tier 3: Future Markets (Months 7-12)', level=2)

tier3_items = [
    ('B2B team wellness. ',
     'HR and People Ops teams deploying Lucid across 20-100 seat teams for aggregate '
     'wellness insights. Higher ACV ($3,600-18,000/year per account) but longer sales cycle. '
     'Requires admin dashboard, team analytics, and SOC 2.'),
    ('Therapist-assisted monitoring. ',
     'Therapists recommending Lucid to clients for between-session monitoring. Each therapist '
     'is a distribution node for 5-15 clients. Requires clinical-specific dashboards and HIPAA '
     'considerations.'),
    ('Graduate students and researchers. ',
     'PhD students in psychology, neuroscience, and public health who want to use voice '
     'biomarker data in their research. Low revenue per user but high advocacy value. Offer '
     'academic pricing ($4.99/mo) in exchange for case studies and publications.'),
]

for bold, rest in tier3_items:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

doc.add_heading('Year 1 Revenue Summary', level=2)

add_styled_table(doc,
    ['Segment', 'SOM (Users)', 'Monthly Revenue', 'Annual Revenue'],
    [
        ['Quantified Self Professionals', '2,000', '$30,000', '$360,000'],
        ['Meeting-Burned Knowledge Workers', '1,000', '$15,000', '$180,000'],
        ['Remote Workers (Tier 2)', '500', '$7,500', '$90,000'],
        ['Mental Health-Aware (Tier 2)', '500', '$7,500', '$90,000'],
        ['Total (Conservative)', '4,000', '$60,000', '$720,000'],
    ]
)

add_callout_box(doc,
    'Reality Check',
    'These numbers assume 50% annual churn and exclude B2B entirely.\n'
    'If B2B closes even one 50-seat pilot at $20/user/mo, that\'s $12,000/year\n'
    'from a single account \u2014 equivalent to 67 consumer subscribers.\n'
    '\n'
    'The Year 1 target of ~$720K ARR is ambitious but achievable\n'
    'with strong execution on distribution. The model isn\'t the bottleneck.\n'
    'Distribution is.',
    'E8F4FD', '2196F3'
)


# ══════════════════════════════════════════════════════════════════════
# SECTION 9: PRODUCT IMPROVEMENTS NEEDED
# ══════════════════════════════════════════════════════════════════════

add_section_header(doc, 9, 'Product Improvements Needed')

doc.add_paragraph(
    'Your product is feature-rich but missing foundational infrastructure. Here\'s every '
    'improvement needed, ranked by impact on conversion and retention. Build in this order \u2014 '
    'resist the temptation to work on fun features before the boring ones.'
)

doc.add_heading('Critical (Blocks Revenue)', level=2)

critical_items = [
    ('Apple Notarization. ',
     'Impact: 10-14x increase in install completion rate. Without this, nothing else matters. '
     'Timeline: 3-5 days. See Section 2.'),
    ('Stripe Integration. ',
     'Impact: Enables revenue. Currently impossible for anyone to pay you. '
     'Timeline: 1-2 days for basic Checkout integration.'),
    ('Analytics (PostHog/Mixpanel). ',
     'Impact: Enables data-driven decisions. You cannot improve what you cannot measure. '
     'Timeline: 1 day for basic event tracking.'),
    ('Onboarding Flow. ',
     'Impact: First-time user experience determines 7-day retention. Right now, a new user '
     'opens the app and sees... what? They need a guided 3-step setup: create voice profile, '
     'make first recording, see first insight. Timeline: 3-5 days.'),
]

for bold, rest in critical_items:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

doc.add_heading('High Priority (Improves Retention)', level=2)

high_items = [
    ('App size reduction. ',
     'From 2.3GB to under 800MB. Investigate model quantization (INT8), strip unused Python '
     'packages, compress assets. Each GB reduction improves download completion by ~15%.'),
    ('Calendar integration. ',
     'Google Calendar first (85%+ of target market uses Google Workspace). Show which meetings '
     'cause stress spikes. This is the "aha moment" feature \u2014 the one that makes users say '
     '"I can\'t go back to not having this."'),
    ('Email/account system. ',
     'Users need accounts for subscription management, cross-device sync (eventually), and '
     'password recovery. Use a simple email + magic link system. Don\'t build auth from scratch \u2014 '
     'use Clerk, Auth0, or Supabase Auth.'),
    ('Auto-update mechanism. ',
     'Users need to receive updates without re-downloading 2.3GB. Implement Electron\'s '
     'autoUpdater with a GitHub Releases backend. Without this, every bug fix requires users '
     'to manually download the full app again.'),
]

for bold, rest in high_items:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

doc.add_heading('Medium Priority (Expands Market)', level=2)

medium_items = [
    ('Health app integrations (Oura, Whoop, Apple Health). ',
     'Your Correlation Explorer already has the UI for this. Shipping actual integrations '
     'unlocks the quantified self market segment \u2014 your highest-value early adopters.'),
    ('Referral mechanism. ',
     '"Give a friend 30 days free, get 30 days free." Simple viral loop. Implement after you '
     'have 50+ users and proven 30-day retention. Before that, referrals just accelerate churn.'),
    ('Team plan / admin dashboard. ',
     'For B2B: aggregate (never individual) wellness metrics across a team. "Your engineering '
     'team\'s collective stress is 23% higher this week than last." Requires careful privacy '
     'design. Don\'t build until you have a signed B2B pilot.'),
]

for bold, rest in medium_items:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

doc.add_heading('Low Priority / Defer', level=2)

defer_items = [
    ('Mobile app (iOS/Android). ',
     'Do not build this in the batch. macOS is your beachhead. Mobile voice analysis has '
     'different constraints (battery, background processing, microphone access). This is a '
     'Series A play, not a batch play.'),
    ('Windows / Linux support. ',
     'macOS market share in your target demographic is sufficient for Year 1. Cross-platform '
     'expands TAM but splits engineering effort. Defer until you have 1,000+ Mac users.'),
    ('Custom ML models (proprietary). ',
     'Building your own voice biomarker model from user data is the long-term moat play. But '
     'it requires 10,000+ hours of labeled speech data that you don\'t have yet. Start '
     'collecting (with consent) now. Build later.'),
    ('Therapist dashboard. ',
     'A clinician-facing view showing client voice biomarker trends between sessions. High '
     'value but requires HIPAA compliance, clinical workflows, and therapist onboarding. '
     'This is a Q4 or Year 2 feature.'),
]

for bold, rest in defer_items:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

add_callout_box(doc,
    'Build Order Discipline',
    'Builders love building features. You have to resist.\n'
    '\n'
    'The Critical items unlock revenue. Build them first.\n'
    'The High items improve retention. Build them second.\n'
    'The Medium items expand market. Build them third.\n'
    'The Low items are distractions disguised as opportunities. Defer them.\n'
    '\n'
    'If you catch yourself working on a Medium item while Critical items remain,\n'
    'stop and ask: "Am I doing the thing that matters most right now?"',
    'FFF3E0', 'FF9800'
)


# ══════════════════════════════════════════════════════════════════════
# SECTION 10: STRATEGIC DIRECTION
# ══════════════════════════════════════════════════════════════════════

add_section_header(doc, 10, 'Strategic Direction: 12-Month Vision')

doc.add_paragraph(
    'Here\'s what the next 12 months look like if you execute well. These milestones '
    'are aggressive but achievable. Each quarter builds on the previous one \u2014 don\'t '
    'skip ahead.'
)

doc.add_heading('Months 1-3: Foundation (Batch)', level=2)
doc.add_paragraph('Theme: Get to 100 paying customers and $1,500 MRR')

m1_3 = [
    ('Week 1-2: ', 'Ship notarization, Stripe, and analytics. These are the three '
     'locks on the door. Remove them.'),
    ('Week 3-4: ', 'Soft launch to 25 hand-selected users. Co-install sessions. '
     'Intensive feedback loop. Ship bug fixes daily.'),
    ('Week 5-6: ', 'Show HN + Reddit launch. Target 500+ signups. Convert 10-20 '
     'to paying. Build onboarding flow based on soft launch learnings.'),
    ('Week 7-8: ', 'First content pieces published. Newsletter launched. Podcast '
     'outreach begins. Target 50 paying customers.'),
    ('Week 9-12: ', 'Iterate on retention. Ship calendar integration. Start B2B '
     'conversations. Target 100 paying customers, $1,500 MRR by Demo Day.'),
]

for bold, rest in m1_3:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

add_metric_box(doc, [
    ('Users', '100'),
    ('MRR', '$1,500'),
    ('Retention (30d)', '>25%'),
    ('NPS', '>40'),
])

doc.add_heading('Months 4-6: Growth', level=2)
doc.add_paragraph('Theme: 500 paying users, $7,500 MRR, first integrations')

m4_6 = [
    ('Ship Oura and Whoop integrations. ', 'Unlock the quantified self market. '
     'This segment has the highest willingness to pay and the strongest word-of-mouth.'),
    ('Launch referral program. ', '"Give a friend 30 days free." Target 20% of new '
     'users coming from referrals.'),
    ('Close first B2B pilot. ', 'One 20-50 seat team at a Mac-heavy company. Use this '
     'to validate team analytics features and B2B pricing.'),
    ('Double content output. ', '4 articles/week, weekly newsletter, 2 podcast appearances '
     '/month. Own the "voice biomarker" search category.'),
    ('Begin seed raise preparation. ', 'Build the deck, compile metrics, create a data '
     'room. You\'ll need: MRR trajectory, retention curves, NPS, and 3-5 user testimonials.'),
]

for bold, rest in m4_6:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

add_metric_box(doc, [
    ('Users', '500'),
    ('MRR', '$7,500'),
    ('B2B Pilots', '1-2'),
    ('Content Pieces', '40+'),
])

doc.add_heading('Months 7-9: Scale', level=2)
doc.add_paragraph('Theme: 2,000 paying users, $30,000 MRR, mobile beta, seed raise')

m7_9 = [
    ('Cross $30K MRR. ', 'This is the inflection point where VCs take you seriously. '
     '$360K ARR run rate, growing 20-30% month-over-month.'),
    ('Begin mobile beta (iOS). ', 'Start with a companion app that shows your existing data, '
     'not a standalone mobile product. Add mobile voice analysis later after validating '
     'the iOS audio pipeline.'),
    ('Raise seed round ($2-4M). ', 'At $30K MRR with strong growth, you\'ll have leverage. '
     'Target: $2-4M at $15-25M post-money valuation. Use funds for mobile engineering, '
     'B2B sales hire, and content marketing.'),
    ('Expand therapist channel. ', 'With 2,000 consumer users as social proof, approach '
     'therapists with data: "X% of our users report improved self-awareness of mood '
     'patterns." Therapists need evidence, and now you have it.'),
]

for bold, rest in m7_9:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

add_metric_box(doc, [
    ('Users', '2,000'),
    ('MRR', '$30,000'),
    ('ARR Run Rate', '$360K'),
    ('Seed Raised', '$2-4M'),
])

doc.add_heading('Months 10-12: Category Leadership', level=2)
doc.add_paragraph('Theme: 5,000+ paying users, $75,000 MRR, define the category')

m10_12 = [
    ('Reach 5,000 paying users. ', 'At this scale, your user data becomes a strategic '
     'asset. Aggregated, anonymized voice biomarker trends across thousands of users is '
     'a dataset that doesn\'t exist anywhere else.'),
    ('Launch "Voice Biomarker Index." ', 'A public, anonymized weekly report on collective '
     'stress and mood trends across your user base. This is content marketing, brand '
     'building, and data moat all in one.'),
    ('Publish first proprietary research. ', 'With 5,000+ users and months of data, you '
     'can publish your own findings on voice biomarker patterns. This transitions you from '
     '"distributing Kintsugi\'s research" to "generating your own."'),
    ('Begin international expansion. ', 'UK, Canada, Australia \u2014 English-speaking markets '
     'with strong health tech adoption and Mac penetration. Same product, different marketing.'),
]

for bold, rest in m10_12:
    add_bullet(doc, rest, bold_prefix=bold)

doc.add_paragraph()

add_metric_box(doc, [
    ('Users', '5,000+'),
    ('MRR', '$75,000'),
    ('ARR Run Rate', '$900K'),
    ('Markets', '4 Countries'),
])


# ══════════════════════════════════════════════════════════════════════
# SECTION 11: 90-DAY TIMELINE
# ══════════════════════════════════════════════════════════════════════

add_section_header(doc, 11, '90-Day Execution Timeline')

doc.add_paragraph(
    'Here is your week-by-week execution plan for the first 90 days of the batch. '
    'Print this out. Put it on your wall. Check items off as you go. If you fall behind '
    'on a week, don\'t try to catch up by working on two weeks simultaneously \u2014 adjust '
    'the plan and keep moving forward.'
)

doc.add_paragraph()

timeline_headers = ['Week', 'Focus Area', 'Key Deliverables', 'Success Metric']

timeline_rows = [
    ['1', 'Infrastructure',
     'Apple Developer enrollment submitted. Stripe account created. PostHog/Mixpanel integrated.',
     'All 3 initiated or completed'],
    ['2', 'Infrastructure',
     'App notarized and signed. Stripe Checkout integrated. Analytics tracking 5 core events.',
     'Signed app installs cleanly on fresh Mac'],
    ['3', 'Soft Launch',
     'First 10 users identified and contacted. Co-install sessions scheduled.',
     '10 users with app installed'],
    ['4', 'Soft Launch',
     'All 25 soft launch users onboarded. First feedback collected. Quick fixes shipped.',
     '25 active users, 5+ daily active'],
    ['5', 'Community Launch',
     'Show HN post published. r/QuantifiedSelf post published. Twitter/X thread posted.',
     '500+ website visits, 50+ downloads'],
    ['6', 'Community Launch',
     'Respond to all community feedback. Ship most-requested fix/feature. Second Reddit post.',
     '100+ total installs, 30+ active users'],
    ['7', 'Conversion',
     'First conversion ask to soft launch users. Stripe payment links sent. Track conversion.',
     '5+ paying customers'],
    ['8', 'Conversion + Content',
     'First 2 blog articles published. Newsletter launched. Podcast outreach (20 cold emails).',
     '15+ paying customers, 2 articles live'],
    ['9', 'Growth',
     'Onboarding flow shipped. Calendar integration MVP started. Referral mechanism designed.',
     '30+ paying customers, onboarding live'],
    ['10', 'Growth',
     'Calendar integration beta. 2 more articles published. First podcast appearance booked.',
     '50+ paying customers'],
    ['11', 'Scale Prep',
     'Calendar integration launched. B2B outreach begins (5 companies contacted). Auto-update shipped.',
     '70+ paying customers, 1 B2B meeting'],
    ['12', 'Scale Prep',
     'Size reduction shipped (<1GB). App store / landing page optimization based on analytics data.',
     '85+ paying customers'],
    ['13', 'Demo Day Prep',
     'Demo Day narrative finalized. Metrics compiled. 100-customer target push. Practice pitch.',
     '100 paying customers, pitch ready'],
]

add_styled_table(doc, timeline_headers, timeline_rows, header_color="E0E0E0")

doc.add_paragraph()

add_callout_box(doc,
    'Timeline Flexibility',
    'This timeline is a framework, not a contract. Some weeks will go faster,\n'
    'some will hit unexpected blockers. The order matters more than the timing.\n'
    '\n'
    'The iron rule: do not skip the infrastructure phase (Weeks 1-2).\n'
    'Everything downstream depends on a signed app, working payments, and analytics.\n'
    'If these take 3 weeks instead of 2, that\'s fine. If you skip them to\n'
    '"move faster," you\'ll spend months paying for it.',
    'E8F4FD', '2196F3'
)


# ══════════════════════════════════════════════════════════════════════
# SECTION 12: MISTAKES TO AVOID
# ══════════════════════════════════════════════════════════════════════

add_section_header(doc, 12, 'Mistakes to Avoid')

doc.add_paragraph(
    'I\'ve watched hundreds of companies go through the batch. These are the mistakes '
    'that kill companies with exactly your profile: strong technology, zero distribution.'
)

doc.add_heading('Mistake #1: Building Features Before Notarizing', level=2)

doc.add_paragraph(
    'You will be tempted to add one more feature before dealing with the "boring" '
    'notarization and signing process. Don\'t. Every feature you build for an app that '
    'nobody can install is a feature built for an audience of zero. The most elegant '
    'feature in the world is worthless if macOS won\'t let users open the app. '
    'Notarize first. Features second.'
)

doc.add_heading('Mistake #2: Leading with "Mental Health" Instead of Clinical Authority', level=2)

doc.add_paragraph(
    'The consumer mental health app market is crowded, skeptical, and burned by years of '
    'unsubstantiated claims. If you position Lucid as "another mental health app," you\'re '
    'competing with Calm, Headspace, BetterHelp, and a hundred others. You lose that fight. '
    'Instead, lead with what makes you genuinely different: a peer-reviewed, clinically '
    'validated model with the strongest evidence base in consumer voice AI. You\'re not a '
    'wellness app. You\'re the output of $30 million in clinical research, now available '
    'on a laptop. That framing changes the conversation entirely.'
)

doc.add_heading('Mistake #3: Comparing to Therapy in Positioning', level=2)

doc.add_paragraph(
    'Your website says "Save 97% vs. therapy." This comparison is useful for anchoring '
    'price perception, but dangerous if it implies that Lucid replaces therapy. It doesn\'t, '
    'and you should never suggest it does. Therapists are potential allies and referral '
    'partners. If they perceive you as positioning against them, you lose the entire clinician '
    'channel. Better framing: "Lucid fills the gap between therapy sessions with objective, '
    'continuous data. It makes therapy more effective, not less necessary."'
)

doc.add_heading('Mistake #4: Trying B2B Before D2C Traction', level=2)

doc.add_paragraph(
    'B2B is seductive because the deal sizes are bigger. But pursuing enterprise deals '
    'without consumer traction is a trap. Enterprise buyers will ask: "How many individual '
    'users do you have?" If the answer is zero, the conversation is over. They need social '
    'proof. They need to know that real people use this tool daily and find it valuable. '
    'Get 100 D2C customers first. Then walk into B2B conversations with: "100 people pay '
    'us $14.99/month. Here\'s what they say about it." That changes everything.'
)

doc.add_heading('Mistake #5: Paid Acquisition Before the Funnel Works', level=2)

doc.add_paragraph(
    'If your install completion rate is 5% (unsigned app, 2.3GB download) and your 7-day '
    'retention is unknown (no analytics), spending money on Google Ads is pouring water '
    'into a bucket with holes. Fix the bucket first: notarize the app, reduce download size, '
    'build onboarding, instrument analytics, measure retention. Only spend money on acquisition '
    'when you know that a user who installs the app has a >20% chance of being active on '
    'Day 7. Otherwise, every ad dollar is subsidizing churn.'
)

doc.add_heading('Mistake #6: Going Mobile Too Early', level=2)

doc.add_paragraph(
    'Every mentor will tell you "you need a mobile app." They\'re right \u2014 eventually. But '
    'building iOS + Android before you have 1,000 Mac users is premature optimization. Mobile '
    'voice analysis has different constraints: battery drain, background processing limits, '
    'microphone permission UX, app store review (which will scrutinize health claims). Each '
    'of these is a multi-week engineering challenge. Meanwhile, your Mac app already works. '
    'Dominate Mac first. Prove the model. Then go mobile with funding, a dedicated mobile '
    'engineer, and proven user demand. If your users are begging for a mobile app, that\'s a '
    'great signal. If you\'re building it because investors asked \u2014 that\'s a mistake.'
)

doc.add_paragraph()

add_callout_box(doc,
    'The Meta-Mistake',
    'The common thread in all six mistakes is the same: doing the exciting thing\n'
    'instead of the important thing. Notarization is boring. Stripe is boring.\n'
    'Analytics is boring. Outreach to 50 people individually is boring.\n'
    '\n'
    'Building features is exciting. Designing a mobile app is exciting.\n'
    'Writing a B2B sales deck is exciting.\n'
    '\n'
    'The difference between companies that make it through the batch and\n'
    'companies that don\'t is the willingness to do the boring work first.\n'
    'Every. Single. Time.',
    'FFF3E0', 'FF9800'
)

# Final page
doc.add_page_break()

for _ in range(6):
    doc.add_paragraph()

closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run('You have built something worth selling.')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1B, 0x1B, 0x1B)

doc.add_paragraph()

closing2 = doc.add_paragraph()
closing2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing2.add_run('Now sell it.')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)

doc.add_paragraph()
doc.add_paragraph()

divider = doc.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = divider.add_run('\u2501' * 40)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
run.font.size = Pt(12)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    'Lucid \u2014 YC S26 CEO Advisory\n'
    'Confidential \u2014 For Internal Use Only\n'
    '\n'
    '$30M in research. 14,898 patients. 244M parameters.\n'
    'Now get it into people\'s hands.'
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.italic = True


# ══════════════════════════════════════════════════════════════════════
# SAVE DOCUMENT
# ══════════════════════════════════════════════════════════════════════

output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, 'Lucid_YC_CEO_Advisory.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.0f} KB')
